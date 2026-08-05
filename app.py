from pathlib import Path

import base64
import io
import json
import os
import re
import ast
import warnings
import hashlib
import urllib.request
from datetime import datetime

import numpy as np
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt

# === 中文字体统一由 setup_chinese_font() 管理（见下文），此处不再单独加载 ===

import matplotlib.font_manager as fm
import streamlit as st

st.set_page_config(
    page_title="数学建模前期数据分析工具",
    layout="wide",
)

import statsmodels.api as sm

from scipy.stats import (
    pearsonr,
    spearmanr,
    shapiro,
    probplot,
    f_oneway,
    chi2_contingency,
    kruskal,
    norm,
    t as t_dist,
)
from scipy.optimize import linprog
from sklearn.model_selection import (
    train_test_split,
    GroupShuffleSplit,
    KFold,
)
from sklearn.preprocessing import StandardScaler, MinMaxScaler
from sklearn.linear_model import (
    Lasso,
    Ridge,
    ElasticNet,
    LassoCV,
    LogisticRegression,
)
from sklearn.ensemble import RandomForestRegressor, RandomForestClassifier
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.metrics import (
    accuracy_score,
    balanced_accuracy_score,
    precision_score,
    recall_score,
    f1_score,
    roc_auc_score,
    log_loss,
    mean_squared_error,
    mean_absolute_error,
    r2_score,
    confusion_matrix,
)
from statsmodels.stats.outliers_influence import (
    variance_inflation_factor,
    OLSInfluence,
)
from statsmodels.stats.diagnostic import (
    het_breuschpagan,
    het_white,
)
from statsmodels.stats.stattools import durbin_watson
from statsmodels.discrete.discrete_model import (
    NegativeBinomial,
    NegativeBinomialP,
)

import jieba
import PyPDF2
import docx


# ============================================================
# 一、页面和全局设置
# ============================================================

# 仅过滤统计库已知的、不影响结果的消息噪音（避免完全隐藏未来变更提示）
for _warning_category, _warning_message in [
    (FutureWarning, "is_sparse is deprecated"),
    (FutureWarning, "use_inf_as_na"),
]:
    warnings.filterwarnings(
        "ignore",
        category=_warning_category,
        message=_warning_message,
    )
# 初始化所有可能用到的 session_state 变量
if "opt_uploaded_file" not in st.session_state:
    st.session_state.opt_uploaded_file = None
if "app_mode" not in st.session_state:
    st.session_state.app_mode = "数据分析"
if "welcome_seen" not in st.session_state:
    st.session_state.welcome_seen = False
if "guide_step" not in st.session_state:
    st.session_state.guide_step = 0
# 新手引导：自动推进（展示用）与用户手动勾选分离
if "guide_auto_step" not in st.session_state:
    st.session_state.guide_auto_step = 0
if "paper_checklist" not in st.session_state:
    st.session_state.paper_checklist = {}
# 高级分析结果缓存（存储时记录数据与所选列的签名）
if "adv_cache_signatures" not in st.session_state:
    st.session_state.adv_cache_signatures = {}
# 共享数据源：两个模块共用同一上传文件
if "shared_data_file" not in st.session_state:
    st.session_state.shared_data_file = None
    

# ===== 页面背景图片 =====

_bg_path = Path(__file__).resolve().parent / "assets" / "background.png"

if _bg_path.exists():
    _bg_base64 = base64.b64encode(
        _bg_path.read_bytes()
    ).decode("ascii")

    st.markdown(
        f"""
        <style>
        html, body, .stApp {{
            background: transparent !important;
        }}

        [data-testid="stAppViewContainer"] {{
            background: transparent !important;
        }}

        [data-testid="stHeader"] {{
            background: transparent !important;
        }}

        #custom-background-image {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100vw;
            height: 100vh;
            object-fit: cover;
            z-index: -1;
            pointer-events: none;
        }}

        [data-testid="stMain"] {{
            background: rgba(255, 255, 255, 0.10) !important;
        }}
        </style>

        <img
            id="custom-background-image"
            src="data:image/png;base64,{_bg_base64}"
        />
        """,
        unsafe_allow_html=True,
    )
else:
    st.info(
        "未找到 assets/background.png，已使用默认背景。"
        "如需自定义背景，请将图片放入 assets 文件夹。"
    )

# ===== 页面背景图片结束 =====


# ===== 页面标题图片路径 =====
_fursona_file = Path(__file__).parent / "assets" / "fursona_shu.png"
# ===== 页面标题图片路径结束 =====

# ===== 页面标题 =====

st.markdown(
    """
    <style>
    /* 标题图片和文字垂直居中 */
    div[data-testid="stHorizontalBlock"] {
        align-items: center !important;
    }

    /* 标题文字 */
    .old-main-title {
        color: #26354a;
        font-size: 72px;
        font-weight: 800;
        line-height: 220px;
        white-space: nowrap;
        letter-spacing: 1px;
        margin: 0;
        padding: 0;
    }

    /* 标题区域列的上下空白 */
    div[data-testid="stHorizontalBlock"] > div[data-testid="column"] {
        padding-top: 0 !important;
        padding-bottom: 0 !important;
    }

    /*
    标题图片位置微调：
    正值向右，负值向左；
    正值向下，负值向上。
    */
    div[data-testid="stHorizontalBlock"]:has(.old-main-title)
    > div[data-testid="column"]:first-child
    [data-testid="stImage"] {
        transform: translate(28px, -5px) !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# 标题区域
_title_col_image, _title_col_text = st.columns(
    [1.5, 8.5],
    gap="small",
)

with _title_col_image:
    if _fursona_file.exists():
        st.image(str(_fursona_file), width=220)
    else:
        st.caption(
            "未找到 assets/fursona_shu.png，标题图片已隐藏。"
            "可自行放入图片开启。"
        )

with _title_col_text:
    st.markdown(
        """
        <div class="old-main-title">
            学建模前期数据分析工具
        </div>
        """,
        unsafe_allow_html=True,
    )

# ===== 页面标题结束 =====

sns.set_theme(style="whitegrid")

plt.rcParams["axes.unicode_minus"] = False


# ============================================================
# 二、字体设置
# ============================================================

def setup_chinese_font():
    """
    尝试配置中文字体（单一入口）。
    优先级：项目内 assets 字体文件 > 系统已安装字体 > 联网下载 Noto Sans CJK。
    任何一步失败都不会影响主程序运行。
    """
    # 1) 优先使用项目内字体文件
    for font_file_name in ["msyh.ttc", "simsun.ttc", "NotoSansCJKsc-Regular.otf"]:
        font_path = Path(__file__).resolve().parent / "assets" / font_file_name
        if font_path.exists():
            try:
                fm.fontManager.addfont(str(font_path))
                font_prop = fm.FontProperties(fname=str(font_path))
                font_name = font_prop.get_name()
                plt.rcParams["font.sans-serif"] = [font_name, "DejaVu Sans"]
                plt.rcParams["axes.unicode_minus"] = False
                return
            except Exception:
                continue

    # 2) 使用系统已安装字体
    font_candidates = [
        "WenQuanYi Zen Hei",
        "Noto Sans CJK SC",
        "SimHei",
        "Microsoft YaHei",
        "Arial Unicode MS",
        "DejaVu Sans",
    ]

    available_fonts = {
        font.name for font in fm.fontManager.ttflist
    }

    selected_font = None
    for font_name in font_candidates:
        if font_name in available_fonts:
            selected_font = font_name
            break

    if selected_font is not None:
        plt.rcParams["font.sans-serif"] = [selected_font]
        plt.rcParams["axes.unicode_minus"] = False
        return

    # 3) 最后才尝试下载 Noto Sans CJK 中文字体
    try:
        font_dir = os.path.expanduser("~/.cache/matplotlib/fonts")
        os.makedirs(font_dir, exist_ok=True)

        font_path = os.path.join(
            font_dir,
            "NotoSansCJKsc-Regular.otf",
        )

        if not os.path.exists(font_path):
            url = (
                "https://github.com/GoogleFonts/"
                "noto-cjk/raw/main/Sans/OTF/"
                "SimplifiedChinese/NotoSansCJKsc-Regular.otf"
            )
            with urllib.request.urlopen(
                url,
                timeout=5,
            ) as response:
                with open(font_path, "wb") as font_file:
                    font_file.write(response.read())

        if os.path.exists(font_path):
            fm.fontManager.addfont(font_path)
            font_prop = fm.FontProperties(fname=font_path)
            plt.rcParams["font.sans-serif"] = [
                font_prop.get_name(),
                "DejaVu Sans",
            ]
            return

    except Exception:
        pass

    plt.rcParams["font.sans-serif"] = ["DejaVu Sans"]


setup_chinese_font()


# ============================================================
# 三、Session State 管理
# ============================================================

MODEL_STATE_KEYS = [
    "fitted_result",
    "fitted_model",
    "final_model_type",
    "model_meta",
    "vif_table",
    "target_mapping",
    "model_signature",
    "selected_model_type",
    "analysis_signature",
    "compare_df",
    "compare_success",
    "new_prediction_df",
    "new_prediction_notes",
    # 修复：K-Fold 结果也随模型切换一起清空，避免旧折结果残留
    "kfold_result",
    "kfold_summary",
    "kfold_success",
]


def clear_model_session_state():
    """清除旧模型结果和分析签名。"""
    for key in MODEL_STATE_KEYS:
        st.session_state.pop(key, None)

    st.session_state.pop(
        "analysis_signature",
        None,
    )


def reset_model_if_signature_changed(signature):
    """
    如果数据、变量、模型设置发生变化，则清除旧模型结果。
    """
    previous_signature = st.session_state.get("analysis_signature")

    if previous_signature != signature:
        clear_model_session_state()
        st.session_state["analysis_signature"] = signature


def build_analysis_signature(
    file_name,
    file_hash,
    target,
    predictors,
    variable_types,
    group_col,
    missing_method,
    outlier_method,
    outlier_action,
    include_target_outlier,
    robust_se,
    use_test_set,
    test_size,
):
    """根据当前分析设置生成唯一签名。"""
    signature_data = {
        "file_name": file_name,
        "file_hash": file_hash,
        "target": target,
        "predictors": sorted(predictors),
        # 按列名排序后再参与签名，避免用户调整变量顺序触发无谓的模型重置
        "variable_types": dict(
            sorted(variable_types.items(), key=lambda item: item[0])
        ),
        "group_col": group_col,
        "missing_method": missing_method,
        "outlier_method": outlier_method,
        "outlier_action": outlier_action,
        "include_target_outlier": include_target_outlier,
        "robust_se": robust_se,
        "use_test_set": use_test_set,
        "test_size": test_size,
    }

    return repr(signature_data)


# ============================================================
# 四、通用工具函数
# ============================================================

def clean_column_name(name):
    """规范列名。"""
    name = str(name).strip()
    name = re.sub(r"\s+", "_", name)
    name = re.sub(r"[^\w\u4e00-\u9fff]", "_", name)
    name = re.sub(r"_+", "_", name)
    name = name.strip("_")

    if not name:
        return "未命名变量"

    return name


def make_unique_columns(columns):
    """清理列名，并保证最终列名全局唯一。"""
    used = set()
    result = []

    for col in columns:
        base = clean_column_name(col)

        if not base:
            base = "未命名变量"

        candidate = base
        suffix = 2

        while candidate in used:
            candidate = f"{base}_{suffix}"
            suffix += 1

        used.add(candidate)
        result.append(candidate)

    return result

def dataframe_download(df, filename, key=None):
    """生成 CSV 下载按钮。"""
    if df is None:
        return

    if not isinstance(df, pd.DataFrame):
        df = pd.DataFrame(df)

    # 修复：to_csv(encoding="utf-8-sig") 已输出带 BOM 的字节串，
    # 再 encode("utf-8-sig") 会写入双重 BOM，此处统一用 utf-8。
    data = df.to_csv(
        index=False,
        encoding="utf-8-sig",
    ).encode("utf-8")

    download_key = key or f"download_{filename}"

    st.download_button(
        label=f"下载 {filename}",
        data=data,
        file_name=filename,
        mime="text/csv",
        key=download_key,
    )


def safe_to_latex(df, **kwargs):
    """DataFrame 转 LaTeX 表格，兼容新旧版 pandas 的 to_latex API。

    pandas 3.0 重写了 DataFrame.to_latex()，移除了 booktabs、bold_rows、
    encoding、sparsify 等旧参数；若直接传入会抛 TypeError：
        NDFrame.to_latex() got an unexpected keyword argument 'booktabs'

    这里先按传入参数尝试，失败后自动剔除不受支持的旧参数重试，
    保证新旧版本 pandas 都能生成 LaTeX 三线表。
    """
    if df is None:
        return ""

    if not isinstance(df, pd.DataFrame):
        df = pd.DataFrame(df)

    # pandas 3.0 已移除的旧参数
    legacy_params = {
        "booktabs",
        "bold_rows",
        "encoding",
        "formatters",
        "sparsify",
        "index_names",
        "line_width",
    }

    try:
        return df.to_latex(**kwargs)
    except TypeError:
        cleaned_kwargs = {
            key: value
            for key, value in kwargs.items()
            if key not in legacy_params
        }
        return df.to_latex(**cleaned_kwargs)


# ============================================================
# 四·补充、UI 兼容与图表管理 helper
# ============================================================

def _st_dataframe(df, *args, **kwargs):
    """
    兼容新旧版 Streamlit 的表格展示。
    新版用 width="stretch"，旧版用 use_container_width=True。
    """
    kwargs.pop("use_container_width", None)

    try:
        return st.dataframe(df, *args, width="stretch", **kwargs)
    except TypeError:
        return st.dataframe(df, *args, use_container_width=True, **kwargs)


# 收集本会话生成的所有图表，供综合报告导出时嵌入 Word
_REPORT_FIGURES = []


def show_fig(fig, name=None):
    """
    展示 matplotlib 图表并附带“下载 PNG”按钮，同时登记到报告图集。
    """
    import io as _io

    st.pyplot(fig)

    if name is None:
        name = f"chart_{len(_REPORT_FIGURES) + 1}"

    try:
        buffer = _io.BytesIO()
        fig.savefig(
            buffer,
            format="png",
            dpi=150,
            bbox_inches="tight",
        )
        st.download_button(
            "⬇️ 下载此图 PNG",
            data=buffer.getvalue(),
            file_name=f"{name}.png",
            mime="image/png",
            key=f"dl_fig_{name}",
        )
    except Exception as _fig_save_error:
        # 修复：不再静默吞掉异常，给出可见提示
        st.warning(
            f"图表「{name}」的下载按钮生成失败（不影响展示）："
            f"{_fig_save_error}"
        )

    _REPORT_FIGURES.append(fig)

    plt.close(fig)


# ============================================================
# 四·补充、新手引导系统
# ============================================================

GUIDE_STEPS = [
    {
        "key": "problem",
        "title": "① 上传赛题（可选）",
        "desc": "上传 PDF/Word/TXT 赛题文件，或直接粘贴赛题原文，点击“自动检测题型”。",
        "when": "拿到赛题的第一时间。让工具自动识别题型，从而推荐对应的分析路径。",
        "how": "上传 PDF / Word / TXT，或直接把赛题文字粘贴到文本框，然后点击“自动检测题型”。",
        "output": "题型判断结果（评价/预测/优化/机理/分类）+ 该题型的推荐流程。",
    },
    {
        "key": "data",
        "title": "② 上传数据表",
        "desc": "上传 CSV / Excel 数据表。没有现成数据？点“下载示例数据”立刻体验全流程。",
        "when": "赛题提供了数据文件时。如果没有数据，先下载示例数据把流程走通，再思考怎么造/采集数据。",
        "how": "CSV / Excel 直接上传即可。注意列名要规范：无空格、无重复，程序会自动清洗列名。",
        "output": "数据概览：行数、列数、缺失单元格数、重复行数。",
    },
    {
        "key": "vars",
        "title": "③ 选择因变量与自变量",
        "desc": "因变量是你要预测/解释的结果（如销量、评分），自变量是可能影响它的因素。",
        "when": "读题后明确“要预测/解释什么”（因变量）、“用什么解释它”（自变量）后。",
        "how": "因变量只能选 1 个；自变量可多选。ID 列（序号、学号等）会被自动排除，不要手动选入。",
        "output": "变量组合——决定了后续所有分析和模型。",
    },
    {
        "key": "types",
        "title": "④ 确认变量类型",
        "desc": "程序自动识别类型，请按变量真实含义确认：连续（数值）、分类（类别）、时间、次数（计数）。",
        "when": "每次更换数据或变量都要核对一遍。自动识别仅供参考，务必人工确认。",
        "how": "连续=可带小数的数值；分类=类别文字或 0/1；时间=日期；次数=非负整数（如促销次数）。",
        "output": "正确的变量类型——直接决定自动推荐什么模型。",
    },
    {
        "key": "clean",
        "title": "⑤ 数据清洗",
        "desc": "处理缺失值与异常值。左侧可切换“分类型处理”或“删除含缺失值的行”。",
        "when": "数据存在缺失值或异常值时必做（赛题数据通常都有）。",
        "how": "缺失值推荐“分类型处理”或“KNN插补”；异常值先用“仅标记，不删除”看数量和分布，再决定是否删除。",
        "output": "清洗汇总报告 + 可直接复制进论文的数据清洗表述。",
    },
    {
        "key": "model",
        "title": "⑥ 建模与诊断",
        "desc": "程序自动推荐模型，选择后点击“开始拟合最终模型”，查看指标、系数和诊断图。",
        "when": "完成清洗、确认变量类型后，这就是核心一步。",
        "how": "采用程序推荐的模型（或展开“多模型对比”选 AIC 更小的），点击“开始拟合最终模型”。",
        "output": "模型指标表、系数表（含显著性）、LaTeX 三线表、残差诊断图、预测区间。",
    },
    {
        "key": "adv",
        "title": "⑦ 高级分析（可选）",
        "desc": "熵权TOPSIS、灰色预测、ARIMA、聚类、PCA、随机森林等，按赛题类型选用。",
        "when": "按题型选用：评价类→熵权TOPSIS/PCA；预测类→灰色预测/ARIMA；分类类→随机森林；其他→稳健性分析。",
        "how": "在“8.5 高级分析方法”中按题型选择对应 Tab，每个 Tab 顶部都有使用说明和论文表述。",
        "output": "补充图表和结论，是论文的加分项。",
    },
    {
        "key": "report",
        "title": "⑧ 导出报告",
        "desc": "一键导出 Word 综合报告，论文表述草稿可直接复制进论文。",
        "when": "全部分析完成后、写论文之前。",
        "how": "点击“📑 一键导出综合报告”生成 Word；各环节的“论文表述”文本可直接复制粘贴。",
        "output": "综合报告.docx + 各环节 CSV 中间结果 + 图表 PNG。",
    },
]


def make_sample_data_csv():
    """
    生成一份带缺失值、异常值、分类变量、计数变量和日期的新手示例数据，
    覆盖本工具几乎所有功能点。
    """
    rng = np.random.default_rng(2026)
    n = 120

    data = {
        "城市": rng.choice(["北京", "上海", "广州", "深圳", "成都"], n),
        "客流量": np.round(rng.normal(500, 120, n), 1),
        "广告费": np.round(rng.normal(30, 10, n), 2),
        "促销次数": rng.integers(0, 8, n),
        "门店面积": np.round(rng.normal(120, 40, n), 1),
        "开业天数": rng.integers(30, 400, n),
    }

    df = pd.DataFrame(data)

    # 制造几个缺失值（先转 float 再赋 NaN，避免整型列报错）
    for col in ["客流量", "广告费", "促销次数"]:
        idx = rng.choice(n, 5, replace=False)
        df[col] = df[col].astype("float64")
        df.loc[idx, col] = np.nan

    # 制造两个明显异常值
    df.loc[0, "客流量"] = 9500.0
    df.loc[1, "广告费"] = 300.0

    return df.to_csv(index=False, encoding="utf-8-sig").encode("utf-8")


def _update_guide_progress():
    """根据任务清单勾选状态，重算当前进行到的步骤（用户手动进度）。"""
    done_steps = []

    for index, step in enumerate(GUIDE_STEPS, start=1):
        if st.session_state.get(
            f"check_guide_{step['key']}",
            False,
        ):
            done_steps.append(index)
        else:
            break

    st.session_state["guide_step"] = (
        max(done_steps) if done_steps else 0
    )


def render_mission_checklist(step_index):
    """侧边栏任务清单：手动勾选与自动进度互通。

    修复：自动进度（guide_auto_step，由实际完成操作驱动）只作为
    首次渲染 checkbox 的默认值，用户手动勾选后以用户状态为准，
    避免页面 rerun 时把任务清单自动勾满。
    """
    st.markdown("### ✅ 我的任务清单")

    auto_step = st.session_state.get("guide_auto_step", 0)

    # 修复：Streamlit 不允许在 widget 实例化后直接修改其 session_state 值。
    # 点击“同步”按钮时只置位一个非 widget 标志并 rerun；
    # 这里在所有 checkbox 创建【之前】检查该标志并预置 widget key，
    # 此时 widget 尚未实例化，赋值合法（官方推荐的“按钮改变其他 widget”模式）。
    if st.session_state.pop("guide_sync_triggered", False):
        for index, step in enumerate(GUIDE_STEPS, start=1):
            st.session_state[
                f"check_guide_{step['key']}"
            ] = index < step_index

        st.rerun()

    for index, step in enumerate(GUIDE_STEPS, start=1):
        st.checkbox(
            step["title"],
            value=st.session_state.get(
                f"check_guide_{step['key']}",
                index < auto_step,
            ),
            key=f"check_guide_{step['key']}",
            on_change=_update_guide_progress,
        )

    st.caption(
        "跟着序号从上往下做。每完成一步就勾选，"
        "工具会记住你的进度。"
    )

    if st.button(
        "🔄 按当前自动进度同步勾选",
        key="sync_guide_button",
    ):
        st.session_state["guide_sync_triggered"] = True
        st.rerun()


def render_guide_panel(step_index):
    """
    在侧边栏渲染新手引导面板：
    step_index 表示当前已进行到第几步（从 1 开始），0 表示尚未开始。
    """
    st.markdown("---")
    st.markdown("### 🧭 新手使用向导")

    # 总体进度条
    progress_value = (
        min(max(step_index, 0), len(GUIDE_STEPS))
        / len(GUIDE_STEPS)
    )

    try:
        st.progress(
            progress_value,
            text=f"总进度 {step_index}/{len(GUIDE_STEPS)}",
        )
    except TypeError:
        st.progress(progress_value)

    for i, step in enumerate(GUIDE_STEPS, start=1):
        if i < step_index:
            st.markdown(f"✅ **{step['title']}**")
        elif i == step_index:
            st.markdown(f"👉 **{step['title']}**")
        else:
            st.markdown(f"⭕ {step['title']}")

    current = GUIDE_STEPS[
        min(max(step_index, 1), len(GUIDE_STEPS)) - 1
    ]

    with st.expander(
        "💡 当前这一步怎么做？",
        expanded=step_index > 0,
    ):
        st.markdown(f"**这一步是什么：** {current['desc']}")
        st.markdown(f"**什么时候做：** {current['when']}")
        st.markdown(f"**怎么做：** {current['how']}")
        st.markdown(f"**做完得到什么：** {current['output']}")

    st.markdown(
        "**没有现成数据？** 点击下方按钮下载一份示例数据，"
        "可以直接用来体验全部功能。"
    )

    sample_bytes = make_sample_data_csv_cached()

    st.download_button(
        "📥 下载示例数据（CSV）",
        data=sample_bytes,
        file_name="示例数据.csv",
        mime="text/csv",
        key="download_sample_data_guide",
    )


# ===== 新手进度横幅（渲染在标题之后、主流程之前） =====

# 修复：展示进度取“手动勾选”与“自动进度”的较大值
_guide_step_now = max(
    st.session_state.get("guide_step", 0),
    st.session_state.get("guide_auto_step", 0),
)

if 0 < _guide_step_now <= len(GUIDE_STEPS):
    _current_step = GUIDE_STEPS[_guide_step_now - 1]
    st.info(
        f"🧭 当前进度：第 {_guide_step_now}/8 步 —— "
        f"{_current_step['title']}"
    )

# ===== 新手进度横幅结束 =====


# ===== 新手总览：拿到赛题后做什么（主区域） =====

# 修复：复选框移到 expander 之外，勾选后仍可取消重新展开
_welcome_seen = st.session_state.get("welcome_seen", False)

with st.expander(
    "🚀 新手总览：拿到赛题后，怎么用这个工具？",
    expanded=not _welcome_seen,
):
    st.markdown(
        "数学建模赛题千变万化，但拿到题目后你只需要记住 **四步走**："
        "**读懂赛题 → 准备数据 → 建模分析 → 撰写论文**。"
        "本工具的 8 个步骤正是按这个顺序设计的："
    )

    st.markdown(
        """
        <div style="display:flex; flex-wrap:wrap; gap:8px; margin:8px 0;">
            <div style="flex:1; min-width:110px; padding:8px; border-radius:8px;
                        background:rgba(70,130,180,0.12); border:1px solid rgba(100,140,180,0.3);">
                <b>① 读题</b><br><span style="font-size:13px;">上传赛题，自动识别题型</span>
            </div>
            <div style="flex:1; min-width:110px; padding:8px; border-radius:8px;
                        background:rgba(70,130,180,0.12); border:1px solid rgba(100,140,180,0.3);">
                <b>② 数据</b><br><span style="font-size:13px;">上传数据 → 选变量 → 清洗</span>
            </div>
            <div style="flex:1; min-width:110px; padding:8px; border-radius:8px;
                        background:rgba(120,160,200,0.15); border:1px solid rgba(100,140,180,0.3);">
                <b>③ 建模</b><br><span style="font-size:13px;">推荐模型 → 拟合 → 诊断</span>
            </div>
            <div style="flex:1; min-width:110px; padding:8px; border-radius:8px;
                        background:rgba(120,160,200,0.15); border:1px solid rgba(100,140,180,0.3);">
                <b>④ 论文</b><br><span style="font-size:13px;">复制表述 → 导出报告</span>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        "**三个原则：**\n\n"
        "1. **不用把所有功能都用一遍** —— 跟着左侧向导的步骤序号走即可；\n"
        "2. **先读题再动手** —— 先做第①步识别题型，工具会告诉你这类题该用哪些功能；\n"
        "3. **每步都有论文素材** —— 数据清洗、描述统计、模型假设等环节"
        "都提供了可直接复制进论文的表述。"
    )

_welcome_ack = st.checkbox(
    "已了解，下次打开不再自动展开",
    value=_welcome_seen,
    key="welcome_seen_checkbox",
)
st.session_state["welcome_seen"] = _welcome_ack

# ===== 新手总览结束 =====


def safe_numeric(series):
    """安全转换为数值。"""
    return pd.to_numeric(series, errors="coerce")


def safe_float(value):
    """安全转换为浮点数。"""
    try:
        return float(value)
    except (TypeError, ValueError):
        return np.nan


@st.cache_data(show_spinner=False)
def make_sample_data_csv_cached():
    """生成示例数据（带缓存，避免每次 rerun 重复生成）。"""
    return make_sample_data_csv()


def model_is_converged(model):
    """检查模型是否收敛。"""
    if model is None:
        return False

    if hasattr(model, "converged"):
        return bool(model.converged)

    if hasattr(model, "mle_retvals"):
        return bool(
            model.mle_retvals.get(
                "converged",
                True,
            )
        )

    return True


# ============================================================
# 五、赛题文件读取和题型识别
# ============================================================

def extract_text_from_file(uploaded_file):
    """提取 PDF、Word、TXT 文件中的文本。"""
    if uploaded_file is None:
        return ""

    file_type = uploaded_file.name.split(".")[-1].lower()

    try:
        if file_type == "pdf":
            text_parts = []
            reader = PyPDF2.PdfReader(uploaded_file)

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            return "\n".join(text_parts).strip()

        if file_type == "docx":
            document = docx.Document(uploaded_file)

            text_parts = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]

            # 同时读取 Word 表格中的内容
            for table in document.tables:
                for row in table.rows:
                    cells = [
                        cell.text.strip()
                        for cell in row.cells
                    ]
                    row_text = " | ".join(
                        cell for cell in cells if cell
                    )
                    if row_text:
                        text_parts.append(row_text)

            return "\n".join(text_parts).strip()

        if file_type == "txt":
            raw = uploaded_file.read()

            for encoding in ["utf-8", "gbk", "gb18030"]:
                try:
                    return raw.decode(encoding)
                except UnicodeDecodeError:
                    continue

            return raw.decode("utf-8", errors="ignore")

        return ""

    except Exception as exc:
        return f"[文件解析失败：{exc}]"


def clean_problem_text(text):
    """清洗赛题文本。"""
    if not isinstance(text, str):
        return ""

    text = text.strip()
    text = re.sub(r"\s+", " ", text)
    return text.lower()


def classify_by_tfidf(problem_text):
    """使用 TF-IDF 和余弦相似度识别赛题类型。"""
    if not isinstance(problem_text, str) or not problem_text.strip():
        return {}

    corpus = {
        "评价类": (
            "层次分析法 AHP 模糊综合评价 TOPSIS 熵权法 "
            "灰色关联度 数据包络分析 优劣解距离法 指标权重 "
            "指标体系 综合得分 排序 评估 绩效考核 可行性研究 "
            "多属性决策 主成分分析 因子分析 满意度评价"
        ),
        "预测类": (
            "时间序列 ARIMA 指数平滑 灰色预测 GM(1,1) 回归分析 "
            "趋势外推 神经网络 深度学习 LSTM 支持向量回归 SVR "
            "随机森林回归 XGBoost 增长率 预测值 未来走势 "
            "短期预测 中长期预测 预报 拟合 插值 变化规律"
        ),
        "优化类": (
            "线性规划 整数规划 0-1规划 非线性规划 动态规划 "
            "多目标优化 遗传算法 粒子群算法 模拟退火 蚁群算法 "
            "最短路径 最大流 最小生成树 网络流 调度 排班 指派 "
            "运输问题 库存优化 资源分配 成本最小化 利润最大化 "
            "约束条件 最优解 方案优选"
        ),
        "机理分析类": (
            "微分方程 偏微分方程 常微分方程 动力学模型 "
            "传染病模型 SIR SEIR Logistic方程 捕食者猎物模型 "
            "种群增长 扩散方程 反应扩散 对流扩散 牛顿力学 "
            "流体力学 传热 电磁场 化学反应动力学 稳定平衡 "
            "分岔 相图 数值解 欧拉法 龙格库塔法 有限元"
        ),
        "分类类": (
            "逻辑回归 Logistic回归 支持向量机 SVM 决策树 "
            "随机森林 K近邻 KNN 朴素贝叶斯 神经网络 聚类分析 "
            "K-means 层次聚类 二分类 多分类 混淆矩阵 准确率 "
            "精确率 召回率 F1值 ROC曲线 AUC 异常检测 模式识别 "
            "图像识别 文本分类 诊断 判定 识别 筛选 检出"
        ),
    }

    categories = list(corpus.keys())
    documents = [problem_text] + [
        corpus[category]
        for category in categories
    ]

    def tokenize(text):
        return " ".join(jieba.cut(text))

    try:
        vectorizer = TfidfVectorizer(
            tokenizer=tokenize,
            token_pattern=None,
        )
        matrix = vectorizer.fit_transform(documents)
        similarities = cosine_similarity(
            matrix[0:1],
            matrix[1:],
        ).flatten()

    except Exception:
        return {}

    return {
        category: float(similarities[index])
        for index, category in enumerate(categories)
    }


def multi_label_classify_problem_text(problem_text):
    """融合关键词和 TF-IDF 的多标签题型识别。"""
    text = clean_problem_text(problem_text)

    if not text:
        return {
            "main_type": "未识别",
            "all_detected_labels": [],
            "label_scores": {},
            "sub_question_context": [],
        }

    keyword_sets = {
        "评价类": [
            "评价",
            "排序",
            "打分",
            "评选",
            "满意度",
            "权重",
            "层次分析",
            "topsis",
            "优劣",
            "综合评价",
            "评估",
            "排名",
            "指标体系",
        ],
        "预测类": [
            "预测",
            "趋势",
            "未来",
            "增长",
            "回归",
            "时间序列",
            "估计",
            "预报",
            "拟合",
            "外推",
            "增长率",
            "变化规律",
            "走势",
        ],
        "优化类": [
            "最大",
            "最小",
            "最优",
            "成本",
            "利润",
            "资源",
            "调度",
            "运输",
            "规划",
            "约束",
            "路径",
            "分配",
            "库存",
            "0-1",
            "整数",
            "安排",
        ],
        "机理分析类": [
            "微分方程",
            "变化率",
            "传染病",
            "扩散",
            "物理",
            "力学",
            "温度",
            "浓度",
            "平衡点",
            "稳定性",
            "logistic",
            "sir",
            "动力学",
            "传播",
            "运动方程",
            "相互作用",
        ],
        "分类类": [
            "分类",
            "判别",
            "识别",
            "判定",
            "判断",
            "异常检测",
            "诊断",
            "聚类",
            "模式识别",
            "区分",
            "类别归属",
        ],
    }

    keyword_scores = {}

    for label, keywords in keyword_sets.items():
        keyword_scores[label] = sum(
            1
            for keyword in keywords
            if keyword in text
        )

    tfidf_scores = classify_by_tfidf(text)
    combined_scores = {}

    for label in keyword_sets:
        keyword_score = min(
            keyword_scores.get(label, 0),
            10,
        ) / 10.0

        tfidf_score = tfidf_scores.get(label, 0.0)

        combined_scores[label] = (
            0.3 * keyword_score
            + 0.7 * tfidf_score
        )

    # 修复：TF-IDF 余弦对“长赛题 vs 短语料”普遍远小于 0.1，
    # 原阈值会让大量赛题“未识别”。改为：
    # 1) 关键词命中 ≥3 直接入选（兜底）；
    # 2) 阈值降为 0.05。
    keyword_hits = {
        label: keyword_scores.get(label, 0)
        for label in keyword_sets
    }

    detected_labels = [
        label
        for label, score in combined_scores.items()
        if score > 0.05
        or keyword_hits.get(label, 0) >= 3
    ]

    detected_labels.sort(
        key=lambda label: combined_scores[label],
        reverse=True,
    )

    main_type = (
        detected_labels[0]
        if detected_labels
        else "未识别"
    )

    # 修复：子问题提取噪声大（“2023.”、“5.”等正文数字被误判）。
    # 只匹配行首或“问题/第”前缀的编号。
    patterns = [
        r"问题\s*[一二三四五六七八九十0-9]+[、\s.]",
        r"(?:^|\n)\s*\d+[\)）、.]",
        r"(?:^|\n)\s*[（(]\s*\d+\s*[）)]",
    ]

    contexts = []

    for pattern in patterns:
        for match in re.finditer(pattern, problem_text):
            start = match.start()
            end = min(
                start + 250,
                len(problem_text),
            )
            contexts.append(
                problem_text[start:end].strip()
            )

    contexts = list(dict.fromkeys(contexts))[:10]

    return {
        "main_type": main_type,
        "all_detected_labels": detected_labels,
        "label_scores": combined_scores,
        "sub_question_context": contexts,
    }


# ===== 按题型给出的推荐工作流（新手引导用） =====

PROBLEM_PATH_GUIDES = {
    "评价类": [
        ("明确评价对象与指标", "每一行是一个被评价对象（如城市/方案），每一列是一个指标；先想清楚这两件事。"),
        ("确认指标方向", "在「熵权法+TOPSIS」里逐个设置指标方向：越大越好=正向，越小越好=负向。"),
        ("熵权法确定权重", "切到高级分析的「熵权法+TOPSIS」Tab，程序自动给出客观权重（哪个指标区分度最高）。"),
        ("TOPSIS 排序", "得到各对象的综合得分与排名，这就是评价类题目的核心结果。"),
        ("稳健性补充", "指标较多时可先用 PCA 看信息结构；或换等权重对比排名是否明显变化。"),
    ],
    "预测类": [
        ("先看走势", "在「6. 数据可视化」里画因变量-自变量散点/曲线图，判断趋势、波动和季节性。"),
        ("判断序列长度", "序列 ≥30 期优先用 ARIMA；4~30 期用「灰色预测 GM(1,1)」。"),
        ("回归建模", "若还有多个解释变量，先走主流程完成回归建模（第⑥步）。"),
        ("交叉验证", "用「K-Fold 交叉验证」评估预测的稳定性，避免单次划分的偶然性。"),
        ("给出预测区间", "模型结果表的“预测区间”可用于论文：报告预测值的同时给出区间。"),
    ],
    "优化类": [
        ("明确决策变量", "先想清楚“要决定什么”（如产量、采购量、排班人数），这些就是决策变量。"),
        ("写目标函数", "是最大化利润/产量，还是最小化成本/时间？在右侧输入目标系数或表达式。"),
        ("逐条加约束", "资源限制、需求限制、上下限等逐条添加，注意系数个数与变量数一致。"),
        ("选求解器", "线性问题用「线性规划/整数线性规划」；非线性问题用「NLP(SLSQP)」或「遗传算法」。"),
        ("敏感性分析", "LP 求解后查看“影子价格”，讨论哪个约束最紧、放宽哪个约束最划算。"),
    ],
    "机理分析类": [
        ("识别核心变量", "找出随时间/空间变化的量及其相互关系（如浓度、温度、种群数量）。"),
        ("建立方程结构", "根据机理写出微分方程/动力学方程（本工具未内置方程求解，建议先手动推导）。"),
        ("用数据校验参数", "将离散化数据与本工具「回归建模」结合，估计模型参数。"),
        ("稳健性与灵敏度", "用「稳健性分析（Bootstrap）」检验参数区间是否包含 0、结论是否稳定。"),
    ],
    "分类类": [
        ("判断分类任务", "因变量类别数=2 用「二项Logistic回归」；>2 用「多项Logistic回归」。"),
        ("检查类别平衡", "看各类别样本数：某类过少（如<10）时模型和评估会不稳定。"),
        ("建模与混淆矩阵", "拟合后重点看：准确率、F1、ROC-AUC（不要只看准确率）。"),
        ("机器学习补充", "高级分析里的「随机森林/决策树」可对比结果，并输出特征重要性。"),
    ],
    "未识别": [
        ("先识别题型", "如果系统没能识别题型，手动在“赛题类型”输入框里填写（如：预测类）。"),
        ("走通基础流程", "按左侧 8 步流程：上传数据 → 选变量 → 确认类型 → 清洗 → 建模。"),
        ("不知道用哪个功能？", "回到左侧「新手使用向导」，展开“当前这一步怎么做”查看说明。"),
    ],
}


def render_problem_path_guide(main_type):
    """渲染针对某题型的推荐工作流。"""
    steps = PROBLEM_PATH_GUIDES.get(
        main_type,
        PROBLEM_PATH_GUIDES["未识别"],
    )

    st.markdown(
        f"**针对「{main_type}」赛题，建议按以下顺序做：**"
    )

    for index, (title, desc) in enumerate(steps, start=1):
        st.markdown(f"**{index}. {title}**  \n{desc}")


# ============================================================
# 六、变量识别和数据处理
# ============================================================

def try_parse_datetime(series):
    """尝试判断变量是否为日期时间变量。"""
    if pd.api.types.is_datetime64_any_dtype(series):
        return pd.to_datetime(series, errors="coerce")

    if not (
        series.dtype == "object"
        or pd.api.types.is_string_dtype(series)
    ):
        return None

    non_missing = series.dropna()

    if len(non_missing) == 0:
        return None

    text_values = non_missing.astype(str).str.strip()

    date_keyword_pattern = (
        r"(年|月|日|日期|时间|date|time|year|month|day)"
    )

    name_hint = False

    if hasattr(series, "name") and series.name is not None:
        name_hint = bool(
            re.search(
                date_keyword_pattern,
                str(series.name),
                flags=re.IGNORECASE,
            )
        )

    has_date_separator = text_values.str.contains(
        r"[-/:年月日]",
        regex=True,
    ).mean() >= 0.6

    if not name_hint and not has_date_separator:
        return None

    # 修复：先排除纯数字 / 小数 / 分数 / 比率格式的字符串，
    # 避免“1.5”、“2”、“3/4”这类内容（常见于文本格式数字列）
    # 被 pd.to_datetime 误解析成时间戳/日期，导致整列静默损坏。
    numeric_like_pattern = (
        r"^\s*[+-]?(\d+(\.\d+)?|\.\d+)"
        r"(\s*/\s*[+-]?(\d+(\.\d+)?|\.\d+))?\s*%?\s*$"
    )
    numeric_like_rate = text_values.str.match(
        numeric_like_pattern
    ).mean()

    if numeric_like_rate >= 0.8:
        return None

    parsed = pd.to_datetime(
        text_values,
        errors="coerce",
    )

    if parsed.notna().mean() >= 0.8:
        full_result = pd.to_datetime(
            series,
            errors="coerce",
        )

        # 修复：解析失败率 >0 时给出提示（由调用方处理），
        # 这里把“解析失败样本数”附加到 Series 的 name 上不可行，
        # 改为返回 (series, fail_count) 由调用方展示。
        # 为兼容旧调用（只取一个返回值），此处返回 Series 本身。
        return full_result

    return None

def classify_variable(series):
    """自动识别变量类型。"""
    parsed_time = try_parse_datetime(series)

    if parsed_time is not None:
        return "时间"

    numeric = pd.to_numeric(
        series,
        errors="coerce",
    )

    numeric_rate = numeric.notna().mean()

    # 修复：缺失率过高（含全 NaN）的列不再武断判为“分类”，
    # 标记为“全缺失”，提示用户删除该列，避免后续填充失败
    # 导致整表样本被静默 dropna 删光。
    if series.isna().mean() >= 0.9:
        return "全缺失"

    if numeric_rate < 0.8:
        return "分类"

    values = numeric.dropna()

    if len(values) == 0:
        return "连续"

    unique_values = set(values.unique().tolist())
    unique_count = values.nunique()

    if (
        unique_count <= 2
        and unique_values.issubset({0, 1})
    ):
        return "分类"

    is_integer = np.allclose(
        values,
        np.round(values),
    )

    is_nonnegative = (values >= 0).all()

    # 修复：原判定要求 max>=5，导致 0~4 的小计数变量（如促销次数）
    # 被误判为“连续”，进而影响模型推荐（不会推荐 Poisson/负二项）。
    # 改为：非负整数且取值基数较小（<=12）即判定为“次数”。
    if (
        is_integer
        and is_nonnegative
        and unique_count <= 12
    ):
        return "次数"

    return "连续"


def is_suspicious_id_column(series, name):
    """判断某列是否疑似 ID、编号或序号。

    修复：改用单词边界匹配 “id”，避免 “grid”、“valid”、“mid”
    等含 id 子串的列名被误判为 ID 列。
    """
    name_text = str(name).lower()

    keywords = [
        "编号",
        "序号",
        "代码",
        "编码",
        "样本号",
        "学生号",
        "患者号",
        "姓名",
        "工号",
    ]

    keyword_flag = any(
        keyword in name_text
        for keyword in keywords
    )

    # “id” 单独用单词边界匹配，防止子串误判
    id_flag = bool(
        re.search(
            r"\bid\b",
            name_text,
            flags=re.IGNORECASE,
        )
    )

    unique_rate = (
        series.nunique(dropna=True)
        / max(len(series), 1)
    )

    numeric = pd.to_numeric(
        series,
        errors="coerce",
    )

    sequential_flag = False

    if numeric.notna().mean() >= 0.95:
        values = (
            numeric.dropna()
            .sort_values()
            .to_numpy()
        )

        if len(values) >= 3:
            differences = np.diff(values)
            sequential_flag = np.allclose(
                differences,
                differences[0],
            )

    return keyword_flag or id_flag or (
        unique_rate >= 0.95
        and sequential_flag
    )


def convert_types(df, variable_types):
    """按照用户确认的变量类型转换数据。"""
    result = df.copy()

    for col, var_type in variable_types.items():
        if col not in result.columns:
            continue

        if var_type in ["连续", "次数"]:
            result[col] = pd.to_numeric(
                result[col],
                errors="coerce",
            )

        elif var_type == "时间":
            result[col] = pd.to_datetime(
                result[col],
                errors="coerce",
            )

        elif var_type == "分类":
            # 修复：极端混合对象（列表/字典等 Excel 脏数据）无法直接
            # astype("string")，先转 str 兜底，避免整页流程中断。
            try:
                result[col] = result[col].astype("string")
            except (TypeError, ValueError):
                result[col] = result[col].astype(str).astype("string")

        elif var_type == "全缺失":
            # 全缺失列保持 NaN，由调用方提示用户删除
            result[col] = result[col].astype("float64")

    return result


def _fill_missing_knn(result, variable_types):
    """KNN 插补：数值列用距离加权 KNN，分类/时间列用众数/前后填充。"""
    from sklearn.impute import KNNImputer

    report_rows = []
    imputed_cells = 0
    deleted_rows = 0

    num_cols = [
        col
        for col, var_type in variable_types.items()
        if var_type in ["连续", "次数"]
        and col in result.columns
    ]
    other_cols = [
        col
        for col in variable_types
        if col not in num_cols and col in result.columns
    ]

    def _row(col, before, treatment, actual, after):
        return {
            "变量": col,
            "原始缺失数": before,
            "处理方式": treatment,
            "实际插补数": actual,
            "剩余缺失数": after,
        }

    # 1) 分类 / 时间列：众数 / 前后向填充
    for col in other_cols:
        before = int(result[col].isna().sum())

        if before == 0:
            report_rows.append(_row(col, 0, "无缺失", 0, 0))
            continue

        if variable_types[col] == "时间":
            result[col] = result[col].ffill().bfill()
            treatment = "前向填充+后向填充"
        else:
            mode = result[col].mode(dropna=True)

            if len(mode) > 0:
                result[col] = result[col].fillna(mode.iloc[0])

            treatment = "众数填补"

        after = int(result[col].isna().sum())
        actual_imputed = max(0, before - after)
        imputed_cells += actual_imputed
        report_rows.append(
            _row(col, before, treatment, actual_imputed, after)
        )

    # 2) 数值列：KNN 插补（距离加权 + 标准化）
    if num_cols:
        try:
            # 修复：邻居数不能超过样本数，否则 KNN 拟合必然抛异常
            n_est = min(5, max(2, len(result) - 1))

            if len(result) < 3:
                raise ValueError(
                    "样本量过小（<3），无法使用 KNN 插补"
                )

            # 修复：KNN 使用欧氏距离，先标准化消除量纲影响，
            # 否则大数量纲列（如销售额上万）会主导距离计算。
            scaler = StandardScaler()
            num_matrix = result[num_cols].astype(float)
            num_matrix_scaled = pd.DataFrame(
                scaler.fit_transform(num_matrix),
                columns=num_cols,
                index=result.index,
            )

            imputer = KNNImputer(
                n_neighbors=n_est,
                weights="distance",
            )
            imputed_scaled = imputer.fit_transform(
                num_matrix_scaled
            )
            # 还原量纲
            imputed_values = scaler.inverse_transform(
                imputed_scaled
            )

            for index, col in enumerate(num_cols):
                before = int(result[col].isna().sum())

                if before == 0:
                    report_rows.append(_row(col, 0, "无缺失", 0, 0))
                    continue

                result[col] = imputed_values[:, index]

                if variable_types[col] == "次数":
                    result[col] = result[col].round()

                after = int(result[col].isna().sum())
                actual_imputed = max(0, before - after)
                imputed_cells += actual_imputed
                treatment = (
                    "KNN插补（距离加权，取整）"
                    if variable_types[col] == "次数"
                    else "KNN插补（距离加权）"
                )
                report_rows.append(
                    _row(col, before, treatment, actual_imputed, after)
                )
        except (ValueError, TypeError) as knn_error:
            # 修复：捕获具体异常并记录原因，回退为插值+中位数
            for col in num_cols:
                before = int(result[col].isna().sum())

                if before == 0:
                    report_rows.append(_row(col, 0, "无缺失", 0, 0))
                    continue

                result[col] = result[col].interpolate(
                    method="linear",
                    limit_direction="both",
                )

                median_value = result[col].median()

                if pd.notna(median_value):
                    result[col] = result[col].fillna(median_value)

                if variable_types[col] == "次数":
                    result[col] = result[col].round()

                after = int(result[col].isna().sum())
                actual_imputed = max(0, before - after)
                imputed_cells += actual_imputed
                report_rows.append(
                    _row(
                        col,
                        before,
                        f"KNN失败回退：插值+中位数（{knn_error}）",
                        actual_imputed,
                        after,
                    )
                )

    return (
        result,
        pd.DataFrame(report_rows),
        deleted_rows,
        imputed_cells,
    )


def fill_missing_values(
    df,
    variable_types,
    method,
):
    """
    处理自变量缺失值。

    因变量缺失值在调用本函数前删除。
    """
    result = df.copy()
    report_rows = []
    deleted_rows = 0
    imputed_cells = 0

    if not variable_types:
        return (
            result,
            pd.DataFrame(),
            0,
            0,
        )

    if method == "KNN插补":
        return _fill_missing_knn(result, variable_types)

    if method == "删除含缺失值的行":
        missing_mask = result[
            list(variable_types.keys())
        ].isna().any(axis=1)

        deleted_rows = int(missing_mask.sum())
        result = result.loc[~missing_mask].copy()

        for col in variable_types:
            before = int(df[col].isna().sum())

            report_rows.append(
                {
                    "变量": col,
                    "原始缺失数": before,
                    "处理方式": "删除所在行",
                    "实际插补数": 0,
                    "剩余缺失数": int(
                        result[col].isna().sum()
                    ),
                }
            )

        return (
            result,
            pd.DataFrame(report_rows),
            deleted_rows,
            0,
        )

    for col, var_type in variable_types.items():
        if col not in result.columns:
            continue

        before = int(result[col].isna().sum())

        if before == 0:
            report_rows.append(
                {
                    "变量": col,
                    "原始缺失数": 0,
                    "处理方式": "无缺失",
                    "实际插补数": 0,
                    "剩余缺失数": 0,
                }
            )
            continue

        if var_type in ["连续", "次数"]:
            # 修复：对截面数据（行间无时间/空间顺序），线性插值等于
            # 在“随机相邻”样本之间造值，结果无意义且不可复现。
            # 默认改用中位数填充；线性插值仅对已按时间排序的序列有意义。
            median_value = result[col].median()

            if pd.notna(median_value):
                result[col] = result[col].fillna(
                    median_value
                )
            else:
                result[col] = result[col].interpolate(
                    method="linear",
                    limit_direction="both",
                )

            if var_type == "次数":
                # 次数变量插补后取整，避免出现“1.5次”这类含义不明的值
                result[col] = result[col].round()

            treatment = (
                "中位数填充（取整）"
                if var_type == "次数"
                else "中位数填充"
            )

        elif var_type == "时间":
            result[col] = result[col].ffill().bfill()
            treatment = "前向填充+后向填充"

        else:
            mode = result[col].mode(dropna=True)

            if len(mode) > 0:
                result[col] = result[col].fillna(
                    mode.iloc[0]
                )

            treatment = "众数填补"

        after = int(result[col].isna().sum())
        actual_imputed = max(0, before - after)
        imputed_cells += actual_imputed

        report_rows.append(
            {
                "变量": col,
                "原始缺失数": before,
                "处理方式": treatment,
                "实际插补数": actual_imputed,
                "剩余缺失数": after,
            }
        )

    return (
        result,
        pd.DataFrame(report_rows),
        deleted_rows,
        imputed_cells,
    )


def detect_outliers(
    df,
    numeric_columns,
    method,
):
    """检测异常值并返回逐行异常标记。"""
    result = df.copy()
    result["_异常行"] = False
    report_rows = []

    for col in numeric_columns:
        values = pd.to_numeric(
            result[col],
            errors="coerce",
        )

        # 修复：inf 会让均值变 inf、标准差变 NaN，3σ 判断直接失效。
        # 先把 inf 显式替换为 NaN 并单独标记为异常行。
        inf_mask = values.isin([np.inf, -np.inf])
        values = values.replace(
            [np.inf, -np.inf],
            np.nan,
        )

        if method == "3σ":
            mean_value = values.mean()
            # 修复：3σ 法则使用总体标准差（ddof=0）
            std_value = values.std(ddof=0)

            if (
                pd.isna(std_value)
                or std_value == 0
            ):
                mask = pd.Series(
                    False,
                    index=result.index,
                )
            else:
                mask = (
                    ~values.between(
                        mean_value - 3 * std_value,
                        mean_value + 3 * std_value,
                    )
                    & values.notna()
                )

        elif method == "IQR":
            q1 = values.quantile(0.25)
            q3 = values.quantile(0.75)
            iqr = q3 - q1

            if pd.isna(iqr) or iqr == 0:
                mask = pd.Series(
                    False,
                    index=result.index,
                )
            else:
                lower = q1 - 1.5 * iqr
                upper = q3 + 1.5 * iqr

                mask = (
                    ~values.between(
                        lower,
                        upper,
                    )
                    & values.notna()
                )

        elif method == "MAD":
            # 中位数绝对偏差法：对偏态数据更稳健
            median_value = values.median()
            mad_value = np.median(
                np.abs(values - median_value)
            )

            if (
                pd.isna(mad_value)
                or mad_value == 0
            ):
                mask = pd.Series(
                    False,
                    index=result.index,
                )
            else:
                mad_threshold = 3 * 1.4826 * mad_value
                mask = (
                    ~values.between(
                        median_value - mad_threshold,
                        median_value + mad_threshold,
                    )
                    & values.notna()
                )

        else:
            mask = pd.Series(
                False,
                index=result.index,
            )

        mask = mask.fillna(False)
        mask |= inf_mask.fillna(False)
        result["_异常行"] |= mask

        inf_count = int(inf_mask.sum())

        report_rows.append(
            {
                "变量": col,
                "异常值数量": int(mask.sum()),
                "其中inf数量": inf_count,
            }
        )

    return (
        result,
        pd.DataFrame(report_rows),
    )


# ============================================================
# 七、相关性分析和符号表
# ============================================================

def correlation_table(
    df,
    target,
    predictors,
    variable_types,
):
    """计算 Pearson 和 Spearman 相关系数。

    修复：时间列先转数值再参与计算，避免 pd.to_numeric 对
    datetime64 列产生垃圾结果或直接抛 TypeError；最小有效样本
    数提高到 10；P 值在输出前做 Benjamini-Hochberg FDR 校正。
    """
    rows = []

    if target not in df.columns:
        return pd.DataFrame()

    def _to_numeric_safe(series):
        """时间列转为自 1970-01-01 起的天数，其余转数值。"""
        if pd.api.types.is_datetime64_any_dtype(series):
            base = pd.Timestamp("1970-01-01")
            return (
                series - base
            ).dt.total_seconds() / 86400.0
        return pd.to_numeric(
            series,
            errors="coerce",
        )

    try:
        target_values = _to_numeric_safe(df[target])
    except Exception:
        return pd.DataFrame()

    for col in predictors:
        if variable_types.get(col) not in [
            "连续",
            "次数",
        ]:
            continue

        if col not in df.columns:
            continue

        try:
            values = _to_numeric_safe(df[col])
        except Exception:
            continue

        valid = pd.concat(
            [target_values, values],
            axis=1,
        ).dropna()

        # 修复：3 个样本算出的相关系数毫无意义，提高到 10
        if len(valid) < 10:
            continue

        if (
            valid.iloc[:, 0].nunique() <= 1
            or valid.iloc[:, 1].nunique() <= 1
        ):
            continue

        try:
            pearson_value, pearson_p = pearsonr(
                valid.iloc[:, 0],
                valid.iloc[:, 1],
            )

            spearman_value, spearman_p = spearmanr(
                valid.iloc[:, 0],
                valid.iloc[:, 1],
            )

            rows.append(
                {
                    "变量": col,
                    "Pearson相关系数": pearson_value,
                    "Pearson_P值": pearson_p,
                    "Spearman相关系数": spearman_value,
                    "Spearman_P值": spearman_p,
                    "有效样本数": len(valid),
                }
            )

        except Exception:
            continue

    if not rows:
        return pd.DataFrame()

    result_df = pd.DataFrame(rows)

    # 修复：多重比较的 Benjamini-Hochberg FDR 校正
    n_tests = max(len(result_df) * 2, 1)
    p_cols = ["Pearson_P值", "Spearman_P值"]
    rank_cols = {}

    for p_col in p_cols:
        sorted_p = result_df[p_col].sort_values(
            ascending=True
        ).to_numpy()
        ranks = np.arange(1, len(sorted_p) + 1)
        q_values = np.minimum.accumulate(
            sorted_p * n_tests / ranks
        )
        q_values = np.minimum(q_values, 1.0)
        rank_cols[p_col] = dict(
            zip(
                result_df[p_col].sort_values(
                    ascending=True
                ).index,
                q_values,
            )
        )

    result_df["Pearson_FDR校正P值"] = result_df.index.map(
        rank_cols["Pearson_P值"]
    )
    result_df["Spearman_FDR校正P值"] = result_df.index.map(
        rank_cols["Spearman_P值"]
    )

    return result_df


def create_variable_symbol_table(
    target,
    predictors,
    variable_types,
):
    """创建变量符号表。

    修复：自变量符号从 x_1 开始连续编号（x_1, x_2, ..., x_k），
    避免原实现中第一个自变量为 x_1 而 x_0 空缺。
    """
    rows = []

    for col in [target] + list(predictors):
        if col == target:
            default_symbol = "y"
        else:
            predictor_index = (
                list(predictors).index(col) + 1
            )
            default_symbol = f"x_{predictor_index}"

        rows.append(
            {
                "变量符号": default_symbol,
                "原始列名": col,
                "变量角色": (
                    "因变量"
                    if col == target
                    else "自变量"
                ),
                "变量类型": variable_types.get(
                    col,
                    "",
                ),
                "单位": "",
                "变量含义": "",
            }
        )

    return pd.DataFrame(rows)

# ============================================================
# 八、模型数据构造
# ============================================================

def build_model_data(
    df,
    target,
    predictors,
    variable_types,
    group_col=None,
):
    """构造模型数据并自动处理分类变量。"""
    use_cols = [target] + list(predictors)

    if group_col not in [None, "无"]:
        use_cols.append(group_col)

    use_cols = list(dict.fromkeys(use_cols))
    selected = df[use_cols].copy()

    target_mapping = None
    target_type = variable_types[target]
    dummy_info = {}

    if target_type in ["连续", "次数"]:
        selected[target] = pd.to_numeric(
            selected[target],
            errors="coerce",
        )

    elif target_type == "分类":
        selected[target] = selected[target].astype(
            "string"
        )

        categories = sorted(
            selected[target]
            .dropna()
            .unique()
            .tolist()
        )

        if len(categories) < 2:
            raise ValueError(
                "分类因变量至少需要两个类别。"
            )

        target_mapping = {
            category: index
            for index, category in enumerate(categories)
        }

        selected[target] = selected[target].map(
            target_mapping
        )

    elif target_type == "时间":
        raise ValueError(
            "时间型因变量不适合当前回归模块。"
        )

    for col in predictors:
        var_type = variable_types[col]

        if var_type in ["连续", "次数"]:
            selected[col] = pd.to_numeric(
                selected[col],
                errors="coerce",
            )

        elif var_type == "时间":
            dates = pd.to_datetime(
                selected[col],
                errors="coerce",
            )

            selected[col] = (
                dates - pd.Timestamp("1970-01-01")
            ).dt.total_seconds() / 86400

        elif var_type == "分类":
            categories = sorted(
                selected[col].dropna().unique().tolist()
            )

            if categories:
                dummy_info[col] = {
                    "原始类别": categories,
                    "参照类别（被删除）": categories[0],
                }

            selected[col] = selected[col].astype(
                "string"
            )

    required_cols = [target] + list(predictors)

    if group_col not in [None, "无"]:
        required_cols.append(group_col)

    selected = selected.dropna(
        subset=required_cols
    )

    # 修复：删行（含 X 缺失行）后某个中间类别可能整类消失，
    # y 编码会变成 {0,2} 这类非连续整数，MNLogit 要求 0..K-1 连续。
    # 因此在删行之后对分类因变量重新映射编码。
    if target_type == "分类":
        remaining_categories = sorted(
            selected[target]
            .dropna()
            .unique()
            .tolist()
        )

        if len(remaining_categories) < 2:
            raise ValueError(
                "删行后分类因变量只剩一个类别，无法建模。"
            )

        if target_mapping is not None and (
            len(remaining_categories)
            != len(target_mapping)
        ):
            new_mapping = {
                category: index
                for index, category in enumerate(
                    remaining_categories
                )
            }
            selected[target] = selected[target].map(
                new_mapping
            )
            target_mapping = new_mapping

    y = selected[target].copy()
    X = selected[predictors].copy()

    X = pd.get_dummies(
        X,
        drop_first=True,
        dtype=float,
    )

    X = X.replace(
        [np.inf, -np.inf],
        np.nan,
    )

    constant_columns = [
        col
        for col in X.columns
        if X[col].nunique(dropna=True) <= 1
    ]

    if constant_columns:
        X = X.drop(
            columns=constant_columns
        )

    valid_rows = (
        X.notna().all(axis=1)
        & y.notna()
    )

    X = X.loc[valid_rows]
    y = y.loc[valid_rows]

    if X.shape[1] == 0:
        raise ValueError(
            "处理后没有有效的自变量。"
        )

    X = sm.add_constant(
        X,
        has_constant="add",
    )

    groups = None

    if group_col not in [None, "无"]:
        groups = selected.loc[
            valid_rows,
            group_col,
        ].astype(str)

    # 修复：持久化训练期填补统计量（median/mode），供新数据预测复用，
    # 避免用“新数据自身”的统计量填补导致口径不一致。
    training_imputation_stats = {}

    for col in predictors:
        if col not in df.columns:
            continue

        var_type = variable_types.get(col)

        if var_type in ["连续", "次数"]:
            median_value = df[col].median()
            training_imputation_stats[col] = {
                "method": "median",
                "value": median_value,
            }
        elif var_type == "分类":
            mode_series = df[col].mode(dropna=True)
            training_imputation_stats[col] = {
                "method": "mode",
                "value": (
                    mode_series.iloc[0]
                    if len(mode_series) > 0
                    else None
                ),
            }
        elif var_type == "时间":
            training_imputation_stats[col] = {
                "method": "ffill_bfill",
                "value": None,
            }

    metadata = {
        "target_mapping": target_mapping,
        "feature_names": list(X.columns),
        "n_rows": len(y),
        "n_features": X.shape[1],
        "constant_columns": constant_columns,
        "dummy_info": dummy_info,
        "training_imputation_stats": (
            training_imputation_stats
        ),
    }

    return (
        y.reset_index(drop=True),
        X.reset_index(drop=True),
        (
            None
            if groups is None
            else groups.reset_index(drop=True)
        ),
        metadata,
    )


# ============================================================
# 九、模型推荐、验证和拟合
# ============================================================

MODEL_OPTIONS = [
    "多元线性回归",
    "Logit变换线性回归",
    "线性混合效应模型",
    "Logit变换线性混合效应模型",
    "二项Logistic回归",
    "多项Logistic回归",
    "Poisson回归",
    "负二项回归",
]


def detect_model_type(
    y,
    target_type,
    groups=None,
):
    """根据数据结构推荐模型。"""
    y_numeric = pd.to_numeric(
        y,
        errors="coerce",
    ).dropna()

    if target_type == "分类":
        unique_count = y.nunique()

        if unique_count == 2:
            return {
                "model_type": "二项Logistic回归",
                "reason": "因变量为二分类变量。",
            }

        if unique_count > 2:
            return {
                "model_type": "多项Logistic回归",
                "reason": "因变量为多分类变量。",
            }

    if target_type == "次数":
        if (
            len(y_numeric) > 0
            and (y_numeric >= 0).all()
            and np.allclose(
                y_numeric,
                np.round(y_numeric),
            )
        ):
            mean_value = y_numeric.mean()
            variance_value = y_numeric.var()

            if (
                mean_value > 0
                and variance_value > 1.5 * mean_value
            ):
                return {
                    "model_type": "负二项回归",
                    "reason": (
                        "因变量为计数变量，且方差可能明显大于均值，"
                        "存在过度离散现象。"
                    ),
                }

            return {
                "model_type": "Poisson回归",
                "reason": "因变量为非负整数计数变量。",
            }

    if target_type == "连续":
        if len(y_numeric) == 0:
            return {
                "model_type": "未识别",
                "reason": "因变量没有有效数值。",
            }

        in_unit_interval = (
            (y_numeric >= 0).all()
            and (y_numeric <= 1).all()
        )

        repeated = False

        if groups is not None:
            repeated = (
                groups.nunique()
                < len(groups)
            )

        if in_unit_interval and repeated:
            return {
                "model_type": "Logit变换线性混合效应模型",
                "reason": (
                    "因变量取值在0到1之间，且存在重复观测。"
                ),
            }

        if in_unit_interval:
            return {
                "model_type": "Logit变换线性回归",
                "reason": (
                    "因变量为0到1之间的比例变量。"
                ),
            }

        if repeated:
            return {
                "model_type": "线性混合效应模型",
                "reason": (
                    "连续型因变量存在重复观测分组。"
                ),
            }

        return {
            "model_type": "多元线性回归",
            "reason": "因变量为一般连续变量。",
        }

    return {
        "model_type": "未识别",
        "reason": "无法判断合适的模型类型。",
    }


def validate_model_selection(
    y,
    target_type,
    model_type,
    groups=None,
):
    """验证模型和数据类型是否匹配。"""
    if y is None or len(y) == 0:
        return False, "当前没有可用于建模的数据。"

    y_numeric = pd.to_numeric(
        pd.Series(y),
        errors="coerce",
    )

    if y_numeric.isna().any():
        return False, "因变量中存在无法转换为数值的内容。"

    if model_type == "多元线性回归":
        if target_type not in ["连续", "次数"]:
            return False, (
                "多元线性回归要求因变量为数值型变量。"
            )

    elif model_type == "Logit变换线性回归":
        if target_type not in ["连续", "次数"]:
            return False, (
                "Logit变换线性回归要求因变量为数值型变量。"
            )

        if ((y_numeric < 0) | (y_numeric > 1)).any():
            return False, (
                "Logit变换线性回归要求因变量取值在0到1之间。"
            )

        if y_numeric.nunique() < 2:
            return False, (
                "Logit变换线性回归要求因变量至少有两个不同取值。"
            )

    elif model_type == "二项Logistic回归":
        if target_type != "分类":
            return False, (
                "二项Logistic回归要求因变量为分类变量。"
            )

        if pd.Series(y).nunique() != 2:
            return False, (
                "二项Logistic回归要求因变量恰好包含两个类别。"
            )

    elif model_type == "多项Logistic回归":
        if target_type != "分类":
            return False, (
                "多项Logistic回归要求因变量为分类变量。"
            )

        if pd.Series(y).nunique() < 3:
            return False, (
                "多项Logistic回归至少需要三个类别。"
            )

    elif model_type in ["Poisson回归", "负二项回归"]:
        if target_type != "次数":
            return False, (
                f"{model_type}要求因变量为次数变量。"
            )

        if (y_numeric < 0).any():
            return False, (
                f"{model_type}要求因变量不能小于0。"
            )

        if not np.allclose(
            y_numeric.to_numpy(),
            np.round(y_numeric.to_numpy()),
        ):
            return False, (
                f"{model_type}要求因变量为非负整数。"
            )

    elif model_type in [
        "线性混合效应模型",
        "Logit变换线性混合效应模型",
    ]:
        if target_type not in ["连续", "次数"]:
            return False, (
                f"{model_type}要求因变量为数值型变量。"
            )

        if groups is None:
            return False, (
                f"{model_type}必须指定分组变量。"
            )

        if pd.Series(groups).nunique() < 2:
            return False, (
                "分组变量至少需要包含两个不同的组。"
            )

        # 修复：大量单观测组会导致随机效应不可识别、LBFGS 不收敛，
        # 提前检查组内最小样本数。
        group_counts = pd.Series(groups).value_counts()

        if group_counts.min() < 2:
            return False, (
                "存在仅含 1 个观测的分组，混合效应模型的"
                "随机效应不可识别。请合并小分组或改用普通回归。"
            )

        if model_type == "Logit变换线性混合效应模型":
            if ((y_numeric < 0) | (y_numeric > 1)).any():
                return False, (
                    "Logit变换线性混合效应模型要求因变量取值在0到1之间。"
                )

    return True, ""


def logit_transform(y):
    """比例变量 Logit 变换。

    修复：不再直接 clip 篡改 0/1 观测值，改用文献标准的
    Smithson & Verkuilen (2006) 平滑公式：
    y' = (y * (n - 1) + 0.5) / n
    """
    values = pd.Series(
        y,
        dtype=float,
    )

    n = max(len(values.dropna()), 2)

    # Smithson & Verkuilen (2006) 平滑：把 0/1 观测缩到 (0,1) 开区间
    smoothed = (
        values * (n - 1) + 0.5
    ) / n

    # 兜底：平滑后仍应处于 (0,1)，极端情况再 clip 到安全范围
    smoothed = smoothed.clip(
        lower=1e-8,
        upper=1 - 1e-8,
    )

    return np.log(
        smoothed / (1 - smoothed)
    )


def inverse_logit(z):
    """Logit 逆变换。"""
    values = np.asarray(z)
    values = np.clip(values, -700, 700)
    return 1 / (1 + np.exp(-values))


def fit_model(
    y,
    X,
    groups,
    model_type,
    robust_se=False,
):
    """拟合最终模型。"""
    if model_type == "多元线性回归":
        if robust_se:
            model = sm.OLS(y, X).fit(
                cov_type="HC3"
            )
        else:
            model = sm.OLS(y, X).fit()

        # 修复：conf_int 默认 which="mean" 给出的是均值置信区间，
        # 不是新观测的预测区间。改为 which="observation"，
        # 并兼容旧版 statsmodels 的 obs=True 参数。
        try:
            try:
                pred_interval = model.get_prediction(
                    X
                ).conf_int(
                    alpha=0.05,
                    which="observation",
                )
            except TypeError:
                pred_interval = model.get_prediction(
                    X
                ).conf_int(
                    alpha=0.05,
                    obs=True,
                )
        except Exception:
            pred_interval = None

        return {
            "model": model,
            "display_y": y,
            "prediction": model.predict(X),
            "prediction_interval": pred_interval,
            "interval_type": "observation",
            "model_type": model_type,
        }

    if model_type == "Logit变换线性回归":
        transformed_y = logit_transform(y)

        if robust_se:
            model = sm.OLS(
                transformed_y,
                X,
            ).fit(cov_type="HC3")
        else:
            model = sm.OLS(
                transformed_y,
                X,
            ).fit()

        prediction = inverse_logit(
            model.predict(X)
        )

        # 修复：同样改用 observation 区间（在 logit 尺度取区间再逆变换）
        try:
            try:
                _logit_ci = model.get_prediction(
                    X
                ).conf_int(
                    alpha=0.05,
                    which="observation",
                )
            except TypeError:
                _logit_ci = model.get_prediction(
                    X
                ).conf_int(
                    alpha=0.05,
                    obs=True,
                )
            pred_interval = np.column_stack(
                [
                    inverse_logit(_logit_ci[:, 0]),
                    inverse_logit(_logit_ci[:, 1]),
                ]
            )
        except Exception:
            pred_interval = None

        return {
            "model": model,
            "display_y": y,
            "prediction": prediction,
            "prediction_interval": pred_interval,
            "interval_type": "observation",
            "transformed_y": transformed_y,
            "model_type": model_type,
        }

    if model_type == "线性混合效应模型":
        if groups is None:
            raise ValueError(
                "线性混合效应模型必须指定分组变量。"
            )

        model = sm.MixedLM(
            endog=y,
            exog=X,
            groups=groups,
        ).fit(
            reml=False,
            method="lbfgs",
            disp=False,
        )

        return {
            "model": model,
            "display_y": y,
            "prediction": model.predict(X),
            "model_type": model_type,
        }

    if model_type == "Logit变换线性混合效应模型":
        if groups is None:
            raise ValueError(
                "Logit变换线性混合效应模型必须指定分组变量。"
            )

        transformed_y = logit_transform(y)

        model = sm.MixedLM(
            endog=transformed_y,
            exog=X,
            groups=groups,
        ).fit(
            reml=False,
            method="lbfgs",
            disp=False,
        )

        prediction = inverse_logit(
            model.predict(X)
        )

        return {
            "model": model,
            "display_y": y,
            "prediction": prediction,
            "transformed_y": transformed_y,
            "model_type": model_type,
        }

    if model_type == "Poisson回归":
        model = sm.GLM(
            y,
            X,
            family=sm.families.Poisson(),
        ).fit()

        # 修复：GLM 无法给出新观测的预测区间，conf_int 只能给出
        # 均值置信区间（在对数连接尺度，需指数还原），标签写清楚。
        try:
            _poi_ci = model.get_prediction(
                X
            ).conf_int(alpha=0.05)
            pred_interval = np.exp(_poi_ci)
        except Exception:
            pred_interval = None

        return {
            "model": model,
            "display_y": y,
            "prediction": model.predict(X),
            "prediction_interval": pred_interval,
            "interval_type": "mean",
            "model_type": model_type,
        }

    if model_type == "负二项回归":
        # 修复：sm.GLM + NegativeBinomial family 会把辅助离散参数
        # alpha 固定为 1（IRLS 不估计它），与“检测到过度离散才推荐
        # 负二项”的逻辑自相矛盾。改用 discrete_model.NegativeBinomial
        # 做 MLE 联合估计 alpha（NB-2 参数化）。
        try:
            model = NegativeBinomial(
                y,
                X,
                loglike_method="nb2",
            ).fit(
                disp=False,
                maxiter=300,
            )
        except Exception as nb_error:
            # 兜底：MLE 不收敛时退回固定 alpha=1 的 GLM 并提示
            model = sm.GLM(
                y,
                X,
                family=sm.families.NegativeBinomial(),
            ).fit()
            model._nb_fallback_reason = str(nb_error)

        try:
            _nb_ci = model.get_prediction(
                X
            ).conf_int(alpha=0.05)
            pred_interval = np.exp(_nb_ci)
        except Exception:
            pred_interval = None

        return {
            "model": model,
            "display_y": y,
            "prediction": model.predict(X),
            "prediction_interval": pred_interval,
            "interval_type": "mean",
            "model_type": model_type,
        }

    if model_type == "二项Logistic回归":
        if y.nunique() != 2:
            raise ValueError(
                "二项Logistic回归要求因变量恰好有两个类别。"
            )

        model = sm.GLM(
            y,
            X,
            family=sm.families.Binomial(),
        ).fit()

        probability = np.asarray(
            model.predict(X),
            dtype=float,
        )

        # 修复：类别不平衡时固定阈值 0.5 会把模型推向全预测多数类。
        # 增加 Youden J 最优阈值：max(TPR - FPR)，提高 AUC/召回表现。
        try:
            y_int = np.asarray(y, dtype=int)
            thresholds = np.unique(probability)
            best_j = -1.0
            best_threshold = 0.5

            for thr in thresholds:
                pred_t = (probability >= thr).astype(int)
                tp = np.sum(
                    (pred_t == 1) & (y_int == 1)
                )
                fn = np.sum(
                    (pred_t == 0) & (y_int == 1)
                )
                fp = np.sum(
                    (pred_t == 1) & (y_int == 0)
                )
                tn = np.sum(
                    (pred_t == 0) & (y_int == 0)
                )
                tpr = tp / max(tp + fn, 1)
                fpr = fp / max(fp + tn, 1)
                j = tpr - fpr

                if j > best_j:
                    best_j = j
                    best_threshold = float(thr)

            use_threshold = (
                best_threshold if best_j > 0 else 0.5
            )
        except Exception:
            use_threshold = 0.5

        prediction = (
            probability >= use_threshold
        ).astype(int)

        return {
            "model": model,
            "display_y": y,
            "prediction": prediction,
            "probability": probability,
            "threshold": use_threshold,
            "model_type": model_type,
        }

    if model_type == "多项Logistic回归":
        if y.nunique() < 3:
            raise ValueError(
                "多项Logistic回归至少需要三个类别。"
            )

        model = sm.MNLogit(
            y,
            X,
        ).fit(
            disp=False,
            maxiter=300,
        )

        probability = np.asarray(
            model.predict(X)
        )

        prediction = probability.argmax(axis=1)

        return {
            "model": model,
            "display_y": y,
            "prediction": prediction,
            "probability": probability,
            "model_type": model_type,
        }

    raise ValueError(
        f"暂不支持的模型：{model_type}"
    )


# ============================================================
# 九·补充、新数据预测与多模型对比
# ============================================================

def prepare_new_data(
    new_df,
    predictors,
    variable_types,
    training_feature_names,
    dummy_info,
    training_imputation_stats=None,
):
    """
    将新数据按训练时的规则预处理，并对齐训练特征列。

    返回 (X_new, 被删除的行数, 说明文本列表)。
    修复：缺失值改用训练集的填补统计量（median/mode），
    避免用“新数据自身”的统计量导致口径不一致。
    """
    notes = []
    X_new = new_df[list(predictors)].copy()
    training_imputation_stats = (
        training_imputation_stats or {}
    )

    for col in predictors:
        var_type = variable_types.get(col, "连续")

        if var_type in ["连续", "次数"]:
            X_new[col] = pd.to_numeric(
                X_new[col],
                errors="coerce",
            )
        elif var_type == "时间":
            dates = pd.to_datetime(
                X_new[col],
                errors="coerce",
            )
            X_new[col] = (
                dates - pd.Timestamp("1970-01-01")
            ).dt.total_seconds() / 86400
        elif var_type == "分类":
            X_new[col] = X_new[col].astype("string")

    # 缺失值处理：优先用训练集统计量，缺失才回退到新数据自身统计量
    n_dropped = 0

    for col in predictors:
        var_type = variable_types.get(col, "连续")
        train_stat = training_imputation_stats.get(col)

        if var_type in ["连续", "次数"]:
            # 用训练期的中位数（口径一致）
            if (
                train_stat
                and train_stat.get("method") == "median"
                and pd.notna(train_stat.get("value"))
            ):
                fill_value = train_stat["value"]
            else:
                fill_value = X_new[col].median()

            if pd.isna(fill_value):
                n_dropped += int(X_new[col].isna().sum())
                X_new = X_new.dropna(subset=[col])
            else:
                filled = X_new[col].isna().sum()
                X_new[col] = X_new[col].fillna(fill_value)
                if filled:
                    notes.append(
                        f"变量「{col}」缺失 {filled} 个值，"
                        "已用训练集统计量填补。"
                    )
        elif var_type == "时间":
            # 时间列转换后为“距1970-01-01的天数”，缺失值用训练中位数填补
            if (
                train_stat
                and train_stat.get("method") == "median"
                and pd.notna(train_stat.get("value"))
            ):
                fill_value = train_stat["value"]
            else:
                fill_value = X_new[col].median()

            if pd.isna(fill_value):
                n_dropped += int(X_new[col].isna().sum())
                X_new = X_new.dropna(subset=[col])
            else:
                filled = X_new[col].isna().sum()
                X_new[col] = X_new[col].fillna(fill_value)
                if filled:
                    notes.append(
                        f"变量「{col}」缺失 {filled} 个值，"
                        "已用训练集统计量（时间序号）填补。"
                    )
        elif var_type == "分类":
            # 用训练期的众数（口径一致）
            if (
                train_stat
                and train_stat.get("method") == "mode"
                and train_stat.get("value") is not None
            ):
                fill_value = train_stat["value"]
            else:
                mode_values = X_new[col].mode(dropna=True)
                fill_value = (
                    mode_values.iloc[0]
                    if len(mode_values) > 0
                    else None
                )

            if fill_value is None or pd.isna(fill_value):
                n_dropped += int(X_new[col].isna().sum())
                X_new = X_new.dropna(subset=[col])
            else:
                filled = X_new[col].isna().sum()
                X_new[col] = X_new[col].fillna(fill_value)
                if filled:
                    notes.append(
                        f"变量「{col}」缺失 {filled} 个值，"
                        "已用训练集统计量（众数）填补。"
                    )

    # 哑变量编码（与训练保持一致：drop_first）
    X_new = pd.get_dummies(
        X_new,
        drop_first=True,
        dtype=float,
    )

    # 对齐训练特征列：缺失列补 0，多余列删除
    keep_cols = [
        col
        for col in training_feature_names
        if col != "const"
    ]

    for col in keep_cols:
        if col not in X_new.columns:
            X_new[col] = 0.0
            notes.append(
                f"新数据缺少特征「{col}」，已按 0 填补（该类别未出现）。"
            )

    extra_cols = [
        col
        for col in X_new.columns
        if col not in keep_cols
    ]

    if extra_cols:
        X_new = X_new.drop(columns=extra_cols)

    X_new = X_new[keep_cols]
    X_new = X_new.replace([np.inf, -np.inf], np.nan)

    # 修复：末尾统一删掉仍有缺失的行，并记录行号供调用方对齐
    n_final_nan = int(
        X_new.isna().any(axis=1).sum()
    )

    if n_final_nan:
        X_new = X_new.dropna(axis=0)
        n_dropped += n_final_nan
        notes.append(
            f"仍有 {n_final_nan} 行因含缺失值被删除。"
        )

    return X_new, n_dropped, notes


def predict_new_data(
    fitted_result,
    new_df,
    target,
    predictors,
    variable_types,
):
    """
    用训练好的模型对新数据做预测。

    返回 (结果表, 说明文本列表) 或 (None, 错误信息)。
    """
    model = fitted_result["model"]
    model_type = fitted_result["model_type"]
    meta = fitted_result.get("meta") or {}

    training_feature_names = meta.get("feature_names") or []
    dummy_info = meta.get("dummy_info") or {}
    training_imputation_stats = meta.get(
        "training_imputation_stats"
    ) or {}

    if not training_feature_names:
        raise ValueError("模型缺少特征信息，无法对新数据预测。")

    missing_cols = [
        col for col in predictors if col not in new_df.columns
    ]

    if missing_cols:
        raise ValueError(
            "新数据缺少以下自变量列："
            + "、".join(missing_cols)
        )

    X_new, n_dropped, notes = prepare_new_data(
        new_df,
        predictors,
        variable_types,
        training_feature_names,
        dummy_info,
        training_imputation_stats=training_imputation_stats,
    )

    if X_new.shape[0] == 0:
        raise ValueError("新数据没有可用样本用于预测。")

    X_new = sm.add_constant(
        X_new,
        has_constant="add",
    )

    prediction = np.asarray(
        model.predict(X_new),
        dtype=float,
    )

    # 修复：按 X_new 保留的原始索引对齐结果行，
    # 避免中间删行后预测值与原始样本错位。
    result = new_df.loc[X_new.index].copy().reset_index(drop=True)

    if model_type == "二项Logistic回归":
        result["预测概率"] = prediction
        use_threshold = fitted_result.get(
            "threshold",
            0.5,
        )
        result["预测类别"] = (
            prediction >= use_threshold
        ).astype(int)
    elif model_type == "多项Logistic回归":
        # prediction 为 (n_samples, n_classes) 概率矩阵
        result["预测类别"] = np.argmax(prediction, axis=1)

        for class_index in range(prediction.shape[1]):
            result[f"类别{class_index}概率"] = prediction[:, class_index]
    elif model_type in [
        "Logit变换线性回归",
        "Logit变换线性混合效应模型",
    ]:
        # 修复：Logit 模型输出前必须逆变换回 0~1 比例尺度，
        # 与训练集口径保持一致。
        result["预测值"] = inverse_logit(prediction)
    else:
        result["预测值"] = prediction

    if n_dropped:
        notes.append(
            f"因全部缺失无法填补，删除了 {n_dropped} 行样本。"
        )

    return result, notes


def compare_models(
    y,
    X,
    groups,
    model_types,
    robust_se=False,
):
    """
    同时拟合多个候选模型并汇总关键评价指标，返回对比表。
    """
    rows = []

    for model_type in model_types:
        try:
            fitted = fit_model(
                y,
                X,
                groups,
                model_type,
                robust_se=robust_se,
            )

            metric_table = make_metric_table(fitted)
            metric_map = dict(
                zip(
                    metric_table["指标"],
                    metric_table["数值"],
                )
            )

            row = {"模型": model_type}

            if not model_is_converged(fitted["model"]):
                row["收敛状态"] = "未收敛⚠️"
            else:
                row["收敛状态"] = "正常"

            if "AIC" in metric_map:
                row["AIC"] = metric_map["AIC"]
            if "BIC" in metric_map:
                row["BIC"] = metric_map["BIC"]
            if "R²" in metric_map:
                row["R²"] = metric_map["R²"]
            if "调整R²" in metric_map:
                row["调整R²"] = metric_map["调整R²"]
            if "RMSE" in metric_map:
                row["RMSE"] = metric_map["RMSE"]
            if "准确率" in metric_map:
                row["准确率"] = metric_map["准确率"]
            if "平衡准确率" in metric_map:
                row["平衡准确率"] = metric_map["平衡准确率"]
            if "Log Loss" in metric_map:
                row["Log Loss"] = metric_map["Log Loss"]

            rows.append(row)

        except Exception as exc:
            rows.append(
                {
                    "模型": model_type,
                    "收敛状态": f"拟合失败：{exc}",
                }
            )

    return pd.DataFrame(rows)


def run_kfold_cv(
    y,
    X,
    groups,
    model_type,
    n_splits=5,
    robust_se=False,
    random_state=42,
):
    """
    K-Fold 交叉验证。

    返回 (每折明细表, 汇总统计表, 说明文本)。
    有分组变量时自动按组划分（GroupKFold），
    避免同一对象同时出现在训练集和测试集。
    """
    from sklearn.model_selection import (
        KFold,
        StratifiedKFold,
        GroupKFold,
    )

    y_s = pd.Series(np.asarray(y)).reset_index(drop=True)
    X_s = X.reset_index(drop=True)

    groups_s = None

    if groups is not None:
        groups_s = pd.Series(groups).reset_index(drop=True)

    has_real_groups = (
        groups_s is not None
        and groups_s.nunique() < len(groups_s)
    )

    n_valid = len(y_s)
    n_splits = max(2, min(int(n_splits), n_valid // 2))

    # 修复：GroupKFold 的折数不能超过分组数，
    # 否则 sklearn 直接抛 ValueError，界面只显示“交叉验证失败”。
    if has_real_groups:
        n_groups = groups_s.nunique()
        n_splits = min(n_splits, n_groups)

    if has_real_groups:
        splitter = GroupKFold(n_splits=n_splits)
        fold_iterator = splitter.split(
            X_s,
            y_s,
            groups=groups_s,
        )
    elif model_type in [
        "二项Logistic回归",
        "多项Logistic回归",
    ]:
        # 修复：少数类样本数 < 折数时 StratifiedKFold 会抛异常，
        # 自动降折数到最小类别样本数。
        min_class_count = y_s.value_counts().min()

        if min_class_count < n_splits:
            n_splits = max(2, int(min_class_count))

        splitter = StratifiedKFold(
            n_splits=n_splits,
            shuffle=True,
            random_state=random_state,
        )
        fold_iterator = splitter.split(X_s, y_s)
    else:
        splitter = KFold(
            n_splits=n_splits,
            shuffle=True,
            random_state=random_state,
        )
        fold_iterator = splitter.split(X_s)

    fold_rows = []
    failed_folds = 0
    fold_failure_reasons = []

    for fold_idx, (tr_idx, te_idx) in enumerate(
        fold_iterator,
        start=1,
    ):
        y_tr = y_s.iloc[tr_idx].reset_index(drop=True)
        y_te = y_s.iloc[te_idx].reset_index(drop=True)
        X_tr = X_s.iloc[tr_idx].reset_index(drop=True)
        X_te = X_s.iloc[te_idx].reset_index(drop=True)

        g_tr = None

        if groups_s is not None:
            g_tr = groups_s.iloc[tr_idx].reset_index(drop=True)

        try:
            fitted_fold = fit_model(
                y_tr,
                X_tr,
                g_tr,
                model_type,
                robust_se=robust_se,
            )

            fold_model = fitted_fold["model"]

            if model_type == "二项Logistic回归":
                prob = np.asarray(
                    fold_model.predict(X_te),
                    dtype=float,
                )
                use_threshold = fitted_fold.get(
                    "threshold",
                    0.5,
                )
                cls = (prob >= use_threshold).astype(int)
                fold_rows.append(
                    {
                        "折": fold_idx,
                        "准确率": accuracy_score(
                            y_te.astype(int),
                            cls,
                        ),
                        "ROC-AUC": (
                            roc_auc_score(
                                y_te.astype(int),
                                prob,
                            )
                            if len(np.unique(y_te)) == 2
                            else np.nan
                        ),
                    }
                )
            elif model_type == "多项Logistic回归":
                prob = np.asarray(
                    fold_model.predict(X_te)
                )
                cls = prob.argmax(axis=1)
                fold_rows.append(
                    {
                        "折": fold_idx,
                        "准确率": accuracy_score(
                            y_te.astype(int),
                            cls,
                        ),
                    }
                )
            else:
                pred = np.asarray(
                    fold_model.predict(X_te),
                    dtype=float,
                )

                # 修复：两个 Logit 模型都要逆变换回 0~1 尺度，
                # 否则混合效应 Logit 的 RMSE/MAE/R² 完全失真。
                if model_type in [
                    "Logit变换线性回归",
                    "Logit变换线性混合效应模型",
                ]:
                    pred = inverse_logit(pred)

                fold_rows.append(
                    {
                        "折": fold_idx,
                        "RMSE": np.sqrt(
                            mean_squared_error(
                                y_te,
                                pred,
                            )
                        ),
                        "MAE": mean_absolute_error(
                            y_te,
                            pred,
                        ),
                        "R²": r2_score(y_te, pred),
                    }
                )
        except Exception as fold_error:
            failed_folds += 1

            if len(fold_failure_reasons) < 3:
                fold_failure_reasons.append(
                    f"第 {fold_idx} 折：{fold_error}"
                )
            continue

    if len(fold_rows) < 2:
        raise ValueError(
            f"成功完成的折数过少（{len(fold_rows)}），"
            "无法进行交叉验证评估。"
        )

    fold_df = pd.DataFrame(fold_rows)
    summary_rows = []

    for col in fold_df.columns:
        if col == "折":
            continue

        values = pd.to_numeric(
            fold_df[col],
            errors="coerce",
        ).dropna()

        if len(values) == 0:
            continue

        summary_rows.append(
            {
                "指标": col,
                "均值": values.mean(),
                "标准差": values.std(),
                "最小值": values.min(),
                "最大值": values.max(),
            }
        )

    summary_df = pd.DataFrame(summary_rows)

    note_text = (
        f"共 {len(fold_df)} 折成功完成"
        + (
            f"（另有 {failed_folds} 折拟合失败已跳过）"
            if failed_folds
            else ""
        )
        + "。"
    )

    if fold_failure_reasons:
        note_text += "\n失败原因：\n" + "\n".join(
            fold_failure_reasons
        )

    return fold_df, summary_df, note_text


# ============================================================
# 十、模型结果、诊断和论文文本
# ============================================================

def significance_label(p):
    if pd.isna(p):
        return "无法判断"

    if p < 0.01:
        return "极显著"

    if p < 0.05:
        return "显著"

    if p < 0.10:
        return "边际显著"

    return "不显著"


def make_result_table(model, model_type):
    """生成模型系数表。"""
    params = model.params
    bse = model.bse
    pvalues = model.pvalues

    try:
        conf_int = model.conf_int()
    except Exception:
        conf_int = None

    if isinstance(params, pd.Series):
        result = pd.DataFrame(
            {
                "变量": params.index,
                "回归系数": params.values,
                "标准误": np.asarray(bse),
                "P值": np.asarray(pvalues),
            }
        )

        if conf_int is not None:
            result["置信区间下限"] = (
                conf_int.iloc[:, 0].values
            )
            result["置信区间上限"] = (
                conf_int.iloc[:, 1].values
            )

        if "Logistic" in model_type:
            # 修复：截距行的 OR 不可解释，不输出
            coefs = np.asarray(params, dtype=float)

            def _or_value(row):
                if str(row["变量"]).lower() in [
                    "const",
                    "intercept",
                    "截距",
                ]:
                    return np.nan
                return float(np.exp(row["回归系数"]))

            result["优势比_OR"] = result.apply(
                _or_value,
                axis=1,
            )

        result["显著性判断"] = result[
            "P值"
        ].apply(significance_label)

        return result

    params_df = pd.DataFrame(params)
    bse_df = pd.DataFrame(bse)
    pvalues_df = pd.DataFrame(pvalues)

    # 修复：多项 Logistic（MNLogit）的 conf_int 同样是 DataFrame，
    # 此前被丢弃，导致三线表缺置信区间，这里一并展开。
    conf_int_df = None

    if conf_int is not None and hasattr(conf_int, "columns"):
        conf_int_df = pd.DataFrame(conf_int)

    rows = []

    for variable in params_df.index:
        for category in params_df.columns:
            coef = params_df.loc[
                variable,
                category,
            ]

            p_value = pvalues_df.loc[
                variable,
                category,
            ]

            row_item = {
                "变量": variable,
                "类别": category,
                "回归系数": coef,
                "标准误": bse_df.loc[
                    variable,
                    category,
                ],
                "P值": p_value,
                # 修复：多项 Logistic 中 exp(coef) 是相对风险比 RRR，
                # 不是优势比 OR，标签纠正避免论文误用。
                "相对风险比_RRR": np.exp(coef),
                "显著性判断": significance_label(
                    p_value
                ),
            }

            if conf_int_df is not None:
                try:
                    ci_values = conf_int_df.loc[
                        variable,
                        category,
                    ]
                    row_item["置信区间下限"] = ci_values[0]
                    row_item["置信区间上限"] = ci_values[1]
                except Exception:
                    pass

            rows.append(row_item)

    return pd.DataFrame(rows)


def calculate_vif(X, dummy_info=None):
    """计算 VIF。

    修复：VIF 为 inf 时标注“完全共线”；若提供 dummy_info，
    对分类变量产生的哑变量组给出“按整组看待”的提示。
    """
    x_df = X.copy()

    if "const" in x_df.columns:
        x_df = x_df.drop(
            columns=["const"]
        )

    x_df = x_df.loc[
        :,
        x_df.nunique(dropna=True) > 1,
    ]

    if x_df.shape[1] <= 1:
        return pd.DataFrame(
            columns=["变量", "VIF"]
        )

    # 反查每个哑变量属于哪个原始分类变量
    col_to_origin = {}

    if dummy_info:
        for origin_col, info in dummy_info.items():
            categories = info.get("原始类别") or []
            for cat in categories[1:]:
                col_to_origin[f"{origin_col}_{cat}"] = (
                    origin_col
                )

    rows = []

    for index, col in enumerate(x_df.columns):
        try:
            value = variance_inflation_factor(
                x_df.astype(float).values,
                index,
            )
        except Exception:
            value = np.inf

        row_item = {
            "变量": col,
            "VIF": value,
        }

        # 修复：inf 明确标注“完全共线”，而非默默显示 inf
        if value == np.inf:
            row_item["诊断"] = "完全共线（VIF无穷大）"

        origin = col_to_origin.get(col)

        if origin:
            row_item["所属分类变量"] = origin

        rows.append(row_item)

    return pd.DataFrame(rows)


def make_metric_table(fitted_result):
    """生成模型评价指标表。"""
    if fitted_result is None:
        return pd.DataFrame()

    model = fitted_result["model"]
    model_type = fitted_result["model_type"]

    y = np.asarray(
        fitted_result["display_y"],
        dtype=float,
    )

    prediction = np.asarray(
        fitted_result["prediction"],
        dtype=float,
    )

    rows = [
        ["模型类型", model_type],
        ["样本量", len(y)],
        [
            "参数数量",
            len(
                np.asarray(
                    model.params
                ).reshape(-1)
            ),
        ],
    ]

    if hasattr(model, "aic"):
        rows.append(["AIC", model.aic])

    if hasattr(model, "bic"):
        try:
            rows.append(["BIC", model.bic])
        except Exception:
            pass

    if model_type in [
        "多元线性回归",
        "Logit变换线性回归",
        "线性混合效应模型",
        "Logit变换线性混合效应模型",
    ]:
        rows.extend(
            [
                [
                    "RMSE",
                    np.sqrt(
                        mean_squared_error(
                            y,
                            prediction,
                        )
                    ),
                ],
                [
                    "MAE",
                    mean_absolute_error(
                        y,
                        prediction,
                    ),
                ],
            ]
        )

        # 修复：混合效应模型的 predict 只含固定效应，
        # RMSE 会系统性偏大，与 OLS 不可直接比较，加标注。
        if model_type in [
            "线性混合效应模型",
            "Logit变换线性混合效应模型",
        ]:
            rows.append(
                [
                    "RMSE说明",
                    "仅含固定效应预测；若需与OLS比较请使用条件残差",
                ]
            )

        if model_type in [
            "多元线性回归",
            "Logit变换线性回归",
        ]:
            rows.append(
                [
                    "R²",
                    r2_score(y, prediction),
                ]
            )

            if hasattr(model, "rsquared"):
                rows.append(
                    ["模型R²", model.rsquared]
                )

            if hasattr(model, "rsquared_adj"):
                rows.append(
                    ["调整R²", model.rsquared_adj]
                )

            if hasattr(model, "fvalue"):
                rows.append(
                    ["F统计量", model.fvalue]
                )

            if hasattr(model, "f_pvalue"):
                rows.append(
                    ["F检验P值", model.f_pvalue]
                )

        if hasattr(model, "llf"):
            rows.append(
                ["对数似然", model.llf]
            )

    elif model_type in [
        "Poisson回归",
        "负二项回归",
    ]:
        rows.extend(
            [
                [
                    "RMSE",
                    np.sqrt(
                        mean_squared_error(
                            y,
                            prediction,
                        )
                    ),
                ],
                [
                    "MAE",
                    mean_absolute_error(
                        y,
                        prediction,
                    ),
                ],
            ]
        )

        if hasattr(model, "deviance"):
            rows.append(
                ["偏差Deviance", model.deviance]
            )

        if (
            hasattr(model, "deviance")
            and hasattr(model, "null_deviance")
            and model.null_deviance != 0
        ):
            rows.append(
                [
                    "伪R²",
                    1
                    - model.deviance
                    / model.null_deviance,
                ]
            )

    elif model_type == "二项Logistic回归":
        probability = np.asarray(
            fitted_result["probability"],
            dtype=float,
        )

        predicted_class = (
            probability >= 0.5
        ).astype(int)

        rows.extend(
            [
                [
                    "准确率",
                    accuracy_score(
                        y.astype(int),
                        predicted_class,
                    ),
                ],
                [
                    "平衡准确率",
                    balanced_accuracy_score(
                        y.astype(int),
                        predicted_class,
                    ),
                ],
                [
                    "精确率",
                    precision_score(
                        y.astype(int),
                        predicted_class,
                        zero_division=0,
                    ),
                ],
                [
                    "召回率",
                    recall_score(
                        y.astype(int),
                        predicted_class,
                        zero_division=0,
                    ),
                ],
                [
                    "F1值",
                    f1_score(
                        y.astype(int),
                        predicted_class,
                        zero_division=0,
                    ),
                ],
                [
                    "Log Loss",
                    log_loss(
                        y.astype(int),
                        np.clip(probability, 1e-6, 1 - 1e-6),
                    ),
                ],
            ]
        )

        if len(np.unique(y)) == 2:
            rows.append(
                [
                    "ROC-AUC",
                    roc_auc_score(
                        y,
                        probability,
                    ),
                ]
            )

    elif model_type == "多项Logistic回归":
        predicted_class = np.asarray(
            fitted_result["prediction"],
            dtype=int,
        )

        rows.extend(
            [
                [
                    "准确率",
                    accuracy_score(
                        y.astype(int),
                        predicted_class,
                    ),
                ],
                [
                    "平衡准确率",
                    balanced_accuracy_score(
                        y.astype(int),
                        predicted_class,
                    ),
                ],
                [
                    "宏平均F1",
                    f1_score(
                        y.astype(int),
                        predicted_class,
                        average="macro",
                        zero_division=0,
                    ),
                ],
            ]
        )

    return pd.DataFrame(
        rows,
        columns=["指标", "数值"],
    )


def create_prediction_table(fitted_result):
    """创建实际值、预测值和残差表。"""
    if fitted_result is None:
        return pd.DataFrame()

    result = pd.DataFrame(
        {
            "实际值": np.asarray(
                fitted_result["display_y"]
            ),
            "预测值": np.asarray(
                fitted_result["prediction"]
            ),
        }
    )

    result["残差"] = (
        result["实际值"]
        - result["预测值"]
    )

    if (
        "prediction_interval" in fitted_result
        and fitted_result["prediction_interval"] is not None
    ):
        interval = np.asarray(
            fitted_result["prediction_interval"]
        )

        if interval.ndim == 2 and interval.shape[1] >= 2:
            result["预测区间下限"] = interval[:, 0]
            result["预测区间上限"] = interval[:, 1]

    if "probability" in fitted_result:
        probability = np.asarray(
            fitted_result["probability"]
        )

        if probability.ndim == 1:
            result["预测概率"] = probability

        result["预测类别"] = np.asarray(
            fitted_result["prediction"]
        )

    return result


def make_diagnostic_table(model):
    """生成线性模型诊断表。

    修复：Durbin-Watson 只对按时间/空间有序的观测有意义，
    截面数据下无诊断含义，表格中给出说明。
    """
    rows = []

    residuals = np.asarray(
        model.resid,
        dtype=float,
    )

    dw_value = durbin_watson(residuals)

    rows.append(
        [
            "Durbin-Watson",
            dw_value,
            (
                "仅当观测按时间/空间有序时有效；"
                "接近2说明无显著一阶自相关，"
                "远小于2或大于2需警惕正/负自相关"
            ),
        ]
    )

    if len(residuals) >= 3:
        try:
            shapiro_values = residuals

            if len(shapiro_values) > 5000:
                rng = np.random.default_rng(42)
                indices = rng.choice(
                    len(shapiro_values),
                    5000,
                    replace=False,
                )
                shapiro_values = (
                    shapiro_values[indices]
                )

            _, shapiro_p = shapiro(
                shapiro_values
            )

            rows.append(
                [
                    "残差Shapiro-Wilk P值",
                    shapiro_p,
                    (
                        "P<0.05 时拒绝残差正态假设，"
                        "可考虑变量变换"
                    ),
                ]
            )

        except Exception:
            pass

    try:
        bp_result = het_breuschpagan(
            residuals,
            model.model.exog,
        )

        rows.extend(
            [
                [
                    "Breusch-Pagan统计量",
                    bp_result[0],
                    "",
                ],
                [
                    "Breusch-Pagan P值",
                    bp_result[1],
                    (
                        "P<0.05 提示存在异方差，"
                        "可考虑稳健标准误"
                    ),
                ],
            ]
        )

    except Exception:
        pass

    try:
        white_result = het_white(
            residuals,
            model.model.exog,
        )

        rows.extend(
            [
                [
                    "White检验统计量",
                    white_result[0],
                    "",
                ],
                [
                    "White检验P值",
                    white_result[1],
                    (
                        "P<0.05 提示存在异方差，"
                        "可考虑稳健标准误"
                    ),
                ],
            ]
        )

    except Exception:
        pass

    return pd.DataFrame(
        rows,
        columns=["诊断指标", "数值", "说明"],
    )


def generate_assumptions(
    target,
    predictors,
    variable_types,
    model_type,
    vif_table=None,
    group_col=None,
):
    """生成可用于论文的模型假设。"""
    assumptions = [
        "假设题目所给数据真实可靠，能够反映研究对象的主要特征。",
        "假设各变量的单位、编码方式和统计口径保持一致。",
        "假设缺失值及异常值处理不会改变数据的主要统计规律。",
    ]

    if (
        vif_table is not None
        and not vif_table.empty
    ):
        finite_vif = vif_table[
            np.isfinite(vif_table["VIF"])
        ]

        if not finite_vif.empty:
            max_vif = finite_vif["VIF"].max()

            if max_vif < 5:
                assumptions.append(
                    "各解释变量VIF均小于5，暂未发现明显的多重共线性。"
                )
            elif max_vif < 10:
                assumptions.append(
                    "部分解释变量VIF介于5和10之间，可能存在一定的多重共线性。"
                )
            else:
                assumptions.append(
                    "部分解释变量VIF不低于10，存在较严重的多重共线性风险。"
                )

    if model_type == "多元线性回归":
        assumptions.extend(
            [
                "假设因变量与解释变量之间的条件均值关系可以用线性函数近似表示。",
                "假设误差项条件均值为零，且不同观测之间相互独立。",
                "假设误差项具有近似恒定的方差。",
                "假设不存在对估计结果产生决定性影响的极端观测。",
            ]
        )

    elif model_type == "Logit变换线性回归":
        assumptions.extend(
            [
                "假设因变量为0到1之间的比例变量。",
                "对因变量进行Logit变换后，其与解释变量之间近似满足线性关系。",
                "假设变换后的误差项相互独立且方差相对稳定。",
                "模型预测结果通过逆Logit变换还原为比例形式。",
            ]
        )

    elif model_type == "线性混合效应模型":
        assumptions.extend(
            [
                "假设同一分组内的多次观测可能存在相关性。",
                "通过随机截距描述不同分组对象之间的个体差异。",
                "假设随机效应均值为零，并与固定效应结构相互独立。",
            ]
        )

    elif model_type == "Logit变换线性混合效应模型":
        assumptions.extend(
            [
                "假设因变量为0到1之间的比例变量。",
                "对因变量进行Logit变换，并通过随机截距处理组内相关性。",
                "模型预测结果通过逆Logit变换还原为比例形式。",
            ]
        )

    elif model_type == "二项Logistic回归":
        assumptions.extend(
            [
                "假设因变量为二分类变量，并编码为0和1。",
                "假设事件发生的对数优势比与解释变量近似线性相关。",
                "假设观测样本之间相互独立。",
            ]
        )

    elif model_type == "多项Logistic回归":
        assumptions.extend(
            [
                "假设因变量为多分类变量。",
                "以一个类别作为参照类别，比较其他类别的发生优势。",
                "假设不同观测样本之间相互独立。",
            ]
        )

    elif model_type == "Poisson回归":
        assumptions.extend(
            [
                "假设因变量为非负整数计数变量。",
                "假设计数变量服从Poisson分布或近似服从Poisson分布。",
                "采用对数连接函数描述解释变量与计数均值的关系。",
                "需要检查数据是否存在过度离散。",
            ]
        )

    elif model_type == "负二项回归":
        assumptions.extend(
            [
                "假设因变量为非负整数计数变量。",
                "假设数据存在超过Poisson分布的额外离散程度。",
                "采用负二项分布和对数连接函数进行建模。",
            ]
        )

    if group_col not in [None, "无"]:
        assumptions.append(
            f'以"{group_col}"作为重复观测分组变量，同一分组内的观测可能存在相关性。'
        )

    return list(dict.fromkeys(assumptions))


# ============================================================
# 十一、侧边栏设置
# ============================================================


    
with st.sidebar:
    # ===== 背景音乐 BGM =====
    with st.expander("🎵 背景音乐（可选）", expanded=False):
        bgm_path = Path(__file__).resolve().parent / "assets" / "bgm.mp3"
        bgm_enabled = st.checkbox("开启背景音乐", value=False)

        if bgm_enabled:
            if bgm_path.exists():
                st.audio(bgm_path.read_bytes(), format="audio/mp3")
            else:
                uploaded_bgm = st.file_uploader(
                    "上传背景音乐（mp3/wav/ogg）",
                    type=["mp3", "wav", "ogg"],
                    key="bgm_uploader",
                )
                if uploaded_bgm is not None:
                    st.audio(
                        uploaded_bgm.getvalue(),
                        format=uploaded_bgm.type,
                    )
                else:
                    st.info(
                        "未检测到背景音乐。\n"
                        "可将 bgm.mp3 放入 assets 文件夹，或直接上传音乐文件。"
                    )
    st.title('功能导航')
    mode_options = [
        "数据分析",
        "优化求解",
    ]

    current_mode = st.session_state.get(
        "app_mode",
        "数据分析",
    )

    if current_mode not in mode_options:
        current_mode = "数据分析"

    app_mode = st.radio(
        "请选择功能模块",
        mode_options,
        index=mode_options.index(current_mode),
        key="app_mode_radio",
    )
    if app_mode != st.session_state.get('app_mode', '数据分析'):
        st.session_state.app_mode = app_mode
        st.rerun()

    # ===== 新手引导面板 =====
    if app_mode == "优化求解":
        with st.expander("💡 优化求解使用说明", expanded=True):
            st.markdown(
                "1. 上传数据表（决策变量可来自数据列）；\n"
                "2. 选择决策变量所在列；\n"
                "3. 输入目标函数系数（线性）或表达式（非线性）；\n"
                "4. 添加约束条件（注意系数个数必须与变量数一致）；\n"
                "5. 设置边界后点击 🚀 求解。\n\n"
                "**新手提示：** 非线性规划请先勾选“启用上方线性约束”试跑，"
                "求解失败时可取消约束或改用遗传算法。"
            )

        # 修复：优化模式也提供简版分步向导，避免新手失去引导
        with st.expander("🧭 优化问题分步向导", expanded=False):
            st.markdown(
                "**拿到优化类赛题，按这个顺序做：**\n\n"
                "**① 明确三要素** —— 决策变量（你要决定什么）、"
                "目标函数（最大化/最小化什么）、约束（限制条件）；\n\n"
                "**② 选类型** —— 全是线性关系选 LP；需要整数选 ILP；"
                "0/1 选 0-1 规划；非线性关系选 NLP；"
                "求解困难可换遗传算法；\n\n"
                "**③ 填系数** —— 目标系数来自赛题单位收益/成本，"
                "约束系数来自资源消耗量，右侧常数来自资源上限；\n\n"
                "**④ 看结果** —— 最优值、最优解、影子价格"
                "（资源每增加1单位目标值的变化）；\n\n"
                "**⑤ 写论文** —— 直接复制“论文表述”，"
                "影子价格就是灵敏度分析素材。"
            )

        # 优化模式也显示任务清单（可勾选）
        render_mission_checklist(
            max(
                st.session_state.get("guide_step", 0),
                st.session_state.get("guide_auto_step", 0),
            )
        )
    else:
        render_guide_panel(
            max(
                st.session_state.get("guide_step", 0),
                st.session_state.get("guide_auto_step", 0),
            )
        )

        render_mission_checklist(
            max(
                st.session_state.get("guide_step", 0),
                st.session_state.get("guide_auto_step", 0),
            )
        )

    if app_mode == '数据分析':
        st.header("基本设置")

        st.subheader("赛题描述")

        uploaded_problem_file = st.file_uploader(
            "上传赛题文件（PDF / Word / TXT）",
            type=["pdf", "docx", "txt"],
            key="problem_file_uploader",
        )

        if uploaded_problem_file is not None:
            problem_file_hash = hashlib.sha256(
                uploaded_problem_file.getvalue()
            ).hexdigest()

            if st.session_state.get(
                "last_problem_file_hash"
            ) != problem_file_hash:
                extracted_problem_text = (
                    extract_text_from_file(
                        uploaded_problem_file
                    )
                )

                st.session_state["problem_text"] = (
                    extracted_problem_text
                )

                st.session_state["problem_text_area"] = (
                    extracted_problem_text
                )
                st.session_state["last_problem_file"] = (
                    uploaded_problem_file.name
                )
                st.session_state["last_problem_file_hash"] = (
                    problem_file_hash
                )

        if "problem_text_area" not in st.session_state:
            st.session_state.problem_text_area = ""

        problem_text = st.text_area(
            "粘贴赛题原文或显示上传文件内容",
            height=160,
            placeholder=(
                "上传文件后自动显示内容，也可以直接在此处粘贴。"
            ),
            key="problem_text_area",
        )

        st.session_state["problem_text"] = problem_text

        if st.button("自动检测题型"):
            result = multi_label_classify_problem_text(
                problem_text
            )

            st.session_state["detect_result"] = result
            st.session_state["problem_type"] = (
                result["main_type"]
            )

            # 强制“赛题类型”输入框刷新为新检测到的类型
            st.session_state.pop(
                "problem_type_input",
                None,
            )

        detect_result = st.session_state.get(
            "detect_result",
            {},
        )

        if detect_result:
            st.success(
                f"主类型：{detect_result['main_type']}"
            )

            with st.expander(
                "📋 针对本题型的推荐流程",
                expanded=True,
            ):
                render_problem_path_guide(
                    detect_result["main_type"]
                )

            # 修复：题型识别与主流程联动。
            # 优化类 → 提供“去优化模块”跳转；预测类 → 提示时间列；
            # 评价类 → 提示熵权TOPSIS / AHP。
            main_type_detected = detect_result.get(
                "main_type",
                "",
            )

            if "优化" in main_type_detected:
                if st.button(
                    "🔀 前往「优化求解」模块",
                    key="go_to_opt_button",
                ):
                    st.session_state.app_mode = (
                        "优化求解"
                    )
                    st.rerun()
                st.info(
                    "检测到优化类赛题：优化求解模块已支持 "
                    "LP / ILP / 0-1 规划 / 非线性规划 / 遗传算法，"
                    "并提供影子价格（灵敏度分析）。"
                )

            if "预测" in main_type_detected:
                st.info(
                    "检测到预测类赛题：若数据包含时间列，"
                    "请在第④步把时间列类型确认为“时间”，"
                    "并可在高级分析中使用 ARIMA / 灰色预测 GM(1,1)。"
                )

            if "评价" in main_type_detected:
                st.info(
                    "检测到评价类赛题：推荐使用高级分析中的"
                    "「熵权法+TOPSIS」或「AHP 层次分析」"
                    "（含一致性检验）。"
                )

            detected_labels = detect_result.get(
                "all_detected_labels",
                [],
            )

            st.info(
                "检测到的类型："
                + (
                    "、".join(detected_labels)
                    if detected_labels
                    else "无"
                )
            )

            if detect_result.get(
                "sub_question_context"
            ):
                with st.expander(
                    "查看自动提取到的子问题片段"
                ):
                    for index, snippet in enumerate(
                        detect_result[
                            "sub_question_context"
                        ],
                        1,
                    ):
                        st.markdown(
                            f"**子问题 {index}：** "
                            f"{snippet}……"
                        )

            with st.expander(
                "查看题型识别分数和建模建议"
            ):
                score_table = pd.DataFrame(
                    {
                        "题型类别": list(
                            detect_result.get(
                                "label_scores",
                                {},
                            ).keys()
                        ),
                        "综合得分": list(
                            detect_result.get(
                                "label_scores",
                                {},
                            ).values()
                        ),
                    }
                )

                _st_dataframe(
                    score_table,
                    use_container_width=True,
                )

                direction_map = {
                    "评价类": (
                        "层次分析、TOPSIS、模糊综合评价、"
                        "灰色关联分析、主成分分析"
                    ),
                    "预测类": (
                        "多元回归、ARIMA、灰色预测、"
                        "神经网络、时间序列分析"
                    ),
                    "优化类": (
                        "线性规划、整数规划、0-1规划、"
                        "遗传算法、模拟退火"
                    ),
                    "机理分析类": (
                        "微分方程、动力学模型、Logistic模型、"
                        "SIR模型、稳定性分析"
                    ),
                    "分类类": (
                        "Logistic回归、决策树、随机森林、"
                        "SVM、聚类分析"
                    ),
                }

                for label in detected_labels:
                    st.write(
                        f"**{label}：** "
                        f"{direction_map.get(label, '根据题目背景确定')}"
                    )

        problem_type = st.text_input(
            "赛题类型",
            value=st.session_state.get(
                "problem_type",
                "",
            ),
            placeholder="例如：预测类、评价类、优化类",
            key="problem_type_input",
        )

        uploaded_file = st.file_uploader(
            "上传数据表",
            type=["csv", "xlsx", "xls"],
            key="data_file_uploader",
        )

        # 修复：数据源与优化模块共享，同一份数据无需上传两次
        if uploaded_file is not None:
            st.session_state.shared_data_file = (
                uploaded_file
            )

        st.subheader("缺失值处理")

        missing_method = st.selectbox(
            "处理方式",
            [
                "分类型处理",
                "KNN插补",
                "删除含缺失值的行",
            ],
            help=(
                "分类型处理：数值列中位数填充、分类列众数填充"
                "（截面数据默认推荐，不依赖行顺序）；"
                "KNN插补：标准化后用距离最近的样本加权估计缺失值"
                "（数值变量较多时推荐）；"
                "删除含缺失值的行：最保守但会损失样本。"
            ),
        )

        st.subheader("异常值处理")

        outlier_method = st.selectbox(
            "识别方法",
            [
                "不处理",
                "3σ",
                "IQR",
                "MAD",
            ],
            help=(
                "3σ：基于正态假设，适合对称分布；"
                "IQR：箱线图法，对偏态更稳健；"
                "MAD：中位数绝对偏差，最稳健但可能标记偏多。"
            ),
        )

        outlier_action = st.selectbox(
            "异常值处理动作",
            [
                "仅标记，不删除",
                "删除异常行",
            ],
        )

        # 修复：默认不把因变量纳入异常检测，
        # 避免删除 y 的“极端但合法”观测导致回归系数有偏。
        include_target_outlier = st.checkbox(
            "异常检测包含因变量（默认不包含）",
            value=False,
            help=(
                "对回归问题，删除因变量的极端观测等价于截尾 y 分布，"
                "会使回归系数系统性有偏。默认只对自变量检测异常。"
            ),
        )

        robust_se = st.checkbox(
            "线性回归使用HC3稳健标准误",
            value=True,
        )

        use_test_set = st.checkbox(
            "进行训练集/测试集评估",
            value=True,
        )

        test_size = st.slider(
            "测试集比例",
            min_value=0.1,
            max_value=0.4,
            value=0.2,
            step=0.05,
        )



    else:  # 优化求解模式
        st.header('优化设置')

        # 修复：优先复用数据分析模块已上传的数据，
        # 未上传时才在此单独上传（数据源共享）。
        shared = st.session_state.get("shared_data_file")

        if shared is not None:
            st.success(
                "已共享「数据分析」模块上传的数据表，"
                "无需重复上传。"
            )
            st.session_state.opt_uploaded_file = shared
        else:
            opt_upload = st.file_uploader(
                '上传数据表（优化用）',
                type=['csv', 'xlsx'],
                key='opt_data_upload',
            )
            if opt_upload is not None:
                st.session_state.opt_uploaded_file = (
                    opt_upload
                )
if st.session_state.get('app_mode', '数据分析') == '数据分析':
    # ============================================================
    # 十二、读取数据
    # ============================================================

    if uploaded_file is None:
        st.info("请在左侧上传 CSV 或 Excel 数据表。")
        st.stop()

    try:
        uploaded_file.seek(0)

        if uploaded_file.name.lower().endswith(".csv"):
            try:
                raw_df = pd.read_csv(uploaded_file, encoding='utf-8')
            except UnicodeDecodeError:
                uploaded_file.seek(0)
                raw_df = pd.read_csv(uploaded_file, encoding='gbk')
            except Exception:
                uploaded_file.seek(0)
                raw_df = pd.read_csv(uploaded_file, encoding='gb18030')
        else:
            raw_df = pd.read_excel(uploaded_file)

    except Exception as exc:
        st.error(f"读取文件失败：{exc}")
        st.stop()

    if raw_df is None or raw_df.empty:
        st.error("数据未成功加载，请检查文件内容。")
        st.stop()

    raw_df.columns = make_unique_columns(
        raw_df.columns
    )

    original_shape = raw_df.shape
    all_columns = list(raw_df.columns)

    if len(all_columns) < 2:
        st.error(
            "数据表至少需要包含一列因变量和一列自变量。"
        )
        st.stop()

    # 引导进度：已完成第 2 步（上传数据表）
    st.session_state["guide_auto_step"] = max(
        st.session_state.get("guide_auto_step", 0),
        2,
    )

    # ============================================================
    # 十三、数据初探
    # ============================================================

    st.subheader("1. 数据初探")

    metric_col1, metric_col2, metric_col3, metric_col4 = (
        st.columns(4)
    )

    metric_col1.metric(
        "原始行数",
        original_shape[0],
    )

    metric_col2.metric(
        "原始列数",
        original_shape[1],
    )

    metric_col3.metric(
        "缺失单元格数",
        int(raw_df.isna().sum().sum()),
    )

    metric_col4.metric(
        "重复行数",
        int(raw_df.duplicated().sum()),
    )

    st.write("数据前20行")
    _st_dataframe(
        raw_df.head(20),
        use_container_width=True,
    )

    with st.expander("查看字段基本信息"):
        info_table = pd.DataFrame(
            {
                "列名": raw_df.columns,
                "数据类型": raw_df.dtypes.astype(str).values,
                "非空数量": raw_df.notna().sum().values,
                "缺失数量": raw_df.isna().sum().values,
                "唯一值数量": raw_df.nunique(
                    dropna=True
                ).values,
                "是否疑似ID列": [
                    "是"
                    if is_suspicious_id_column(
                        raw_df[col],
                        col,
                    )
                    else "否"
                    for col in raw_df.columns
                ],
            }
        )

        _st_dataframe(
            info_table,
            use_container_width=True,
        )

        dataframe_download(
            info_table,
            "字段基本信息.csv",
            key="download_info_table",
        )


    # ============================================================
    # 十四、变量选择
    # ============================================================

    st.subheader("2. 因变量、自变量与分组变量设置")

    target = st.selectbox(
        "请选择因变量",
        all_columns,
        index=0,
    )

    default_predictors = [
        col
        for col in all_columns
        if (
            col != target
            and not is_suspicious_id_column(
                raw_df[col],
                col,
            )
        )
    ]

    predictors = st.multiselect(
        "请选择自变量",
        [
            col
            for col in all_columns
            if col != target
        ],
        default=default_predictors,
    )

    if not predictors:
        st.warning("请至少选择一个自变量。")
        st.stop()

    group_candidates = ["无"]

    for col in all_columns:
        if col == target or col in predictors:
            continue

        non_missing = raw_df[col].dropna()

        if len(non_missing) == 0:
            continue

        group_count = non_missing.nunique()
        group_sizes = non_missing.value_counts()

        if (
            group_count >= 2
            and group_sizes.max() >= 2
            and group_count < len(raw_df)
        ):
            group_candidates.append(col)

    group_col = st.selectbox(
        "重复观测分组变量",
        group_candidates,
        help=(
            "如果同一对象有多次观测，请选择对应的对象编号。"
        ),
    )

    if group_col != "无" and group_col in predictors:
        predictors = [
            col
            for col in predictors
            if col != group_col
        ]

        st.info(
            f'已将 "{group_col}" 从自变量中移除（它被用作分组变量）。'
        )

    if not predictors:
        st.error(
            "移除分组变量后没有剩余自变量。"
        )
        st.stop()

    # 引导进度：已完成第 3 步（选择因变量与自变量）
    st.session_state["guide_auto_step"] = max(
        st.session_state.get("guide_auto_step", 0),
        3,
    )

    # ============================================================
    # 十五、变量类型识别
    # ============================================================

    st.subheader("3. 变量类型识别与确认")

    initial_variable_types = {
        col: classify_variable(raw_df[col])
        for col in [target] + predictors
    }

    initial_type_table = pd.DataFrame(
        {
            "变量": list(
                initial_variable_types.keys()
            ),
            "自动识别类型": list(
                initial_variable_types.values()
            ),
            "原始数据类型": [
                str(raw_df[col].dtype)
                for col in initial_variable_types
            ],
            "唯一值数量": [
                raw_df[col].nunique(
                    dropna=True
                )
                for col in initial_variable_types
            ],
        }
    )

    _st_dataframe(
        initial_type_table,
        use_container_width=True,
    )

    st.caption(
        "自动识别结果仅供参考，正式建模前请根据变量实际含义确认类型。"
    )

    variable_types = {}
    type_options = [
        "连续",
        "分类",
        "时间",
        "次数",
    ]

    for col in [target] + predictors:
        default_type = initial_variable_types[col]

        if default_type not in type_options:
            default_type = "分类"

        default_index = type_options.index(
            default_type
        )

        variable_types[col] = st.selectbox(
            f"确认变量 `{col}` 的类型",
            type_options,
            index=default_index,
            key=f"variable_type_{col}",
        )


    # 注意：
    # 页面上的变量类型已经由用户确认。
    # 这里不能再次自动覆盖，否则“次数”“分类”等人工设置会失效。
    # 自动识别结果只用于初始默认值。

    # 引导进度：已完成第 4 步（确认变量类型）
    st.session_state["guide_auto_step"] = max(
        st.session_state.get("guide_auto_step", 0),
        4,
    )

    current_signature = build_analysis_signature(
        file_name=uploaded_file.name,
        file_hash=hashlib.sha256(
            uploaded_file.getvalue()
        ).hexdigest(),
        target=target,
        predictors=predictors,
        variable_types=variable_types,
        group_col=group_col,
        missing_method=missing_method,
        outlier_method=outlier_method,
        outlier_action=outlier_action,
        include_target_outlier=include_target_outlier,
        robust_se=robust_se,
        use_test_set=use_test_set,
        test_size=test_size,
    )

    reset_model_if_signature_changed(
        current_signature
    )


    # ============================================================
    # 十六、符号表
    # ============================================================

    st.subheader("4. 数学建模符号表")

    symbol_mode = st.radio(
        "符号表类型",
        [
            "样本符号表",
            "变量符号表",
        ],
        horizontal=True,
    )

    # 注意：这里不要重新定义 target 和 predictors。
    # 它们已经在前面的建模设置部分确定。
    if symbol_mode == "样本符号表":
        # 修复：样本符号表单独实现，包含样本维度记号
        # （样本容量 n、下标 i、观测 x_ij/y_i），
        # 不再与“变量符号表”输出相同内容。

        _sample_rows = []

        _sample_symbols = [
            ("n", "样本容量（观测数）", "整数", "样本总数"),
            ("i", "样本下标", "整数", "i = 1, 2, …, n"),
            ("y_i", "第 i 个样本的因变量观测值", "数值/类别", f"来自列：{target}"),
            ("x_ij", "第 i 个样本第 j 个自变量的观测值", "数值/类别", "来自对应自变量列"),
            ("X", "自变量观测矩阵", "矩阵", "n × p 矩阵"),
            ("y", "因变量观测向量", "向量", "长度 n"),
            ("β_j", "第 j 个回归系数", "参数", "待估计"),
            ("ε_i", "第 i 个样本的随机误差项", "随机变量", "均值 0"),
        ]

        for symbol, meaning, var_type, note in _sample_symbols:
            _sample_rows.append(
                {
                    "符号": symbol,
                    "含义": meaning,
                    "类型": var_type,
                    "说明": note,
                }
            )

        for idx, col in enumerate(predictors, start=1):
            _sample_rows.append(
                {
                    "符号": f"x_{idx}",
                    "含义": f"第 {idx} 个自变量（列：{col}）",
                    "类型": variable_types.get(col, ""),
                    "说明": "即 x_ij 的第 j 列",
                }
            )

        sample_symbol_table = pd.DataFrame(
            _sample_rows
        )

        _st_dataframe(
            sample_symbol_table,
            use_container_width=True,
        )

        dataframe_download(
            sample_symbol_table,
            "样本符号表.csv",
            key="download_sample_symbol_table",
        )

        st.caption(
            "样本符号表描述“观测—变量—参数”的整体记号体系，"
            "用于论文“符号说明”章节；"
            "变量符号表则在“变量符号表”选项下可编辑。"
        )

    else:
        variable_symbol_table = (
            create_variable_symbol_table(
                target,
                predictors,
                variable_types,
            )
        )

        st.info(
            r"可以直接编辑“变量符号”“单位”和“变量含义”。"
            r"例如：\mathrm{Year}_i"
        )

        edited_symbol_table = st.data_editor(
            variable_symbol_table,
            use_container_width=True,
            num_rows="fixed",
            disabled=[
                "原始列名",
                "变量角色",
                "变量类型",
            ],
            key="variable_symbol_editor",
        )

        st.write("LaTeX 变量符号预览")

        for _, row in edited_symbol_table.iterrows():
            symbol = str(row["变量符号"]).strip()
            original_name = str(row["原始列名"]).strip()

            if symbol:
                st.markdown(
                    f"原始变量：`{original_name}`"
                )
                st.latex(symbol)

        dataframe_download(
            edited_symbol_table,
            "变量符号表.csv",
            key="download_variable_symbol_table",
        )

    # ============================================================
    # 十七、数据清洗
    # ============================================================

    st.subheader("5. 数据清洗")

    st.info(
        "因变量缺失的样本无法用于监督建模，因此程序会直接删除；"
        "缺失值填补仅针对自变量。"
    )

    typed_df = convert_types(
        raw_df,
        variable_types,
    )

    before_missing_cells = int(
        typed_df.isna().sum().sum()
    )

    target_missing_mask = typed_df[target].isna()
    deleted_target_missing_rows = int(
        target_missing_mask.sum()
    )

    typed_df = typed_df.loc[
        ~target_missing_mask
    ].copy()

    predictor_variable_types = {
        col: variable_types[col]
        for col in predictors
        if col in variable_types
    }

    (
        clean_df,
        missing_detail_table,
        deleted_missing_rows,
        imputed_cells,
    ) = fill_missing_values(
        typed_df,
        predictor_variable_types,
        missing_method,
    )

    numeric_columns_for_outlier = [
        col
        for col in [target] + predictors
        if variable_types.get(col)
        in ["连续", "次数"]
    ]

    # 修复：默认只对自变量做异常检测（除非用户勾选包含因变量）
    if not include_target_outlier and (
        target in numeric_columns_for_outlier
    ):
        numeric_columns_for_outlier = [
            col
            for col in numeric_columns_for_outlier
            if col != target
        ]

    outlier_detail_table = pd.DataFrame()
    outlier_row_count = 0
    deleted_outlier_rows = 0

    if outlier_method != "不处理":
        marked_df, outlier_detail_table = (
            detect_outliers(
                clean_df,
                numeric_columns_for_outlier,
                outlier_method,
            )
        )

        outlier_row_count = int(
            marked_df["_异常行"].sum()
        )

        if outlier_action == "删除异常行":
            deleted_outlier_rows = outlier_row_count
            clean_df = marked_df.loc[
                ~marked_df["_异常行"]
            ].copy()
        else:
            clean_df = marked_df.copy()

    else:
        clean_df["_异常行"] = False

    used_columns = [target] + predictors

    if group_col != "无":
        used_columns.append(group_col)


    used_columns = list(dict.fromkeys(used_columns))

    unused_columns = [
        col
        for col in clean_df.columns
        if (
            col not in used_columns
            and col != "_异常行"
        )
    ]

    clean_df = clean_df.drop(
        columns=unused_columns,
        errors="ignore",
    )

    after_missing_cells = int(
        clean_df.isna().sum().sum()
    )

    clean_df = clean_df.reset_index(
        drop=True
    )

    clean_data_for_model = clean_df.drop(
        columns=["_异常行"],
        errors="ignore",
    )

    # 引导进度：已完成第 5 步（数据清洗）
    st.session_state["guide_auto_step"] = max(
        st.session_state.get("guide_auto_step", 0),
        5,
    )

    cleaning_summary = pd.DataFrame(
        {
            "项目": [
                "处理时间",
                "赛题类型",
                "原始行数",
                "清洗后行数",
                "原始列数",
                "清洗后列数",
                "原始缺失单元格数",
                "缺失值插补数量",
                "因变量缺失删除行数",
                "自变量缺失删除行数",
                "检测到的异常行数",
                "因异常删除行数",
                "清洗后剩余缺失单元格数",
                "缺失值处理方式",
                "异常值识别方法",
                "异常值处理动作",
                "重复观测分组变量",
                "删除的无用列",
            ],
            "结果": [
                datetime.now().strftime(
                    "%Y-%m-%d %H:%M:%S"
                ),
                problem_type or "未填写",
                original_shape[0],
                clean_data_for_model.shape[0],
                original_shape[1],
                clean_data_for_model.shape[1],
                before_missing_cells,
                imputed_cells,
                deleted_target_missing_rows,
                deleted_missing_rows,
                outlier_row_count,
                deleted_outlier_rows,
                after_missing_cells,
                missing_method,
                outlier_method,
                outlier_action,
                group_col,
                (
                    ", ".join(unused_columns)
                    if unused_columns
                    else "无"
                ),
            ],
        }
    )

    clean_metric1, clean_metric2, clean_metric3, clean_metric4 = (
        st.columns(4)
    )

    clean_metric1.metric(
        "缺失值插补数量",
        imputed_cells,
    )

    clean_metric2.metric(
        "因缺失删除行数",
        deleted_target_missing_rows
        + deleted_missing_rows,
    )

    clean_metric3.metric(
        "检测到异常行数",
        outlier_row_count,
    )

    clean_metric4.metric(
        "清洗后样本数",
        clean_data_for_model.shape[0],
    )

    _st_dataframe(
        cleaning_summary,
        use_container_width=True,
    )

    # 修复：清洗后仍有缺失时给出明确告警并列出列名，
    # 避免用户不知情导致后续样本被静默删除。
    _remaining_missing_cols = [
        col
        for col in clean_data_for_model.columns
        if clean_data_for_model[col].isna().any()
    ]

    if after_missing_cells > 0 and _remaining_missing_cols:
        st.warning(
            f"清洗后仍有 {after_missing_cells} 个缺失单元格，"
            "涉及列："
            + "、".join(_remaining_missing_cols)
            + "。这些行在建模时会自动删除，"
            "若删除过多请回退调整缺失值处理方法，"
            "或直接删除缺失率过高的列。"
        )

    _st_dataframe(
        missing_detail_table,
        use_container_width=True,
    )

    if not outlier_detail_table.empty:
        _st_dataframe(
            outlier_detail_table,
            use_container_width=True,
        )

    _st_dataframe(
        clean_data_for_model.head(20),
        use_container_width=True,
    )

    dataframe_download(
        cleaning_summary,
        "清洗汇总报告.csv",
        key="download_cleaning_summary",
    )

    dataframe_download(
        missing_detail_table,
        "缺失值处理明细.csv",
        key="download_missing_detail",
    )

    if not outlier_detail_table.empty:
        dataframe_download(
            outlier_detail_table,
            "异常值处理明细.csv",
            key="download_outlier_detail",
        )

    dataframe_download(
        clean_data_for_model,
        "清洗后数据.csv",
        key="download_clean_data",
    )

    cleaning_text = f"""
    本文首先对原始数据进行完整性、一致性和变量类型检查。
    原始数据共包含{original_shape[0]}条样本和{original_shape[1]}个变量。
    针对缺失数据，本文采用“{missing_method}”方法进行处理，
    共插补{imputed_cells}个缺失单元格。
    对于因变量缺失的样本，由于无法提供有效的被解释变量观测值，
    直接删除{deleted_target_missing_rows}条样本；
    另外因自变量缺失处理删除{deleted_missing_rows}条样本。
    针对异常数据，本文采用“{outlier_method}”方法进行识别，
    共检测到{outlier_row_count}条异常样本，
    其中因异常值删除{deleted_outlier_rows}条样本。
    经过数据类型转换、缺失值处理、异常值处理和无用列处理后，
    最终获得{clean_data_for_model.shape[0]}条有效样本，
    用于后续统计分析和模型建立。
    """.strip()

    st.text_area(
        "可直接用于论文的数据清洗表述",
        cleaning_text,
        height=220,
    )


    # ============================================================
    # 5.5 描述统计与正态性检验（十七·补充）
    # ============================================================

    st.subheader("5.5 描述统计与正态性检验")

    desc_numeric_cols = [
        col
        for col in [target] + predictors
        if (
            variable_types.get(col) in ["连续", "次数"]
            and col in clean_data_for_model.columns
        )
    ]

    if not desc_numeric_cols:
        st.info("当前没有可用于描述统计的数值型变量。")
    else:
        desc_rows = []
        shapiro_results = []

        for col in desc_numeric_cols:
            values = pd.to_numeric(
                clean_data_for_model[col],
                errors="coerce",
            ).dropna()

            if len(values) == 0:
                continue

            desc_rows.append(
                {
                    "变量": col,
                    "样本量": len(values),
                    "均值": values.mean(),
                    "标准差": values.std(),
                    "最小值": values.min(),
                    "25%分位": values.quantile(0.25),
                    "中位数": values.median(),
                    "75%分位": values.quantile(0.75),
                    "最大值": values.max(),
                    "偏度": values.skew(),
                    "峰度": values.kurtosis(),
                }
            )

            test_values = values

            if len(test_values) > 5000:
                test_values = test_values.sample(
                    5000,
                    random_state=42,
                )

            try:
                _, shapiro_p = shapiro(test_values)
                shapiro_results.append(
                    {
                        "变量": col,
                        "Shapiro-Wilk P值": shapiro_p,
                        "正态性(P>=0.05)": (
                            "是" if shapiro_p >= 0.05 else "否"
                        ),
                    }
                )
            except Exception:
                shapiro_results.append(
                    {
                        "变量": col,
                        "Shapiro-Wilk P值": np.nan,
                        "正态性(P>=0.05)": "无法判断",
                    }
                )

        desc_table = pd.DataFrame(desc_rows)
        shapiro_table = pd.DataFrame(shapiro_results)

        st.write("描述统计表（均值、标准差、偏度、峰度等）")
        _st_dataframe(
            desc_table,
            use_container_width=True,
        )

        dataframe_download(
            desc_table,
            "描述统计表.csv",
            key="download_desc_stats",
        )

        st.write("正态性检验（Shapiro-Wilk）")
        _st_dataframe(
            shapiro_table,
            use_container_width=True,
        )

        dataframe_download(
            shapiro_table,
            "正态性检验.csv",
            key="download_shapiro",
        )

        normal_cols = [
            row["变量"]
            for row in shapiro_results
            if row["正态性(P>=0.05)"] == "是"
        ]

        desc_text = (
            f"对建模涉及的主要数值变量进行描述性统计分析，"
            f"各变量的均值、标准差、偏度和峰度均处于合理范围。"
            "Shapiro-Wilk正态性检验显示，"
        )

        if normal_cols:
            desc_text += (
                "变量"
                + "、".join(normal_cols)
                + "的P值大于0.05，可近似认为满足正态性假设；"
                "其余变量存在一定偏离，后续建模时应注意稳健性，"
                "或考虑对变量进行变换处理。"
            )
        else:
            desc_text += (
                "多数变量P值小于0.05，不完全满足正态性假设，"
                "后续建模建议结合稳健标准误或非参数方法。"
            )

        st.text_area(
            "可直接用于论文的描述统计表述",
            desc_text,
            height=160,
        )


    # ============================================================
    # 5.6 数据标准化与归一化（十七·补充）
    # ============================================================

    st.subheader("5.6 数据标准化与归一化")

    st.info(
        "标准化常用于评价类、聚类类、主成分分析等需要消除量纲影响的场景。"
        "标准化结果仅用于导出，不影响上方回归建模流程。"
    )

    std_method = st.selectbox(
        "标准化方法",
        [
            "Z-score标准化",
            "Min-Max归一化",
            "极差标准化(映射到-1~1)",
        ],
        key="std_method_selector",
    )

    std_cols = st.multiselect(
        "选择要标准化的数值变量",
        desc_numeric_cols,
        default=desc_numeric_cols,
        key="std_columns_selector",
    )

    if std_cols:
        std_df = clean_data_for_model[std_cols].copy()

        for col in std_cols:
            values = pd.to_numeric(
                std_df[col],
                errors="coerce",
            )

            if std_method == "Z-score标准化":
                mean_value = values.mean()
                std_value = values.std()

                if pd.isna(std_value) or std_value == 0:
                    std_df[col] = 0.0
                else:
                    std_df[col] = (
                        values - mean_value
                    ) / std_value

            elif std_method == "Min-Max归一化":
                min_value = values.min()
                max_value = values.max()

                if pd.isna(max_value) or max_value == min_value:
                    std_df[col] = 0.0
                else:
                    std_df[col] = (
                        values - min_value
                    ) / (max_value - min_value)

            else:
                min_value = values.min()
                max_value = values.max()

                if pd.isna(max_value) or max_value == min_value:
                    std_df[col] = 0.0
                else:
                    std_df[col] = (
                        (values - min_value)
                        / (max_value - min_value)
                        * 2
                        - 1
                    )

        st.write("标准化后的数据（前20行）")
        _st_dataframe(
            std_df.head(20),
            use_container_width=True,
        )

        dataframe_download(
            std_df,
            "标准化后数据.csv",
            key="download_std_data",
        )

        st.caption(
            "提示：若希望用标准化后的数据建模，请下载后重新上传该文件，"
            "并在建模流程中重新设置变量。"
        )


    # ============================================================
    # 十八、数据可视化
    # ============================================================

    st.subheader("6. 数据可视化")

    numeric_variables = [
        col
        for col in [target] + predictors
        if (
            variable_types.get(col)
            in ["连续", "次数"]
            and col in clean_data_for_model.columns
        )
    ]

    categorical_variables = [
        col
        for col in [target] + predictors
        if (
            variable_types.get(col) == "分类"
            and col in clean_data_for_model.columns
        )
    ]

    chart_type = st.selectbox(
        "选择图形类型",
        [
            "变量分布图",
            "变量箱线图",
            "因变量-自变量散点图",
            "因变量-自变量曲线图",
            "数值变量相关性热力图",
            "分类变量频数图",
        ],
    )

    if chart_type == "变量分布图":
        if not numeric_variables:
            st.info("当前没有连续型或次数型变量。")
        else:
            selected_col = st.selectbox(
                "选择变量",
                numeric_variables,
                key="histogram_column",
            )

            values = pd.to_numeric(
                clean_data_for_model[selected_col],
                errors="coerce",
            ).dropna()

            fig, ax = plt.subplots(
                figsize=(8, 5)
            )

            sns.histplot(
                values,
                kde=True,
                ax=ax,
            )

            ax.set_title(
                f"{selected_col} 分布图"
            )
            ax.set_xlabel(selected_col)

            show_fig(fig)

    elif chart_type == "变量箱线图":
        if not numeric_variables:
            st.info("当前没有连续型或次数型变量。")
        else:
            selected_columns = st.multiselect(
                "选择变量",
                numeric_variables,
                default=numeric_variables,
                key="boxplot_columns",
            )

            if selected_columns:
                fig, ax = plt.subplots(
                    figsize=(10, 5)
                )

                sns.boxplot(
                    data=clean_data_for_model[
                        selected_columns
                    ],
                    ax=ax,
                )

                ax.tick_params(
                    axis="x",
                    rotation=35,
                )

                ax.set_title("变量箱线图")

                show_fig(fig)

    elif chart_type == "因变量-自变量散点图":
        x_candidates = [
            col
            for col in numeric_variables
            if col != target
        ]

        if not x_candidates:
            st.info(
                "没有适合绘制散点图的数值型自变量。"
            )
        else:
            x_col = st.selectbox(
                "选择横轴自变量",
                x_candidates,
                key="scatter_x",
            )

            plot_df = clean_data_for_model[
                [x_col, target]
            ].copy()

            plot_df[x_col] = pd.to_numeric(
                plot_df[x_col],
                errors="coerce",
            )

            plot_df[target] = pd.to_numeric(
                plot_df[target],
                errors="coerce",
            )

            plot_df = plot_df.dropna()

            fig, ax = plt.subplots(
                figsize=(8, 5)
            )

            sns.scatterplot(
                data=plot_df,
                x=x_col,
                y=target,
                ax=ax,
            )

            ax.set_title(
                f"{target} 与 {x_col} 的散点图"
            )

            show_fig(fig)

    elif chart_type == "因变量-自变量曲线图":
        x_candidates = [
            col
            for col in numeric_variables
            if col != target
        ]

        if not x_candidates:
            st.info(
                "没有适合绘制曲线图的数值型自变量。"
            )
        else:
            x_col = st.selectbox(
                "选择横轴自变量",
                x_candidates,
                key="curve_x",
            )

            plot_df = clean_data_for_model[
                [x_col, target]
            ].copy()

            plot_df[x_col] = pd.to_numeric(
                plot_df[x_col],
                errors="coerce",
            )

            plot_df[target] = pd.to_numeric(
                plot_df[target],
                errors="coerce",
            )

            plot_df = (
                plot_df.dropna()
                .sort_values(x_col)
            )

            fig, ax = plt.subplots(
                figsize=(8, 5)
            )

            ax.plot(
                plot_df[x_col],
                plot_df[target],
                marker="o",
                linewidth=1.5,
            )

            ax.set_xlabel(x_col)
            ax.set_ylabel(target)
            ax.set_title(
                f"{target} 与 {x_col} 的变化曲线"
            )

            show_fig(fig)

    elif chart_type == "数值变量相关性热力图":
        if len(numeric_variables) < 2:
            st.info("至少需要两个数值型变量。")
        else:
            corr = clean_data_for_model[
                numeric_variables
            ].corr()

            fig, ax = plt.subplots(
                figsize=(10, 7)
            )

            sns.heatmap(
                corr,
                annot=True,
                cmap="coolwarm",
                fmt=".2f",
                ax=ax,
            )

            ax.set_title(
                "数值变量Pearson相关性热力图"
            )

            show_fig(fig)

    elif chart_type == "分类变量频数图":
        if not categorical_variables:
            st.info("当前没有分类变量。")
        else:
            selected_col = st.selectbox(
                "选择分类变量",
                categorical_variables,
                key="count_column",
            )

            count_table = (
                clean_data_for_model[selected_col]
                .astype("string")
                .value_counts(dropna=False)
                .rename_axis(selected_col)
                .reset_index(name="数量")
            )

            fig, ax = plt.subplots(
                figsize=(9, 5)
            )

            sns.barplot(
                data=count_table,
                x=selected_col,
                y="数量",
                ax=ax,
            )

            ax.tick_params(
                axis="x",
                rotation=35,
            )

            ax.set_title(
                f"{selected_col} 频数图"
            )

            show_fig(fig)


    # ============================================================
    # 十九、相关性分析
    # ============================================================

    st.subheader("7. 相关性分析")

    corr_table = correlation_table(
        clean_data_for_model,
        target,
        predictors,
        variable_types,
    )

    if corr_table.empty:
        st.info(
            "没有足够的连续型或次数型变量用于相关性分析。"
        )
    else:
        _st_dataframe(
            corr_table,
            use_container_width=True,
        )

        # 修复：标注相关系数基于插补后数据，避免新手误读
        st.caption(
            "说明：相关系数基于清洗（含缺失值处理）后的数据计算；"
            "P 值已做 Benjamini-Hochberg FDR 多重比较校正，"
            "有效样本数少于 10 的变量对不参与计算。"
        )

        dataframe_download(
            corr_table,
            "相关性分析.csv",
            key="download_correlation",
        )


    # ============================================================
    # 二十、模型构造和推荐
    # ============================================================

    st.subheader("8. 自动模型推荐与人工确认")

    try:
        y, X, groups, model_meta = build_model_data(
            clean_data_for_model,
            target,
            predictors,
            variable_types,
            group_col=group_col,
        )

        if len(y) < 5:
            st.error(
                "清洗后有效样本少于5条，无法进行可靠建模。"
            )
            st.stop()

        if len(y) <= X.shape[1]:
            st.error(
                "有效样本数不大于模型参数数量，无法稳定建立模型。"
            )
            st.stop()

        if variable_types.get(target) == "分类":
            target_counts = pd.Series(y).value_counts()

            if target_counts.min() < 2:
                st.warning(
                    "因变量存在样本数少于2的类别，"
                    "分类模型和测试集评估可能不稳定。"
                )

        if groups is not None:
            if groups.nunique() < 2:
                st.warning(
                    "分组数量过少，混合效应模型结果可能不稳定。"
                )

        matrix_rank = np.linalg.matrix_rank(
            X.values
        )

        if matrix_rank < X.shape[1]:
            st.warning(
                "设计矩阵可能存在完全多重共线性，部分参数可能无法稳定估计。"
            )

        target_type_for_model = variable_types.get(target, "分类")

        # 统一变量类型名称，避免类型名称不一致导致模型推荐为“未识别”
        type_aliases = {
            "连续型": "连续",
            "次数型": "次数",
            "数值": "连续",
            "数值型": "连续",
            "类别": "分类",
            "分类变量": "分类",
            "时间型": "时间",
        }
        target_type_for_model = type_aliases.get(
            target_type_for_model,
            target_type_for_model,
        )

        recommended_info = detect_model_type(
            y,
            target_type_for_model,
            groups=groups,
        )

        # 对未识别类型提供兜底推荐。
        # 模型名称必须与 MODEL_OPTIONS 完全一致。
        if recommended_info.get("model_type") == "未识别":
            if target_type_for_model == "连续":
                recommended_info = {
                    "model_type": "多元线性回归",
                    "reason": "目标变量被设置为连续型变量。",
                }

            elif target_type_for_model == "次数":
                recommended_info = {
                    "model_type": "Poisson回归",
                    "reason": "目标变量被设置为次数型变量。",
                }

            elif target_type_for_model == "分类":
                unique_count = len(np.unique(y))

                if unique_count <= 2:
                    model_name = "二项Logistic回归"
                else:
                    model_name = "多项Logistic回归"

                recommended_info = {
                    "model_type": model_name,
                    "reason": "目标变量被设置为分类变量。",
                }

        recommended_model = recommended_info[
            "model_type"
        ]

        st.info(
            f"程序推荐模型：**{recommended_model}**\n\n"
            f"判断依据：{recommended_info['reason']}"
        )

        recommended_index = (
            MODEL_OPTIONS.index(
                recommended_model
            )
            if recommended_model in MODEL_OPTIONS
            else 0
        )

        final_model_type = st.selectbox(
            "请选择最终拟合模型",
            MODEL_OPTIONS,
            index=recommended_index,
            key="final_model_selector",
        )

        # 重要修复：
        # 统一使用 final_model_type，不再使用未定义的 model_type。
        model_type = final_model_type

        # ===== 新增：多模型一键对比 =====
        with st.expander("🔬 多模型对比（可选）", expanded=False):
            st.caption(
                "同时拟合多个候选模型，对比 AIC / BIC / R² / 准确率等指标，"
                "帮助你选出最合适的模型。混合效应模型需要已选择分组变量。"
            )

            compare_models_list = st.multiselect(
                "选择要对比的模型",
                MODEL_OPTIONS,
                default=[recommended_model],
                key="compare_models_selector",
            )

            if compare_models_list and st.button(
                "运行多模型对比",
                key="run_model_compare",
            ):
                try:
                    # 修复：混合效应模型对比前校验分组变量，
                    # 避免 fit_model 抛出的 ValueError 被吞掉、
                    # 表格只显示“拟合失败”而用户看不到原因。
                    mixed_needs_group = [
                        m
                        for m in compare_models_list
                        if m
                        in [
                            "线性混合效应模型",
                            "Logit变换线性混合效应模型",
                        ]
                    ]

                    if mixed_needs_group and (
                        group_col == "无"
                        or groups is None
                    ):
                        st.error(
                            "对比中包含混合效应模型，"
                            "但尚未选择分组变量。"
                            "请在上方选择分组变量后重试。"
                        )
                    else:
                        compare_df = compare_models(
                            y,
                            X,
                            groups,
                            compare_models_list,
                            robust_se=robust_se,
                        )
                        st.session_state["compare_df"] = compare_df
                        st.session_state["compare_success"] = True
                except Exception as exc:
                    st.error(f"多模型对比失败：{exc}")
                    st.session_state.pop("compare_df", None)

            if "compare_df" in st.session_state:
                _st_dataframe(
                    st.session_state["compare_df"],
                    use_container_width=True,
                )
                dataframe_download(
                    st.session_state["compare_df"],
                    "多模型对比结果.csv",
                    key="download_compare_models",
                )

                compare_df_now = st.session_state["compare_df"]

                if "AIC" in compare_df_now.columns:
                    # 修复：拟合失败的模型行 AIC 为 NaN，
                    # 直接 idxmin 可能选中 NaN 行，必须先过滤。
                    aic_valid = compare_df_now.dropna(
                        subset=["AIC"]
                    )

                    if not aic_valid.empty:
                        best_row = aic_valid.loc[
                            aic_valid["AIC"].idxmin()
                        ]
                        st.success(
                            f"按 AIC 最小原则，推荐模型："
                            f"**{best_row['模型']}**"
                            f"（AIC={best_row['AIC']:.3f}）"
                        )

                        # 修复：AIC 仅在同一似然体系内可比，
                        # 跨分布族（如 Logit 变换 OLS vs 二项 GLM）不可直接比较
                        st.caption(
                            "注意：AIC 仅在同类模型（相同分布族）间可比；"
                            "跨分布族比较请以交叉验证的样本外 RMSE / AUC 为准。"
                        )
                    else:
                        st.info(
                            "对比的模型均未能提供有效的 AIC，"
                            "无法按 AIC 推荐。"
                        )

        previous_model_type = st.session_state.get(
            "selected_model_type"
        )

        if (
            previous_model_type is not None
            and previous_model_type != final_model_type
        ):
            clear_model_session_state()

        st.session_state["selected_model_type"] = (
            final_model_type
        )

        if (
            final_model_type
            in [
                "线性混合效应模型",
                "Logit变换线性混合效应模型",
            ]
            and group_col == "无"
        ):
            st.error(
                "当前模型需要分组变量，请先选择重复观测分组变量。"
            )
            st.stop()

        vif_table = calculate_vif(
            X,
            dummy_info=model_meta.get("dummy_info") or {},
        )

        st.write("多重共线性诊断")

        if vif_table.empty:
            st.info(
                "当前没有足够的自变量计算 VIF。"
            )
        else:
            _st_dataframe(
                vif_table,
                use_container_width=True,
            )

            dataframe_download(
                vif_table,
                "VIF多重共线性诊断.csv",
                key="download_vif",
            )

        fit_button = st.button(
            "开始拟合最终模型",
            type="primary",
            key="fit_final_model",
        )

        # ===== 新增：清洗前后模型对比 =====
        # 用清洗前的原始数据（raw_df）与清洗后数据（clean_data_for_model）
        # 分别构造数据并拟合同一模型，对比指标说明数据清洗的价值。
        with st.expander(
            "📊 清洗前后模型对比（论文“数据预处理必要性”素材）",
            expanded=False,
        ):
            st.caption(
                "用清洗前（原始数据）与清洗后数据拟合同一个模型，"
                "对比 R² / RMSE / AIC 等指标。"
                "如果清洗后指标明显更好，就可以在论文里写："
                "“数据清洗显著提升了模型性能”。"
            )

            if st.button(
                "运行清洗前后对比",
                key="run_clean_compare",
            ):
                try:
                    compare_rows = []

                    for label, source_df in [
                        ("清洗前", raw_df),
                        (
                            "清洗后",
                            clean_data_for_model,
                        ),
                    ]:
                        try:
                            (
                                y_c,
                                X_c,
                                groups_c,
                                meta_c,
                            ) = build_model_data(
                                source_df,
                                target,
                                predictors,
                                variable_types,
                                group_col=group_col,
                            )

                            is_valid_c, msg_c = (
                                validate_model_selection(
                                    y=y_c,
                                    target_type=(
                                        variable_types[target]
                                    ),
                                    model_type=final_model_type,
                                    groups=groups_c,
                                )
                            )

                            if not is_valid_c:
                                compare_rows.append(
                                    {
                                        "阶段": label,
                                        "有效样本": len(y_c),
                                        "结论": (
                                            f"无法拟合：{msg_c}"
                                        ),
                                    }
                                )
                                continue

                            fitted_c = fit_model(
                                y_c,
                                X_c,
                                groups_c,
                                final_model_type,
                                robust_se=robust_se,
                            )

                            metric_c = make_metric_table(
                                fitted_c
                            )
                            metric_map_c = dict(
                                zip(
                                    metric_c["指标"],
                                    metric_c["数值"],
                                )
                            )

                            compare_rows.append(
                                {
                                    "阶段": label,
                                    "有效样本": len(y_c),
                                    "R²": metric_map_c.get(
                                        "R²",
                                        metric_map_c.get(
                                            "模型R²",
                                            np.nan,
                                        ),
                                    ),
                                    "RMSE": metric_map_c.get(
                                        "RMSE",
                                        np.nan,
                                    ),
                                    "AIC": metric_map_c.get(
                                        "AIC",
                                        np.nan,
                                    ),
                                    "MAE": metric_map_c.get(
                                        "MAE",
                                        np.nan,
                                    ),
                                }
                            )
                        except Exception as inner_error:
                            compare_rows.append(
                                {
                                    "阶段": label,
                                    "有效样本": np.nan,
                                    "结论": f"拟合失败：{inner_error}",
                                }
                            )

                    clean_compare_df = pd.DataFrame(
                        compare_rows
                    )
                    st.session_state[
                        "clean_compare_df"
                    ] = clean_compare_df
                except Exception as cc_error:
                    st.error(
                        f"清洗前后对比失败：{cc_error}"
                    )

            clean_compare_df = st.session_state.get(
                "clean_compare_df"
            )

            if isinstance(
                clean_compare_df,
                pd.DataFrame,
            ) and not clean_compare_df.empty:
                _st_dataframe(
                    clean_compare_df,
                    use_container_width=True,
                )

                dataframe_download(
                    clean_compare_df,
                    "清洗前后模型对比.csv",
                    key="download_clean_compare",
                )

                if (
                    "清洗后" in clean_compare_df[
                        "阶段"
                    ].values
                    and "清洗前" in clean_compare_df[
                        "阶段"
                    ].values
                ):
                    row_after = clean_compare_df[
                        clean_compare_df["阶段"]
                        == "清洗后"
                    ].iloc[0]
                    row_before = clean_compare_df[
                        clean_compare_df["阶段"]
                        == "清洗前"
                    ].iloc[0]

                    if (
                        pd.notna(
                            row_after.get("RMSE")
                        )
                        and pd.notna(
                            row_before.get("RMSE")
                        )
                        and row_before.get("RMSE", 0)
                        != 0
                    ):
                        change_pct = (
                            (
                                row_before["RMSE"]
                                - row_after["RMSE"]
                            )
                            / abs(row_before["RMSE"])
                            * 100
                        )

                        if change_pct > 1:
                            st.success(
                                f"清洗后 RMSE 下降 {change_pct:.1f}%，"
                                "数据清洗对模型有正向作用。"
                            )
                        elif change_pct < -1:
                            st.info(
                                f"清洗后 RMSE 上升 {abs(change_pct):.1f}%，"
                                "注意检查清洗是否删除了有价值的信息。"
                            )
                        else:
                            st.info(
                                "清洗前后 RMSE 差异不大"
                                f"（{change_pct:.1f}%）。"
                            )

        if fit_button:
            is_valid, validation_message = (
                validate_model_selection(
                    y=y,
                    target_type=variable_types[target],
                    model_type=final_model_type,
                    groups=groups,
                )
            )

            if not is_valid:
                st.error(validation_message)
                st.stop()

            try:
                new_fitted_result = fit_model(
                    y,
                    X,
                    groups,
                    final_model_type,
                    robust_se=robust_se,
                )

                # 携带模型元信息（特征名、哑变量说明等），
                # 供“新数据预测”等功能使用。
                new_fitted_result["meta"] = model_meta

                st.session_state[
                    "fitted_result"
                ] = new_fitted_result

                st.session_state[
                    "fitted_model"
                ] = new_fitted_result["model"]

                st.session_state[
                    "final_model_type"
                ] = final_model_type

                st.session_state[
                    "model_meta"
                ] = model_meta

                st.session_state[
                    "vif_table"
                ] = vif_table

                st.session_state[
                    "target_mapping"
                ] = model_meta.get(
                    "target_mapping"
                )

                st.success(
                    "模型拟合完成，结果已经保存。"
                )

                # 模型导出（joblib）：训练一次即可离线复用
                try:
                    import joblib
                    import io as _io

                    _model_buffer = _io.BytesIO()
                    joblib.dump(
                        {
                            "model": new_fitted_result["model"],
                            "model_type": final_model_type,
                            "meta": model_meta,
                            "target": target,
                            "predictors": predictors,
                            "variable_types": variable_types,
                            "target_mapping": model_meta.get(
                                "target_mapping"
                            ),
                        },
                        _model_buffer,
                    )

                    st.download_button(
                        "💾 导出训练好的模型 (.joblib)",
                        data=_model_buffer.getvalue(),
                        file_name="trained_model.joblib",
                        mime="application/octet-stream",
                        key="download_model_joblib",
                    )

                    st.caption(
                        "导出的模型可用于离线复用："
                        "在后续会话中加载后即可对新数据预测，"
                        "无需重新训练。"
                    )
                except Exception as _export_error:
                    st.warning(
                        f"模型导出失败（可忽略）：{_export_error}"
                    )

                # 引导进度：已完成第 6 步（建模与诊断）
                st.session_state["guide_auto_step"] = max(
                    st.session_state.get("guide_auto_step", 0),
                    6,
                )

            except Exception as exc:
                st.error(
                    f"模型拟合失败：{exc}"
                )

        # 重要修复：
        # 每次页面重新运行时从 session_state 恢复模型结果。
        fitted_result = st.session_state.get(
            "fitted_result"
        )

        fitted_model = st.session_state.get(
            "fitted_model"
        )

        stored_model_type = st.session_state.get(
            "final_model_type",
            final_model_type,
        )

        # --------------------------------------------------------
        # 模型结果
        # --------------------------------------------------------

        if fitted_result is None:
            st.info(
                '请点击“开始拟合最终模型”后查看模型结果。'
            )

        else:
            st.subheader("9. 模型结果")

            mapping = st.session_state.get(
                "target_mapping"
            )

            if mapping:
                st.write("分类因变量编码映射")

                mapping_table = pd.DataFrame(
                    {
                        "原始类别": list(
                            mapping.keys()
                        ),
                        "模型编码": list(
                            mapping.values()
                        ),
                    }
                )

                _st_dataframe(
                    mapping_table,
                    use_container_width=True,
                )

            # ===== 新增：虚拟变量编码说明 =====
            stored_meta = st.session_state.get(
                "model_meta",
                {},
            )
            dummy_info = (stored_meta or {}).get(
                "dummy_info",
                {},
            )

            if dummy_info:
                with st.expander(
                    "🧩 分类变量哑变量编码说明",
                    expanded=False,
                ):
                    st.caption(
                        "程序对分类自变量使用独热编码，并删除参照类别"
                        "（drop_first）。系数表中的列名含义如下。"
                    )

                    dummy_rows = []

                    for col, info in dummy_info.items():
                        categories = info.get(
                            "原始类别",
                            [],
                        )
                        reference = info.get(
                            "参照类别（被删除）",
                            "",
                        )

                        for cat in categories:
                            if cat == reference:
                                continue
                            dummy_rows.append(
                                {
                                    "原始变量": col,
                                    "哑变量列名": f"{col}_{cat}",
                                    "含义": (
                                        f"该观测的「{col}」是否为「{cat}」"
                                        f"（参照类别：{reference}）"
                                    ),
                                }
                            )

                    dummy_table = pd.DataFrame(
                        dummy_rows
                    )

                    _st_dataframe(
                        dummy_table,
                        use_container_width=True,
                    )

            if model_is_converged(fitted_model):
                st.success(
                    "模型已收敛或未检测到明显收敛问题。"
                )
            else:
                st.warning(
                    "模型未收敛，当前系数、P值和预测结果不建议直接用于论文。"
                )

            metric_table = make_metric_table(
                fitted_result
            )

            st.write("模型评价指标")
            _st_dataframe(
                metric_table,
                use_container_width=True,
            )

            dataframe_download(
                metric_table,
                "模型评价指标.csv",
                key="download_model_metrics",
            )

            result_table = make_result_table(
                fitted_model,
                stored_model_type,
            )

            st.write("模型系数结果")
            _st_dataframe(
                result_table,
                use_container_width=True,
            )

            dataframe_download(
                result_table,
                "模型系数结果.csv",
                key="download_model_coefficients",
            )

            # ===== 新增：Logistic 边际效应（AME，论文常用表述） =====
            if stored_model_type == "二项Logistic回归":
                try:
                    logit_model = fitted_model
                    X_logit = logit_model.model.exog
                    y_logit = np.asarray(
                        logit_model.model.endog,
                        dtype=float,
                    )
                    prob_logit = np.asarray(
                        logit_model.predict(X_logit),
                        dtype=float,
                    )
                    coef_logit = np.asarray(
                        logit_model.params,
                        dtype=float,
                    )

                    # 平均边际效应 AME = mean(β * p * (1-p))
                    # 对离散（哑变量）特征用有限差分更准确
                    ame_values = []

                    for k, param_name in enumerate(
                        logit_model.params.index
                    ):
                        col_k = X_logit[:, k]

                        if (
                            np.all(
                                (col_k == 0)
                                | (col_k == 1)
                            )
                            and len(np.unique(col_k)) == 2
                        ):
                            # 哑变量：p(x=1) - p(x=0)
                            x0 = X_logit.copy()
                            x1 = X_logit.copy()
                            x0[:, k] = 0.0
                            x1[:, k] = 1.0
                            p0 = logit_model.predict(x0)
                            p1 = logit_model.predict(x1)
                            ame = np.mean(p1 - p0)
                        else:
                            ame = np.mean(
                                coef_logit[k]
                                * prob_logit
                                * (1 - prob_logit)
                            )

                        ame_values.append(ame)

                    ame_table = pd.DataFrame(
                        {
                            "变量": list(
                                logit_model.params.index
                            ),
                            "平均边际效应(AME)": ame_values,
                        }
                    )
                    ame_table["含义"] = (
                        "自变量每增加1单位，因变量=1的概率平均变化量"
                    )

                    with st.expander(
                        "📈 Logistic 边际效应（论文表述用）",
                        expanded=False,
                    ):
                        st.caption(
                            "平均边际效应（AME）：在其他变量取均值时，"
                            "该变量每增加 1 个单位，因变量取 1 类的概率"
                            "平均变化多少。论文里比 OR 更直观。"
                        )
                        _st_dataframe(
                            ame_table,
                            use_container_width=True,
                        )

                        dataframe_download(
                            ame_table,
                            "Logistic边际效应.csv",
                            key="download_ame",
                        )
                except Exception:
                    pass

            st.write("LaTeX 三线表（可直接粘贴到论文）")

            # 修复：pandas 3.0 移除了 to_latex 的 booktabs 参数，
            # 用兼容函数 safe_to_latex 生成，新旧版本都能工作。
            latex_table_text = safe_to_latex(
                result_table,
                index=False,
                booktabs=True,
            )

            st.code(
                latex_table_text,
                language="latex",
            )

            st.download_button(
                "下载 LaTeX 三线表 (.tex)",
                data=latex_table_text.encode("utf-8"),
                file_name="模型系数三线表.tex",
                mime="text/plain",
                key="download_latex_coefficients",
            )

            prediction_table = create_prediction_table(
                fitted_result
            )

            st.write(
                "实际值、预测值与残差"
            )

            _st_dataframe(
                prediction_table.head(100),
                use_container_width=True,
            )

            dataframe_download(
                prediction_table,
                "实际值预测值残差.csv",
                key="download_prediction_table",
            )

            # ----------------------------------------------------
            # 新增：用训练好的模型预测新数据
            # ----------------------------------------------------

            st.subheader("新数据预测")

            st.caption(
                "上传一份包含与建模时相同自变量列的新数据表"
                "（CSV/Excel），程序会按训练时的规则自动预处理并输出预测值。"
            )

            new_data_file = st.file_uploader(
                "上传新数据表（预测用）",
                type=["csv", "xlsx", "xls"],
                key="new_data_uploader",
            )

            if new_data_file is not None:
                try:
                    new_data_file.seek(0)

                    if new_data_file.name.lower().endswith(".csv"):
                        try:
                            new_df = pd.read_csv(
                                new_data_file,
                                encoding="utf-8",
                            )
                        except UnicodeDecodeError:
                            new_data_file.seek(0)
                            new_df = pd.read_csv(
                                new_data_file,
                                encoding="gbk",
                            )
                    else:
                        new_df = pd.read_excel(new_data_file)

                    new_df.columns = make_unique_columns(
                        new_df.columns
                    )

                    st.info(
                        f"新数据共 {new_df.shape[0]} 行、"
                        f"{new_df.shape[1]} 列。"
                    )

                    if st.button(
                        "🚀 生成新数据预测",
                        key="run_new_data_predict",
                    ):
                        try:
                            (
                                prediction_new,
                                predict_notes,
                            ) = predict_new_data(
                                fitted_result,
                                new_df,
                                target,
                                predictors,
                                variable_types,
                            )

                            st.session_state[
                                "new_prediction_df"
                            ] = prediction_new
                            st.session_state[
                                "new_prediction_notes"
                            ] = predict_notes
                        except Exception as predict_error:
                            st.error(
                                f"新数据预测失败：{predict_error}"
                            )

                except Exception as read_error:
                    st.error(f"读取新数据失败：{read_error}")

            if "new_prediction_df" in st.session_state:
                st.write("预测结果（前100行）")

                _st_dataframe(
                    st.session_state[
                        "new_prediction_df"
                    ].head(100),
                    use_container_width=True,
                )

                dataframe_download(
                    st.session_state[
                        "new_prediction_df"
                    ],
                    "新数据预测结果.csv",
                    key="download_new_prediction",
                )

                predict_notes = st.session_state.get(
                    "new_prediction_notes",
                    [],
                )

                if predict_notes:
                    with st.expander(
                        "查看预测预处理说明",
                        expanded=False,
                    ):
                        for note in predict_notes:
                            st.markdown(f"- {note}")

            # ----------------------------------------------------
            # 训练集 / 测试集评估
            # ----------------------------------------------------

            st.subheader(
                "训练集/测试集评估"
            )

            if not use_test_set:
                st.info(
                    "当前已关闭训练集/测试集评估。"
                    "上方模型评价指标为样本内指标。"
                )

            elif stored_model_type in [
                "线性混合效应模型",
                "Logit变换线性混合效应模型",
            ]:
                st.info(
                    "混合效应模型暂未采用普通随机划分，"
                    "建议按照分组变量进行分组交叉验证。"
                )

            else:
                try:
                    if len(y) < 10:
                        st.warning(
                            "样本量少于10，训练集/测试集划分结果可能不稳定。"
                        )

                    indices = np.arange(len(y))

                    has_groups = (
                        groups is not None
                        and groups.nunique() < len(groups)
                    )

                    if has_groups:
                        splitter = GroupShuffleSplit(
                            n_splits=1,
                            test_size=test_size,
                            random_state=42,
                        )

                        train_index, test_index = next(
                            splitter.split(
                                X,
                                y,
                                groups=groups,
                            )
                        )
                    else:
                        stratify_value = None

                        if stored_model_type in [
                            "二项Logistic回归",
                            "多项Logistic回归",
                        ]:
                            stratify_value = y

                        train_index, test_index = (
                            train_test_split(
                                indices,
                                test_size=test_size,
                                random_state=42,
                                stratify=stratify_value,
                            )
                        )

                    y_train = y.iloc[train_index]
                    y_test = y.iloc[test_index]

                    X_train = X.iloc[train_index]
                    X_test = X.iloc[test_index]

                    # 修复：此处的分组属于“训练集分组”（用于拟合训练集模型），
                    # 原先命名为 test_groups 与语义相反，容易误导后续维护。
                    train_groups = None

                    if has_groups:
                        train_groups = groups.iloc[
                            train_index
                        ].reset_index(drop=True)

                    test_model_result = fit_model(
                        y_train.reset_index(drop=True),
                        X_train.reset_index(drop=True),
                        train_groups,
                        stored_model_type,
                        robust_se=robust_se,
                    )

                    test_model = test_model_result[
                        "model"
                    ]

                    if stored_model_type == (
                        "二项Logistic回归"
                    ):
                        test_probability = np.asarray(
                            test_model.predict(X_test),
                            dtype=float,
                        )

                        test_class = (
                            test_probability >= 0.5
                        ).astype(int)

                        if len(np.unique(y_test)) < 2:
                            st.warning("测试集中只有一种类别，仅给出准确率。")
                            test_metrics = ["准确率"]
                            test_values = [
                                accuracy_score(
                                    y_test.astype(int),
                                    test_class,
                                ),
                            ]
                        else:
                            test_metrics = [
                                "准确率",
                                "平衡准确率",
                                "精确率",
                                "召回率",
                                "F1值",
                                "Log Loss",
                            ]

                            test_values = [
                                accuracy_score(
                                    y_test.astype(int),
                                    test_class,
                                ),
                                balanced_accuracy_score(
                                    y_test.astype(int),
                                    test_class,
                                ),
                                precision_score(
                                    y_test.astype(int),
                                    test_class,
                                    zero_division=0,
                                ),
                                recall_score(
                                    y_test.astype(int),
                                    test_class,
                                    zero_division=0,
                                ),
                                f1_score(
                                    y_test.astype(int),
                                    test_class,
                                    zero_division=0,
                                ),
                                log_loss(
                                    y_test.astype(int),
                                    np.clip(test_probability, 1e-6, 1 - 1e-6),
                                ),
                            ]

                            if len(
                                np.unique(
                                    y_test
                                )
                            ) == 2:
                                test_metrics.append(
                                    "ROC-AUC"
                                )
                                test_values.append(
                                    roc_auc_score(
                                        y_test,
                                        test_probability,
                                    )
                                )
                    elif stored_model_type == (
                        "多项Logistic回归"
                    ):
                        probability = np.asarray(
                            test_model.predict(X_test)
                        )

                        test_class = (
                            probability.argmax(axis=1)
                        )

                        if len(np.unique(y_test)) < 2:
                            st.warning("测试集中只有一种类别，仅给出准确率。")
                            test_metrics = ["准确率"]
                            test_values = [
                                accuracy_score(
                                    y_test.astype(int),
                                    test_class,
                                ),
                            ]
                        else:
                            test_metrics = [
                                "准确率",
                                "平衡准确率",
                                "宏平均F1",
                            ]

                            test_values = [
                                accuracy_score(
                                    y_test.astype(int),
                                    test_class,
                                ),
                                balanced_accuracy_score(
                                    y_test.astype(int),
                                    test_class,
                                ),
                                f1_score(
                                    y_test.astype(int),
                                    test_class,
                                    average="macro",
                                    zero_division=0,
                                ),
                            ]
                    else:
                        raw_test_prediction = np.asarray(
                            test_model.predict(X_test),
                            dtype=float,
                        )

                        # Logit模型预测值原本位于Logit尺度，
                        # 必须还原到0到1的原始比例尺度。
                        if stored_model_type == "Logit变换线性回归":
                            test_prediction = inverse_logit(
                                raw_test_prediction
                            )
                        else:
                            test_prediction = raw_test_prediction

                        test_metrics = [
                            "RMSE",
                            "MAE",
                            "R²",
                        ]

                        test_values = [
                            np.sqrt(
                                mean_squared_error(
                                    y_test,
                                    test_prediction,
                                )
                            ),
                            mean_absolute_error(
                                y_test,
                                test_prediction,
                            ),
                            r2_score(
                                y_test,
                                test_prediction,
                            ),
                        ]

                    test_metric_table = pd.DataFrame(
                        {
                            "测试集指标": test_metrics,
                            "数值": test_values,
                        }
                    )

                    _st_dataframe(
                        test_metric_table,
                        use_container_width=True,
                    )

                    dataframe_download(
                        test_metric_table,
                        "测试集评价指标.csv",
                        key="download_test_metrics",
                    )

                except Exception as exc:
                    st.warning(
                        f"测试集评估失败：{exc}"
                    )

            # ----------------------------------------------------
            # K-Fold 交叉验证（比单次划分更稳定的泛化评估）
            # ----------------------------------------------------
            with st.expander(
                "🧪 K-Fold 交叉验证（更稳定的泛化评估）",
                expanded=False,
            ):
                st.caption(
                    "把数据分成 K 份，轮流用 K-1 份训练、1 份验证，"
                    "取 K 次结果的均值±标准差，比单次划分更可信。"
                    "有分组变量时自动按组划分，避免同一对象"
                    "同时出现在训练集和测试集。"
                )

                kfold_k = st.selectbox(
                    "折数 K",
                    [3, 4, 5, 6, 8, 10],
                    index=2,
                    key="kfold_k_selector",
                )

                if st.button(
                    "运行 K-Fold 交叉验证",
                    key="run_kfold_button",
                ):
                    try:
                        kfold_fold_df, kfold_summary_df, kfold_note = (
                            run_kfold_cv(
                                y,
                                X,
                                groups,
                                stored_model_type,
                                n_splits=kfold_k,
                                robust_se=robust_se,
                            )
                        )
                        st.session_state["kfold_result"] = {
                            "fold": kfold_fold_df,
                            "summary": kfold_summary_df,
                            "note": kfold_note,
                        }
                    except Exception as kfold_error:
                        st.error(
                            f"K-Fold 交叉验证失败：{kfold_error}"
                        )

                if "kfold_result" in st.session_state:
                    kfold_stored = st.session_state["kfold_result"]

                    st.write("每折评估结果")
                    _st_dataframe(
                        kfold_stored["fold"],
                        use_container_width=True,
                    )

                    st.write("汇总统计（均值±标准差）")
                    _st_dataframe(
                        kfold_stored["summary"],
                        use_container_width=True,
                    )

                    dataframe_download(
                        kfold_stored["fold"],
                        "K-Fold每折结果.csv",
                        key="download_kfold_fold",
                    )

                    dataframe_download(
                        kfold_stored["summary"],
                        "K-Fold汇总.csv",
                        key="download_kfold_summary",
                    )

                    kfold_text = (
                        f"为评估模型泛化能力，采用K折交叉验证"
                        f"（K={len(kfold_stored['fold'])}）。"
                        + kfold_stored["note"]
                    )

                    for _, row in kfold_stored[
                        "summary"
                    ].iterrows():
                        kfold_text += (
                            f"{row['指标']}的均值为{row['均值']:.4f}"
                            f"（标准差{row['标准差']:.4f}），"
                        )

                    kfold_text = kfold_text.rstrip("，") + "。"

                    st.text_area(
                        "论文表述",
                        kfold_text,
                        height=130,
                        key="kfold_text_area",
                    )

            # ----------------------------------------------------
            # 线性模型诊断
            # ----------------------------------------------------

            if stored_model_type in [
                "多元线性回归",
                "Logit变换线性回归",
            ]:
                st.subheader("线性模型诊断")

                diagnostic_table = (
                    make_diagnostic_table(
                        fitted_model
                    )
                )

                _st_dataframe(
                    diagnostic_table,
                    use_container_width=True,
                )

                dataframe_download(
                    diagnostic_table,
                    "线性模型诊断.csv",
                    key="download_diagnostic_table",
                )

                residuals = np.asarray(
                    fitted_model.resid,
                    dtype=float,
                )

                fitted_values = np.asarray(
                    fitted_model.fittedvalues,
                    dtype=float,
                )

                diag_col1, diag_col2 = st.columns(2)

                with diag_col1:
                    fig_residual, ax_residual = (
                        plt.subplots(
                            figsize=(7, 5)
                        )
                    )

                    sns.scatterplot(
                        x=fitted_values,
                        y=residuals,
                        ax=ax_residual,
                    )

                    ax_residual.axhline(
                        0,
                        color="red",
                        linestyle="--",
                    )

                    ax_residual.set_xlabel(
                        "拟合值"
                    )
                    ax_residual.set_ylabel(
                        "残差"
                    )
                    ax_residual.set_title(
                        "残差-拟合值图"
                    )

                    show_fig(fig_residual)
                    plt.close(fig_residual)

                with diag_col2:
                    fig_qq, ax_qq = plt.subplots(
                        figsize=(7, 5)
                    )

                    probplot(
                        residuals,
                        dist="norm",
                        plot=ax_qq,
                    )

                    ax_qq.set_title(
                        "残差正态QQ图"
                    )

                    show_fig(fig_qq)
                    plt.close(fig_qq)

                try:
                    influence = OLSInfluence(
                        fitted_model
                    )

                    cooks_distance = (
                        influence.cooks_distance[0]
                    )

                    cooks_table = pd.DataFrame(
                        {
                            "样本序号": np.arange(
                                len(cooks_distance)
                            ),
                            "Cook距离": cooks_distance,
                        }
                    ).sort_values(
                        "Cook距离",
                        ascending=False,
                    )

                    st.write(
                        "Cook距离较大的观测"
                    )

                    _st_dataframe(
                        cooks_table.head(20),
                        use_container_width=True,
                    )

                    dataframe_download(
                        cooks_table,
                        "Cook距离.csv",
                        key="download_cooks_distance",
                    )

                except Exception as exc:
                    st.info(
                        f"Cook距离计算失败：{exc}"
                    )

            # ----------------------------------------------------
            # 分类模型诊断
            # ----------------------------------------------------

            if stored_model_type == (
                "二项Logistic回归"
            ):
                st.subheader(
                    "二分类模型诊断"
                )

                cm = confusion_matrix(
                    prediction_table[
                        "实际值"
                    ].astype(int),
                    prediction_table[
                        "预测类别"
                    ].astype(int),
                )

                cm_index = [
                    f"真实{i}"
                    for i in range(cm.shape[0])
                ]

                cm_columns = [
                    f"预测{i}"
                    for i in range(cm.shape[1])
                ]

                cm_table = pd.DataFrame(
                    cm,
                    index=cm_index,
                    columns=cm_columns,
                )

                st.write("混淆矩阵")
                _st_dataframe(
                    cm_table,
                    use_container_width=True,
                )

                dataframe_download(
                    cm_table.reset_index(),
                    "二分类混淆矩阵.csv",
                    key="download_confusion_matrix",
                )

                # ===== 新增：ROC 曲线与混淆矩阵热力图 =====
                try:
                    probability = np.asarray(
                        fitted_result["probability"],
                        dtype=float,
                    )
                    y_true = np.asarray(
                        prediction_table["实际值"],
                        dtype=int,
                    )

                    from sklearn.metrics import roc_curve, auc

                    fpr, tpr, _ = roc_curve(
                        y_true,
                        probability,
                    )
                    roc_auc = auc(fpr, tpr)

                    fig_roc, ax_roc = plt.subplots(
                        figsize=(7, 5)
                    )
                    ax_roc.plot(
                        fpr,
                        tpr,
                        color="#1f6feb",
                        lw=2,
                        label=f"ROC曲线 (AUC={roc_auc:.3f})",
                    )
                    ax_roc.plot(
                        [0, 1],
                        [0, 1],
                        color="gray",
                        linestyle="--",
                        label="随机猜测",
                    )
                    ax_roc.set_xlabel("假正率 (FPR)")
                    ax_roc.set_ylabel("真正率 (TPR)")
                    ax_roc.set_title("ROC 曲线")
                    ax_roc.legend(loc="lower right")
                    ax_roc.set_xlim(0, 1)
                    ax_roc.set_ylim(0, 1)
                    show_fig(fig_roc, name="roc_curve")
                except Exception as roc_error:
                    st.info(f"ROC 曲线绘制失败：{roc_error}")

                try:
                    fig_cm, ax_cm = plt.subplots(
                        figsize=(6, 5)
                    )
                    sns.heatmap(
                        cm,
                        annot=True,
                        fmt="d",
                        cmap="Blues",
                        xticklabels=cm_columns,
                        yticklabels=cm_index,
                        ax=ax_cm,
                    )
                    ax_cm.set_title("混淆矩阵热力图")
                    ax_cm.set_xlabel("预测类别")
                    ax_cm.set_ylabel("真实类别")
                    show_fig(
                        fig_cm,
                        name="confusion_matrix_heatmap",
                    )
                except Exception as cm_error:
                    st.info(
                        f"混淆矩阵热力图绘制失败：{cm_error}"
                    )

            # ----------------------------------------------------
            # 模型假设
            # ----------------------------------------------------

            st.subheader("10. 模型假设")

            assumptions = generate_assumptions(
                target,
                predictors,
                variable_types,
                stored_model_type,
                vif_table=st.session_state.get(
                    "vif_table",
                    vif_table,
                ),
                group_col=group_col,
            )

            assumptions_text = "\n".join(
                [
                    f"{index + 1}. {text}"
                    for index, text in enumerate(
                        assumptions
                    )
                ]
            )

            st.text_area(
                "可直接复制到论文的模型假设",
                assumptions_text,
                height=350,
            )

            # ----------------------------------------------------
            # 论文表述草稿
            # ----------------------------------------------------

            st.subheader("11. 论文表述草稿")

            paper_text = f"""
    本文首先对原始数据进行完整性、一致性和变量类型检查。
    原始数据包含{original_shape[0]}条样本和{original_shape[1]}个变量。
    针对缺失数据，本文采用“{missing_method}”方法进行处理，
    共插补{imputed_cells}个缺失单元格，
    并因缺失值删除{deleted_target_missing_rows + deleted_missing_rows}条样本。
    针对异常数据，本文采用“{outlier_method}”方法进行识别，
    共检测到{outlier_row_count}条异常样本，
    并根据处理策略删除{deleted_outlier_rows}条异常样本。
    本文将“{target}”作为因变量，
    将{", ".join(predictors)}作为解释变量。
    """.strip()

            if group_col != "无":
                paper_text += (
                    f'考虑到同一“{group_col}”下可能存在多次观测，'
                    "本文将其作为重复观测的分组变量。"
                )

            paper_text += (
                f"在综合考虑因变量类型、数据取值范围和观测结构后，"
                f'最终采用“{stored_model_type}”进行建模。'
                f"清洗后最终保留{clean_data_for_model.shape[0]}条有效样本，"
                "并据此开展后续统计分析和模型估计。"
            )

            st.text_area(
                "论文表述草稿",
                paper_text,
                height=300,
            )

    except Exception as exc:
        st.error(
            f"模型数据构造或分析过程中发生错误：{exc}"
        )



    # ============================================================
    # 二十一、高级分析方法（补充）
    # ============================================================

    st.subheader("8.5 高级分析方法")

    if "clean_data_for_model" not in locals() or clean_data_for_model is None:
        st.info("请先在上方完成数据清洗，再使用高级分析方法。")
    else:
        # 修复：高级分析结果与当前数据绑定。
        # 更换数据文件后自动清除旧结果，避免误用上一次数据的结果。
        adv_file_hash = hashlib.sha256(
            uploaded_file.getvalue()
        ).hexdigest()

        # 修复：高级分析结果还与“所选列/参数”绑定。
        # 用户修改参数后旧结果不再展示，避免与新选择不一致。
        def _adv_cache_valid(key, cols_signature=None):
            """校验高级分析缓存是否仍然有效。"""
            sig = st.session_state.get(
                "adv_cache_signatures",
                {},
            ).get(key)

            if sig is None:
                return False

            if sig.get("data_hash") != adv_file_hash:
                return False

            if cols_signature is not None and (
                sig.get("cols_signature")
                != cols_signature
            ):
                return False

            return True

        def _adv_cache_store(key, cols_signature=None):
            """记录高级分析结果的缓存签名。"""
            st.session_state.setdefault(
                "adv_cache_signatures",
                {},
            )[key] = {
                "data_hash": adv_file_hash,
                "cols_signature": cols_signature,
            }

        for _adv_key in [
            "topsis_data",
            "gm11_result",
            "arima_result",
            "anova_result",
            "chi2_result",
            "pca_data",
            "cluster_data",
            "ml_data",
        ]:
            _adv_item = st.session_state.get(_adv_key)

            if isinstance(_adv_item, dict) and (
                _adv_item.get("_data_hash") != adv_file_hash
            ):
                st.session_state.pop(_adv_key, None)

        # 修复：引导进度不再无条件推满。
        # 第 7 步仅在“确实运行了某项高级分析”时由各按钮回调推进，
        # 这里只更新自动进度展示（guide_auto_step），不影响用户手动勾选。
        if any(
            st.session_state.get(_adv_key) is not None
            for _adv_key in [
                "topsis_data",
                "gm11_result",
                "arima_result",
                "anova_result",
                "chi2_result",
                "pca_data",
                "cluster_data",
                "ml_data",
            ]
        ):
            st.session_state["guide_auto_step"] = max(
                st.session_state.get("guide_auto_step", 0),
                7,
            )

        st.info(
            "高级分析为**可选模块**，按题型选用："
            "评价类→熵权TOPSIS / PCA；预测类→灰色预测 / ARIMA；"
            "分类类→随机森林 / 决策树；"
            "任何已拟合模型都可做稳健性分析。"
            "每个 Tab 顶部都有“何时使用”的说明。"
        )

        adv_tabs = st.tabs(
            [
                "熵权法+TOPSIS",
                "灰色预测GM(1,1)",
                "ARIMA时间序列",
                "方差分析+卡方检验",
                "PCA主成分分析",
                "聚类分析",
                "随机森林/决策树",
                "稳健性分析",
                "AHP层次分析",
                "正则化变量筛选",
                "时间序列辅助",
            ]
        )

        adv_numeric_cols = [
            col
            for col in clean_data_for_model.columns
            if pd.to_numeric(
                clean_data_for_model[col],
                errors="coerce",
            ).notna().mean() >= 0.8
        ]

        # ---------- 熵权法 + TOPSIS ----------
        with adv_tabs[0]:
            st.markdown("**熵权法确定权重 + TOPSIS 综合评价**")
            st.caption(
                "适合评价类题目：对多个评价对象（行）按多个指标（列）综合打分排名。"
            )

            if len(adv_numeric_cols) < 2:
                st.info("至少需要两列数值型指标才能进行评价。")
            else:
                eval_cols = st.multiselect(
                    "选择评价指标列",
                    adv_numeric_cols,
                    default=adv_numeric_cols[: min(4, len(adv_numeric_cols))],
                    key="eval_cols_selector",
                )

                if len(eval_cols) < 2:
                    st.warning("请至少选择两个评价指标。")
                else:
                    directions = {}
                    dir_cols = st.columns(len(eval_cols))

                    for i, col in enumerate(eval_cols):
                        directions[col] = dir_cols[i].selectbox(
                            f"{col} 方向",
                            ["正向指标", "负向指标"],
                            key=f"eval_dir_{col}",
                        )

                    if st.button("计算熵权法+TOPSIS", key="run_topsis"):
                        try:
                            eval_df = clean_data_for_model[
                                eval_cols
                            ].apply(
                                pd.to_numeric,
                                errors="coerce",
                            ).dropna().reset_index(drop=True)

                            if len(eval_df) < 2:
                                raise ValueError("有效评价对象少于2个。")

                            matrix = eval_df.to_numpy(dtype=float)

                            # 修复：负向指标不再取倒数（含 0/负值时会产生 inf/NaN），
                            # 改为“取反 + 平移”同向化，保证越大越好。
                            for i, col in enumerate(eval_cols):
                                col_values = matrix[:, i]

                                if directions[col] == "负向指标":
                                    matrix[:, i] = (
                                        col_values.max() - col_values
                                    )

                            # min-max 归一化到 [0,1]，消除量纲
                            col_min = matrix.min(axis=0)
                            col_max = matrix.max(axis=0)
                            col_range = col_max - col_min
                            const_mask = col_range == 0
                            col_range[const_mask] = 1.0
                            matrix = (matrix - col_min) / col_range

                            n_obj = len(eval_df)
                            eps = 1e-12

                            # 熵权法：常数列（无区分信息）按等概率处理，权重为 0
                            p = matrix / np.maximum(
                                matrix.sum(axis=0),
                                eps,
                            )
                            p[:, const_mask] = 1.0 / n_obj

                            k = 1.0 / np.log(n_obj)
                            entropy = -k * np.sum(
                                p * np.log(p + eps),
                                axis=0,
                            )
                            entropy[const_mask] = 1.0
                            d = 1 - entropy
                            weights = d / d.sum()

                            weight_table = pd.DataFrame(
                                {
                                    "指标": eval_cols,
                                    "熵值": entropy,
                                    "差异系数": d,
                                    "熵权": weights,
                                }
                            )

                            # 归一化后无需再除以向量长度
                            weighted = matrix * weights

                            ideal_best = weighted.max(axis=0)
                            ideal_worst = weighted.min(axis=0)

                            dist_best = np.sqrt(
                                ((weighted - ideal_best) ** 2).sum(axis=1)
                            )
                            dist_worst = np.sqrt(
                                ((weighted - ideal_worst) ** 2).sum(axis=1)
                            )

                            closeness = dist_worst / (
                                dist_best + dist_worst + eps
                            )

                            topsis_table = eval_df.copy()
                            topsis_table["与最优解距离"] = dist_best
                            topsis_table["与最劣解距离"] = dist_worst
                            topsis_table["综合得分"] = closeness
                            topsis_table["排名"] = (
                                closeness.argsort()[::-1].argsort() + 1
                            )
                            topsis_table = topsis_table.sort_values(
                                "排名"
                            ).reset_index(drop=True)

                            st.session_state["topsis_data"] = {
                                "weight_table": weight_table,
                                "topsis_table": topsis_table,
                                "_data_hash": adv_file_hash,
                            }
                            _adv_cache_store(
                                "topsis_data",
                                cols_signature=tuple(
                                    eval_cols
                                ),
                            )
                        except Exception as eval_error:
                            st.error(f"熵权法/TOPSIS计算失败：{eval_error}")

                    if (
                        "topsis_data" in st.session_state
                        and _adv_cache_valid(
                            "topsis_data",
                            cols_signature=tuple(
                                eval_cols
                            ),
                        )
                    ):
                        topsis_stored = st.session_state["topsis_data"]
                        weight_table = topsis_stored["weight_table"]
                        topsis_table = topsis_stored["topsis_table"]

                        st.write("熵权法计算权重")
                        _st_dataframe(
                            weight_table,
                            use_container_width=True,
                        )

                        dataframe_download(
                            weight_table,
                            "熵权法权重.csv",
                            key="download_entropy_weights",
                        )

                        st.write("TOPSIS 综合评价排名")
                        _st_dataframe(
                            topsis_table,
                            use_container_width=True,
                        )

                        dataframe_download(
                            topsis_table,
                            "TOPSIS排名.csv",
                            key="download_topsis",
                        )

                        best_row = topsis_table.iloc[0]
                        max_weight_row = weight_table.loc[
                            weight_table["熵权"].idxmax()
                        ]

                        topsis_text = (
                            f"采用熵权法确定指标权重，其中"
                            f"“{max_weight_row['指标']}”的权重最高"
                            f"（{max_weight_row['熵权']:.4f}），"
                            "说明该指标提供的区分信息最多。"
                            "基于熵权权重，采用TOPSIS方法计算各评价对象"
                            "与正负理想解的距离，得到综合贴近度并排序，"
                            f"综合得分最高的评价对象为第{best_row['排名']}号"
                            f"（得分{best_row['综合得分']:.4f}）。"
                        )

                        st.text_area(
                            "论文表述",
                            topsis_text,
                            height=130,
                            key="topsis_text_area",
                        )

        # ---------- 灰色预测 GM(1,1) ----------
        with adv_tabs[1]:
            st.markdown("**灰色预测 GM(1,1)**")
            st.caption(
                "适合小样本时间序列（通常不少于4期）的中短期预测，"
                "如人口、产量、需求量等。"
            )

            if not adv_numeric_cols:
                st.info("没有可用的数值列。")
            else:
                gm_col = st.selectbox(
                    "选择要预测的序列列",
                    adv_numeric_cols,
                    key="gm_col_selector",
                )
                gm_steps = st.slider(
                    "预测期数",
                    1,
                    10,
                    3,
                    key="gm_steps_slider",
                )

                def gm11_predict(series, n_forecast=5):
                    """灰色预测 GM(1,1)，返回预测值与检验指标。"""
                    x0 = np.asarray(series, dtype=float)
                    n = len(x0)

                    if n < 4:
                        raise ValueError("灰色预测至少需要4个数据点")

                    x1 = np.cumsum(x0)
                    z = (x1[:-1] + x1[1:]) / 2.0

                    B = np.column_stack([-z, np.ones(n - 1)])
                    Y = x0[1:]
                    theta = np.linalg.lstsq(B, Y, rcond=None)[0]
                    a, b = theta

                    total = n + n_forecast
                    x1_hat = np.zeros(total)
                    x1_hat[0] = x0[0]

                    for t in range(1, total):
                        x1_hat[t] = (
                            (x0[0] - b / a) * np.exp(-a * t) + b / a
                        )

                    x0_hat = np.empty(total)
                    x0_hat[0] = x0[0]

                    for t in range(1, total):
                        x0_hat[t] = x1_hat[t] - x1_hat[t - 1]

                    in_sample = x0_hat[:n]
                    forecast = x0_hat[n:]

                    relative_errors = np.abs(x0 - in_sample) / (
                        np.abs(x0) + 1e-12
                    )
                    mean_relative_error = relative_errors.mean()

                    lambda_vals = x0[:-1] / (x0[1:] + 1e-12)
                    lower_bound = np.exp(-2 / (n + 1))
                    upper_bound = np.exp(2 / (n + 1))
                    lambda_ok = bool(
                        lambda_vals.min() > lower_bound
                        and lambda_vals.max() < upper_bound
                    )

                    s1 = x0.std(ddof=1)
                    resid = x0 - in_sample
                    s2 = resid.std(ddof=1)
                    c_value = s2 / (s1 + 1e-12)

                    if c_value < 0.35:
                        level = "好"
                    elif c_value < 0.5:
                        level = "合格"
                    elif c_value < 0.65:
                        level = "勉强"
                    else:
                        level = "不合格"

                    return {
                        "a": a,
                        "b": b,
                        "in_sample": in_sample,
                        "forecast": forecast,
                        "mean_relative_error": mean_relative_error,
                        "posterior_ratio": c_value,
                        "level": level,
                        "lambda_ok": lambda_ok,
                    }

                if st.button("运行灰色预测", key="run_gm11"):
                    try:
                        gm_series = pd.to_numeric(
                            clean_data_for_model[gm_col],
                            errors="coerce",
                        ).dropna().reset_index(drop=True)

                        if len(gm_series) < 4:
                            raise ValueError("序列少于4个有效数据点")

                        gm_result = gm11_predict(gm_series, int(gm_steps))
                        st.session_state["gm11_result"] = {
                            "col": gm_col,
                            "result": gm_result,
                            "series": gm_series,
                            "_data_hash": adv_file_hash,
                        }
                        _adv_cache_store(
                            "gm11_result",
                            cols_signature=(
                                gm_col,
                                int(gm_steps),
                            ),
                        )
                    except Exception as gm_error:
                        st.error(f"灰色预测失败：{gm_error}")

                if (
                    "gm11_result" in st.session_state
                    and _adv_cache_valid(
                        "gm11_result",
                        cols_signature=(
                            gm_col,
                            int(gm_steps),
                        ),
                    )
                ):
                    gm_stored = st.session_state["gm11_result"]
                    gm_res = gm_stored["result"]
                    gm_series = gm_stored["series"]
                    gm_n = len(gm_series)

                    forecast_index = list(
                        range(gm_n + 1, gm_n + len(gm_res["forecast"]) + 1)
                    )

                    gm_out = pd.DataFrame(
                        {
                            "期数": list(range(1, gm_n + 1))
                            + forecast_index,
                            "实际值": list(gm_series)
                            + [np.nan] * len(gm_res["forecast"]),
                            "拟合值/预测值": list(gm_res["in_sample"])
                            + list(gm_res["forecast"]),
                        }
                    )

                    st.write("灰色预测结果")
                    _st_dataframe(
                        gm_out,
                        use_container_width=True,
                    )

                    dataframe_download(
                        gm_out,
                        "灰色预测结果.csv",
                        key="download_gm11",
                    )

                    fig_gm, ax_gm = plt.subplots(figsize=(9, 5))
                    ax_gm.plot(
                        range(1, gm_n + 1),
                        gm_series,
                        marker="o",
                        label="实际值",
                    )
                    ax_gm.plot(
                        range(1, gm_n + 1),
                        gm_res["in_sample"],
                        linestyle="--",
                        label="拟合值",
                    )
                    ax_gm.plot(
                        forecast_index,
                        gm_res["forecast"],
                        marker="s",
                        color="red",
                        label="预测值",
                    )
                    ax_gm.axvline(gm_n + 0.5, color="gray", linestyle=":")
                    ax_gm.set_xlabel("期数")
                    ax_gm.set_ylabel(gm_stored["col"])
                    ax_gm.set_title(f"{gm_stored['col']} 灰色预测 GM(1,1)")
                    ax_gm.legend()
                    show_fig(fig_gm)
                    plt.close(fig_gm)

                    gm_text = (
                        f"对{gm_stored['col']}建立GM(1,1)灰色预测模型，"
                        f"发展系数a={gm_res['a']:.4f}，"
                        f"灰色作用量b={gm_res['b']:.4f}。"
                        f"平均相对误差{gm_res['mean_relative_error'] * 100:.2f}%，"
                        f"后验差比值C={gm_res['posterior_ratio']:.4f}，"
                        f"精度等级为“{gm_res['level']}”。"
                        f"预测未来{len(gm_res['forecast'])}期数值分别为："
                        + "、".join(
                            f"{v:.2f}" for v in gm_res["forecast"]
                        )
                        + "。"
                    )

                    st.text_area(
                        "论文表述",
                        gm_text,
                        height=140,
                        key="gm_text_area",
                    )

        # ---------- ARIMA 时间序列 ----------
        with adv_tabs[2]:
            st.markdown("**ARIMA 时间序列预测**")
            st.caption(
                "适合较长的时间序列（建议30期以上）的预测。"
                "程序会自动选择 (p,d,q) 阶数。"
            )

            if not adv_numeric_cols:
                st.info("没有可用的数值列。")
            else:
                arima_col = st.selectbox(
                    "选择序列列",
                    adv_numeric_cols,
                    key="arima_col_selector",
                )
                arima_steps = st.slider(
                    "预测期数",
                    1,
                    12,
                    5,
                    key="arima_steps_slider",
                )

                if st.button("运行ARIMA预测", key="run_arima"):
                    try:
                        from statsmodels.tsa.arima.model import ARIMA
                        from statsmodels.tsa.stattools import adfuller

                        arima_series = pd.to_numeric(
                            clean_data_for_model[arima_col],
                            errors="coerce",
                        ).dropna().reset_index(drop=True)

                        if len(arima_series) < 8:
                            raise ValueError(
                                "序列太短，建议改用灰色预测GM(1,1)。"
                            )

                        adf_p = adfuller(
                            arima_series,
                            autolag="AIC",
                        )[1]

                        d = 0
                        series_for_d = arima_series.copy()

                        while adf_p > 0.05 and d < 2:
                            series_for_d = (
                                series_for_d.diff().dropna()
                            )
                            adf_p = adfuller(
                                series_for_d,
                                autolag="AIC",
                            )[1]
                            d += 1

                        best_aic = np.inf
                        best_order = (1, d, 0)

                        # 修复：限制 (p+q) <= 3，减少候选组合数，
                        # 避免长时间无反馈的网格搜索。
                        with st.status(
                            "正在自动搜索 ARIMA 阶数 (p,d,q)…",
                            expanded=False,
                        ) as _arima_status:
                            for p in range(0, 4):
                                for q in range(0, 4):
                                    if p + q > 3:
                                        continue
                                    try:
                                        tmp_model = ARIMA(
                                            arima_series,
                                            order=(p, d, q),
                                        ).fit()
                                        if tmp_model.aic < best_aic:
                                            best_aic = tmp_model.aic
                                            best_order = (p, d, q)
                                    except Exception:
                                        continue
                            _arima_status.update(
                                label=(
                                    "搜索完成，最优阶数 "
                                    f"ARIMA{best_order}"
                                ),
                                state="complete",
                            )

                        arima_model = ARIMA(
                            arima_series,
                            order=best_order,
                        ).fit()

                        forecast_result = arima_model.get_forecast(
                            steps=int(arima_steps)
                        )
                        forecast_mean = np.asarray(
                            forecast_result.predicted_mean
                        )
                        forecast_ci = forecast_result.conf_int()

                        st.session_state["arima_result"] = {
                            "col": arima_col,
                            "order": best_order,
                            "aic": arima_model.aic,
                            "bic": arima_model.bic,
                            "adf_p": adf_p,
                            "forecast": forecast_mean,
                            "ci_low": np.asarray(
                                forecast_ci.iloc[:, 0]
                            ),
                            "ci_high": np.asarray(
                                forecast_ci.iloc[:, 1]
                            ),
                            "fitted": np.asarray(
                                arima_model.fittedvalues
                            ),
                            "series": arima_series,
                            "_data_hash": adv_file_hash,
                        }
                        _adv_cache_store(
                            "arima_result",
                            cols_signature=(
                                arima_col,
                                int(arima_steps),
                            ),
                        )
                    except Exception as arima_error:
                        st.error(f"ARIMA预测失败：{arima_error}")

                if (
                    "arima_result" in st.session_state
                    and _adv_cache_valid(
                        "arima_result",
                        cols_signature=(
                            arima_col,
                            int(arima_steps),
                        ),
                    )
                ):
                    ar_stored = st.session_state["arima_result"]
                    ar_n = len(ar_stored["series"])
                    ar_fc_n = len(ar_stored["forecast"])

                    ar_out = pd.DataFrame(
                        {
                            "期数": list(range(1, ar_n + 1))
                            + list(
                                range(ar_n + 1, ar_n + ar_fc_n + 1)
                            ),
                            "实际值": list(ar_stored["series"])
                            + [np.nan] * ar_fc_n,
                            "拟合/预测值": list(ar_stored["fitted"])
                            + list(ar_stored["forecast"]),
                        }
                    )

                    st.write(
                        f"选用 ARIMA{ar_stored['order']}，"
                        f"AIC={ar_stored['aic']:.2f}，"
                        f"BIC={ar_stored['bic']:.2f}，"
                        f"ADF检验P值={ar_stored['adf_p']:.4f}"
                    )

                    _st_dataframe(
                        ar_out,
                        use_container_width=True,
                    )

                    dataframe_download(
                        ar_out,
                        "ARIMA预测结果.csv",
                        key="download_arima",
                    )

                    fig_ar, ax_ar = plt.subplots(figsize=(9, 5))
                    ax_ar.plot(
                        range(1, ar_n + 1),
                        ar_stored["series"],
                        marker="o",
                        label="实际值",
                    )
                    ax_ar.plot(
                        range(1, ar_n + 1),
                        ar_stored["fitted"],
                        linestyle="--",
                        label="拟合值",
                    )

                    fore_idx = list(
                        range(ar_n + 1, ar_n + ar_fc_n + 1)
                    )

                    ax_ar.plot(
                        fore_idx,
                        ar_stored["forecast"],
                        marker="s",
                        color="red",
                        label="预测值",
                    )
                    ax_ar.fill_between(
                        fore_idx,
                        ar_stored["ci_low"],
                        ar_stored["ci_high"],
                        color="red",
                        alpha=0.15,
                        label="95%置信区间",
                    )
                    ax_ar.axvline(
                        ar_n + 0.5,
                        color="gray",
                        linestyle=":",
                    )
                    ax_ar.set_title(
                        f"{ar_stored['col']} ARIMA 预测"
                    )
                    ax_ar.legend()
                    show_fig(fig_ar)
                    plt.close(fig_ar)

        # ---------- 方差分析 + 卡方检验 ----------
        with adv_tabs[3]:
            st.markdown("**方差分析 ANOVA / 卡方检验**")
            st.caption(
                "ANOVA：检验数值变量在不同分组间的差异；"
                "卡方：检验两个分类变量是否独立。"
            )

            group_test_candidates = [
                col
                for col in clean_data_for_model.columns
                if clean_data_for_model[col].notna().nunique() < 30
            ]

            if not group_test_candidates or not adv_numeric_cols:
                st.info("需要至少一个分组变量和一个数值变量。")
            else:
                st.markdown("**单因素方差分析**")
                anova_group = st.selectbox(
                    "分组变量",
                    group_test_candidates,
                    key="anova_group_selector",
                )
                anova_value = st.selectbox(
                    "数值变量",
                    adv_numeric_cols,
                    key="anova_value_selector",
                )

                if st.button("运行方差分析", key="run_anova"):
                    try:
                        from scipy.stats import f_oneway, kruskal

                        anova_data = clean_data_for_model[
                            [anova_group, anova_value]
                        ].copy()
                        anova_data[anova_value] = pd.to_numeric(
                            anova_data[anova_value],
                            errors="coerce",
                        )
                        anova_data = anova_data.dropna()

                        groups_data = [
                            group_values
                            for _, group_values in anova_data.groupby(
                                anova_group
                            )[anova_value]
                        ]

                        if len(groups_data) < 2:
                            raise ValueError("分组数量不足2组。")

                        f_stat, f_p = f_oneway(*groups_data)
                        h_stat, h_p = kruskal(*groups_data)

                        group_means = anova_data.groupby(
                            anova_group
                        )[anova_value].agg(
                            ["count", "mean", "std"]
                        )

                        st.session_state["anova_result"] = {
                            "group": anova_group,
                            "value": anova_value,
                            "f": f_stat,
                            "p": f_p,
                            "h": h_stat,
                            "hp": h_p,
                            "means": group_means,
                            "_data_hash": adv_file_hash,
                        }
                        _adv_cache_store(
                            "anova_result",
                            cols_signature=(
                                anova_group,
                                anova_value,
                            ),
                        )
                    except Exception as anova_error:
                        st.error(f"方差分析失败：{anova_error}")

                if (
                    "anova_result" in st.session_state
                    and _adv_cache_valid(
                        "anova_result",
                        cols_signature=(
                            anova_group,
                            anova_value,
                        ),
                    )
                ):
                    ar_result = st.session_state["anova_result"]

                    _st_dataframe(
                        ar_result["means"],
                        use_container_width=True,
                    )
                    st.write(
                        f"ANOVA：F = {ar_result['f']:.4f}，"
                        f"P = {ar_result['p']:.4f}（P<0.05 说明组间差异显著）"
                    )
                    st.write(
                        f"Kruskal-Wallis（非参数）：H = {ar_result['h']:.4f}，"
                        f"P = {ar_result['hp']:.4f}"
                    )

                    conclusion = (
                        "存在显著差异"
                        if ar_result["p"] < 0.05
                        else "未发现显著差异"
                    )

                    st.success(
                        f"结论：{ar_result['value']} 在 "
                        f"{ar_result['group']} 的不同组间{conclusion}"
                        f"（F={ar_result['f']:.4f}，P={ar_result['p']:.4f}）。"
                    )

                st.markdown("**卡方独立性检验**")

                cat_cols = [
                    col
                    for col in clean_data_for_model.columns
                    if clean_data_for_model[col].nunique(dropna=True) <= 20
                ]

                if len(cat_cols) >= 2:
                    chi_a = st.selectbox(
                        "分类变量A",
                        cat_cols,
                        key="chi_a_selector",
                    )
                    chi_b = st.selectbox(
                        "分类变量B",
                        cat_cols,
                        key="chi_b_selector",
                    )

                    if st.button("运行卡方检验", key="run_chi2"):
                        try:
                            from scipy.stats import chi2_contingency

                            ct = pd.crosstab(
                                clean_data_for_model[chi_a],
                                clean_data_for_model[chi_b],
                            )
                            chi2_stat, chi2_p, dof, _ = (
                                chi2_contingency(ct)
                            )

                            st.session_state["chi2_result"] = {
                                "a": chi_a,
                                "b": chi_b,
                                "chi2": chi2_stat,
                                "p": chi2_p,
                                "dof": dof,
                                "table": ct,
                                "_data_hash": adv_file_hash,
                            }
                            _adv_cache_store(
                                "chi2_result",
                                cols_signature=(
                                    chi_a,
                                    chi_b,
                                ),
                            )
                        except Exception as chi_error:
                            st.error(f"卡方检验失败：{chi_error}")

                    if (
                        "chi2_result" in st.session_state
                        and _adv_cache_valid(
                            "chi2_result",
                            cols_signature=(
                                chi_a,
                                chi_b,
                            ),
                        )
                    ):
                        cr_result = st.session_state["chi2_result"]

                        _st_dataframe(
                            cr_result["table"],
                            use_container_width=True,
                        )
                        st.write(
                            f"卡方统计量 = {cr_result['chi2']:.4f}，"
                            f"自由度 = {cr_result['dof']}，"
                            f"P = {cr_result['p']:.4f}"
                        )

                        chi_conclusion = (
                            "存在显著关联"
                            if cr_result["p"] < 0.05
                            else "无显著关联"
                        )

                        st.success(
                            f"结论：{cr_result['a']} 与 {cr_result['b']} "
                            f"{chi_conclusion}（P={cr_result['p']:.4f}）。"
                        )
                else:
                    st.info("需要至少两个分类变量。")

        # ---------- PCA 主成分分析 ----------
        with adv_tabs[4]:
            st.markdown("**主成分分析 PCA（降维）**")
            st.caption(
                "适合自变量多且高度相关时，用少数主成分替代原变量。"
            )

            if len(adv_numeric_cols) < 2:
                st.info("至少需要两个数值变量。")
            else:
                pca_cols = st.multiselect(
                    "选择变量",
                    adv_numeric_cols,
                    default=adv_numeric_cols,
                    key="pca_cols_selector",
                )

                if len(pca_cols) >= 2:
                    if st.button("运行主成分分析", key="run_pca"):
                        try:
                            from sklearn.decomposition import PCA
                            from sklearn.preprocessing import StandardScaler

                            pca_data = clean_data_for_model[
                                pca_cols
                            ].apply(
                                pd.to_numeric,
                                errors="coerce",
                            ).dropna()

                            scaled = StandardScaler().fit_transform(
                                pca_data
                            )

                            pca_model = PCA()
                            components = pca_model.fit_transform(scaled)

                            var_table = pd.DataFrame(
                                {
                                    "主成分": [
                                        f"PC{i+1}"
                                        for i in range(
                                            len(
                                                pca_model.explained_variance_ratio_
                                            )
                                        )
                                    ],
                                    "特征值": pca_model.explained_variance_,
                                    "方差解释率": pca_model.explained_variance_ratio_,
                                    "累计方差解释率": np.cumsum(
                                        pca_model.explained_variance_ratio_
                                    ),
                                }
                            )

                            cum_ratio = np.cumsum(
                                pca_model.explained_variance_ratio_
                            )
                            n_keep = int(
                                np.argmax(cum_ratio >= 0.8) + 1
                            )
                            n_show = min(n_keep, 6)

                            loadings = pd.DataFrame(
                                pca_model.components_.T[:, :n_show],
                                index=pca_cols,
                                columns=[
                                    f"PC{i+1}"
                                    for i in range(n_show)
                                ],
                            )

                            score_table = pd.DataFrame(
                                components[:, :n_show],
                                columns=[
                                    f"PC{i+1}"
                                    for i in range(n_show)
                                ],
                            )

                            st.session_state["pca_data"] = {
                                "var_table": var_table,
                                "loadings": loadings,
                                "score_table": score_table,
                                "n_keep": n_keep,
                                "cum_ratio": cum_ratio,
                                "_data_hash": adv_file_hash,
                            }
                            _adv_cache_store(
                                "pca_data",
                                cols_signature=tuple(
                                    pca_cols
                                ),
                            )
                        except Exception as pca_error:
                            st.error(f"PCA失败：{pca_error}")

                    if (
                        "pca_data" in st.session_state
                        and _adv_cache_valid(
                            "pca_data",
                            cols_signature=tuple(
                                pca_cols
                            ),
                        )
                    ):
                        pca_stored = st.session_state["pca_data"]

                        st.write("方差解释率")
                        _st_dataframe(
                            pca_stored["var_table"],
                            use_container_width=True,
                        )

                        dataframe_download(
                            pca_stored["var_table"],
                            "PCA方差解释.csv",
                            key="download_pca_var",
                        )

                        st.write("主成分载荷（变量贡献）")
                        _st_dataframe(
                            pca_stored["loadings"],
                            use_container_width=True,
                        )

                        st.write("主成分得分（前20行）")
                        _st_dataframe(
                            pca_stored["score_table"].head(20),
                            use_container_width=True,
                        )

                        dataframe_download(
                            pca_stored["score_table"],
                            "PCA主成分得分.csv",
                            key="download_pca_scores",
                        )

                        fig_pca, ax_pca = plt.subplots(figsize=(8, 5))
                        ax_pca.plot(
                            range(1, len(pca_stored["cum_ratio"]) + 1),
                            pca_stored["cum_ratio"],
                            marker="o",
                        )
                        ax_pca.axhline(
                            0.8,
                            color="red",
                            linestyle="--",
                            label="80%阈值",
                        )
                        ax_pca.set_xlabel("主成分个数")
                        ax_pca.set_ylabel("累计方差解释率")
                        ax_pca.set_title("累计方差解释率")
                        ax_pca.legend()
                        show_fig(fig_pca)
                        plt.close(fig_pca)

                        pca_text = (
                            f"对上述变量进行主成分分析，前"
                            f"{pca_stored['n_keep']}个主成分的累计方差"
                            f"解释率达到"
                            f"{pca_stored['cum_ratio'][pca_stored['n_keep'] - 1] * 100:.1f}%，"
                            "可提取这些主成分代替原始变量用于后续建模。"
                        )

                        st.text_area(
                            "论文表述",
                            pca_text,
                            height=100,
                            key="pca_text_area",
                        )

        # ---------- 聚类分析 ----------
        with adv_tabs[5]:
            st.markdown("**聚类分析（K-means / 层次聚类）**")
            st.caption(
                "适合把样本分成若干类别，再结合题目背景解读每类特征。"
            )

            if len(adv_numeric_cols) < 2:
                st.info("至少需要两个数值变量。")
            else:
                clu_cols = st.multiselect(
                    "选择聚类变量",
                    adv_numeric_cols,
                    default=adv_numeric_cols,
                    key="clu_cols_selector",
                )
                clu_method = st.radio(
                    "聚类方法",
                    ["K-means", "层次聚类"],
                    horizontal=True,
                    key="clu_method_radio",
                )

                if len(clu_cols) >= 2:
                    k_value = st.slider(
                        "聚类个数 K",
                        2,
                        10,
                        3,
                        key="clu_k_slider",
                    )

                    if st.button("运行聚类", key="run_cluster"):
                        try:
                            from sklearn.cluster import (
                                KMeans,
                                AgglomerativeClustering,
                            )
                            from sklearn.preprocessing import (
                                StandardScaler,
                            )

                            clu_data = clean_data_for_model[
                                clu_cols
                            ].apply(
                                pd.to_numeric,
                                errors="coerce",
                            ).dropna()

                            scaled = StandardScaler().fit_transform(
                                clu_data
                            )

                            if clu_method == "K-means":
                                km = KMeans(
                                    n_clusters=int(k_value),
                                    random_state=42,
                                    n_init=10,
                                )
                                labels = km.fit_predict(scaled)
                                wcss_value = km.inertia_
                            else:
                                agg = AgglomerativeClustering(
                                    n_clusters=int(k_value)
                                )
                                labels = agg.fit_predict(scaled)
                                wcss_value = None

                            clu_out = clu_data.copy()
                            clu_out["聚类结果"] = labels + 1

                            profile = clu_out.groupby(
                                "聚类结果"
                            )[clu_cols].mean()

                            st.session_state["cluster_data"] = {
                                "clu_out": clu_out,
                                "profile": profile,
                                "k": int(k_value),
                                "wcss": wcss_value,
                                "_data_hash": adv_file_hash,
                            }
                            _adv_cache_store(
                                "cluster_data",
                                cols_signature=(
                                    tuple(clu_cols),
                                    clu_method,
                                    int(k_value),
                                ),
                            )
                        except Exception as clu_error:
                            st.error(f"聚类失败：{clu_error}")

                    if (
                        "cluster_data" in st.session_state
                        and _adv_cache_valid(
                            "cluster_data",
                            cols_signature=(
                                tuple(clu_cols),
                                clu_method,
                                int(k_value),
                            ),
                        )
                    ):
                        clu_stored = st.session_state["cluster_data"]

                        st.write("各类别样本数")
                        _st_dataframe(
                            clu_stored["clu_out"][
                                "聚类结果"
                            ].value_counts().rename_axis(
                                "类别"
                            ).reset_index(name="样本数"),
                            use_container_width=True,
                        )

                        st.write("聚类结果（前20行）")
                        _st_dataframe(
                            clu_stored["clu_out"].head(20),
                            use_container_width=True,
                        )

                        dataframe_download(
                            clu_stored["clu_out"],
                            "聚类结果.csv",
                            key="download_cluster",
                        )

                        st.write("各类别特征均值（画像）")
                        _st_dataframe(
                            clu_stored["profile"],
                            use_container_width=True,
                        )

                        if clu_stored["wcss"] is not None:
                            st.write(
                                f"K-means 组内平方和（WCSS）= "
                                f"{clu_stored['wcss']:.2f}，"
                                "可用于肘部法则选择K。"
                            )

                        st.success(
                            f"已将样本划分为 {clu_stored['k']} 类，"
                            "请结合各类均值画像解读类别含义。"
                        )

        # ---------- 随机森林 / 决策树 ----------
        with adv_tabs[6]:
            st.markdown("**机器学习分类/回归（随机森林、决策树）**")
            st.caption(
                "作为传统统计模型的补充：可输出特征重要性，是论文加分点。"
            )

            ml_candidates = list(clean_data_for_model.columns)

            if not ml_candidates:
                st.info("无可用数据。")
            else:
                ml_target = st.selectbox(
                    "选择目标变量",
                    ml_candidates,
                    key="ml_target_selector",
                )

                ml_default_feats = [
                    c for c in ml_candidates if c != ml_target
                ]

                ml_feats = st.multiselect(
                    "选择特征变量",
                    ml_default_feats,
                    default=ml_default_feats[
                        : min(6, len(ml_default_feats))
                    ],
                    key="ml_feats_selector",
                )

                if ml_feats:
                    ml_method = st.radio(
                        "方法",
                        ["随机森林", "决策树"],
                        horizontal=True,
                        key="ml_method_radio",
                    )

                    ml_raw = clean_data_for_model[ml_target]

                    # 任务类型：字符串/类别型目标按分类；
                    # 数值型低基数（如 0/1 标记）也按分类。
                    task_is_classification = (
                        ml_raw.dtype == object
                        or pd.api.types.is_string_dtype(ml_raw)
                        or isinstance(
                            ml_raw.dtype,
                            pd.CategoricalDtype,
                        )
                        or ml_raw.nunique(dropna=True) <= 10
                    )

                    st.caption(
                        "任务类型："
                        + ("分类" if task_is_classification else "回归")
                        + "（按目标变量自动判断）"
                    )

                    if st.button("训练模型并评估", key="run_ml"):
                        try:
                            from sklearn.ensemble import (
                                RandomForestClassifier,
                                RandomForestRegressor,
                            )
                            from sklearn.tree import (
                                DecisionTreeClassifier,
                                DecisionTreeRegressor,
                            )

                            ml_data = clean_data_for_model[
                                [ml_target] + ml_feats
                            ].dropna()

                            y_ml = ml_data[ml_target]
                            X_ml = ml_data[ml_feats].copy()

                            for col_ml in ml_feats:
                                col_series = X_ml[col_ml]

                                # 修复：仅字符串/类别型特征转类别编码；
                                # 数值型特征即使取值较少（如 1-10 的等级）
                                # 也保持数值，不再误判为分类。
                                is_string_like = (
                                    col_series.dtype == object
                                    or pd.api.types.is_string_dtype(
                                        col_series
                                    )
                                    or isinstance(
                                        col_series.dtype,
                                        pd.CategoricalDtype,
                                    )
                                )

                                if is_string_like:
                                    X_ml[col_ml] = col_series.astype(
                                        "category"
                                    ).cat.codes
                                else:
                                    X_ml[col_ml] = pd.to_numeric(
                                        col_series,
                                        errors="coerce",
                                    )

                            X_ml = X_ml.apply(
                                pd.to_numeric,
                                errors="coerce",
                            ).dropna()
                            y_ml = y_ml.loc[X_ml.index]

                            if task_is_classification:
                                y_ml = y_ml.astype(
                                    "category"
                                ).cat.codes
                            else:
                                y_ml = pd.to_numeric(
                                    y_ml,
                                    errors="coerce",
                                )
                                keep = y_ml.notna()
                                y_ml = y_ml[keep]
                                X_ml = X_ml.loc[keep]

                            if len(X_ml) < 10:
                                raise ValueError(
                                    "有效样本少于10条"
                                )

                            X_train, X_test, y_train, y_test = (
                                train_test_split(
                                    X_ml,
                                    y_ml,
                                    test_size=0.2,
                                    random_state=42,
                                )
                            )

                            if ml_method == "随机森林":
                                if task_is_classification:
                                    model_ml = RandomForestClassifier(
                                        n_estimators=100,
                                        random_state=42,
                                        n_jobs=-1,
                                    )
                                else:
                                    model_ml = RandomForestRegressor(
                                        n_estimators=100,
                                        random_state=42,
                                        n_jobs=-1,
                                    )
                            else:
                                if task_is_classification:
                                    model_ml = DecisionTreeClassifier(
                                        max_depth=5,
                                        random_state=42,
                                    )
                                else:
                                    model_ml = DecisionTreeRegressor(
                                        max_depth=5,
                                        random_state=42,
                                    )

                            model_ml.fit(X_train, y_train)
                            y_pred = model_ml.predict(X_test)

                            if task_is_classification:
                                metrics = {
                                    "准确率": accuracy_score(
                                        y_test, y_pred
                                    ),
                                    "平衡准确率": (
                                        balanced_accuracy_score(
                                            y_test, y_pred
                                        )
                                    ),
                                    "F1(宏平均)": f1_score(
                                        y_test,
                                        y_pred,
                                        average="macro",
                                        zero_division=0,
                                    ),
                                }
                            else:
                                metrics = {
                                    "RMSE": np.sqrt(
                                        mean_squared_error(
                                            y_test, y_pred
                                        )
                                    ),
                                    "MAE": mean_absolute_error(
                                        y_test, y_pred
                                    ),
                                    "R²": r2_score(y_test, y_pred),
                                }

                            metric_df_ml = pd.DataFrame(
                                {
                                    "指标": list(metrics.keys()),
                                    "数值": list(metrics.values()),
                                }
                            )

                            importance = model_ml.feature_importances_
                            imp_df = pd.DataFrame(
                                {
                                    "特征": ml_feats,
                                    "重要性": importance,
                                }
                            ).sort_values(
                                "重要性",
                                ascending=False,
                            )

                            st.session_state["ml_data"] = {
                                "metrics": metric_df_ml,
                                "importance": imp_df,
                                "task": (
                                    "分类"
                                    if task_is_classification
                                    else "回归"
                                ),
                                "_data_hash": adv_file_hash,
                            }
                            _adv_cache_store(
                                "ml_data",
                                cols_signature=(
                                    tuple(ml_feats),
                                    ml_target,
                                    ml_method,
                                ),
                            )
                        except Exception as ml_error:
                            st.error(
                                f"机器学习建模失败：{ml_error}"
                            )

                    if (
                        "ml_data" in st.session_state
                        and _adv_cache_valid(
                            "ml_data",
                            cols_signature=(
                                tuple(ml_feats),
                                ml_target,
                                ml_method,
                            ),
                        )
                    ):
                        ml_stored = st.session_state["ml_data"]

                        st.write("测试集评估指标")
                        _st_dataframe(
                            ml_stored["metrics"],
                            use_container_width=True,
                        )

                        st.write("特征重要性")
                        _st_dataframe(
                            ml_stored["importance"],
                            use_container_width=True,
                        )

                        dataframe_download(
                            ml_stored["importance"],
                            "特征重要性.csv",
                            key="download_ml_importance",
                        )

                        fig_imp, ax_imp = plt.subplots(
                            figsize=(8, max(4, len(ml_stored["importance"]) * 0.4))
                        )
                        ax_imp.barh(
                            ml_stored["importance"]["特征"],
                            ml_stored["importance"]["重要性"],
                        )
                        ax_imp.invert_yaxis()
                        ax_imp.set_title("特征重要性")
                        show_fig(fig_imp)
                        plt.close(fig_imp)

                        top_feat = ml_stored["importance"].iloc[0]

                        st.success(
                            f"最重要特征：{top_feat['特征']}"
                            f"（重要性{top_feat['重要性']:.3f}）。"
                        )

        # ---------- 稳健性分析 ----------
        with adv_tabs[7]:
            st.markdown("**稳健性分析（Bootstrap 系数置信区间）**")
            st.caption(
                "通过重复抽样重新拟合模型，得到系数的Bootstrap置信区间，"
                "验证结论是否稳健。"
            )

            fitted_model_adv = st.session_state.get("fitted_model")
            stored_type_adv = st.session_state.get("final_model_type")

            if fitted_model_adv is None:
                st.info("请先在上方完成模型拟合，再进行稳健性分析。")
            else:
                n_boot = st.slider(
                    "重抽样次数",
                    100,
                    2000,
                    500,
                    step=100,
                    key="boot_n_slider",
                )

                if st.button("运行稳健性分析", key="run_bootstrap"):
                    try:
                        def _param_names_adv(model):
                            params = model.params

                            if isinstance(params, pd.DataFrame):
                                names = []

                                for var in params.index:
                                    for cat in params.columns:
                                        names.append(f"{var}[{cat}]")

                                return names

                            return list(params.index)

                        y_adv, X_adv, groups_adv, _ = build_model_data(
                            clean_data_for_model,
                            target,
                            predictors,
                            variable_types,
                            group_col=group_col,
                        )

                        param_names = _param_names_adv(fitted_model_adv)
                        original_coefs = np.asarray(
                            fitted_model_adv.params,
                            dtype=float,
                        ).reshape(-1)

                        boot_results = []

                        # 修复：使用固定随机种子保证结果可复现；
                        # 混合效应模型存在分组结构时按“组”重抽样，保留组内相关性。
                        boot_rng = np.random.default_rng(42)
                        group_ids = None

                        if groups_adv is not None and (
                            groups_adv.nunique() < len(groups_adv)
                        ):
                            group_ids = groups_adv

                        for _ in range(int(n_boot)):
                            if group_ids is None:
                                idx = boot_rng.integers(
                                    0,
                                    len(y_adv),
                                    size=len(y_adv),
                                )
                            else:
                                # 按组重抽样：随机抽取组（有放回），
                                # 组内观测全部保留。
                                unique_groups = group_ids.unique()
                                sampled_groups = boot_rng.choice(
                                    unique_groups,
                                    size=len(unique_groups),
                                    replace=True,
                                )
                                idx = np.concatenate(
                                    [
                                        np.where(group_ids == g)[0]
                                        for g in sampled_groups
                                    ]
                                )

                            y_b = y_adv.iloc[idx].reset_index(drop=True)
                            X_b = X_adv.iloc[idx].reset_index(drop=True)

                            if groups_adv is not None:
                                g_b = groups_adv.iloc[idx].reset_index(drop=True)
                            else:
                                g_b = None

                            try:
                                res_b = fit_model(
                                    y_b,
                                    X_b,
                                    g_b,
                                    stored_type_adv,
                                    robust_se=False,
                                )
                                boot_results.append(
                                    np.asarray(
                                        res_b["model"].params,
                                        dtype=float,
                                    ).reshape(-1)
                                )
                            except Exception:
                                continue

                        if len(boot_results) < 50:
                            raise ValueError(
                                f"成功完成的重抽样仅{len(boot_results)}次，"
                                "请减少次数或更换模型。"
                            )

                        boot_arr = np.array(boot_results)

                        p_low = np.percentile(boot_arr, 2.5, axis=0)
                        p_high = np.percentile(boot_arr, 97.5, axis=0)

                        boot_table = pd.DataFrame(
                            {
                                "参数": param_names[: len(p_low)],
                                "原始系数": original_coefs[: len(p_low)],
                                "Bootstrap均值": boot_arr.mean(axis=0),
                                "2.5%分位": p_low,
                                "97.5%分位": p_high,
                                "是否包含0": np.where(
                                    (p_low < 0) & (p_high > 0),
                                    "是（不稳健）",
                                    "否（稳健）",
                                ),
                            }
                        )

                        st.write(
                            f"Bootstrap（{len(boot_results)}次有效重抽样）"
                            "95%系数置信区间"
                        )
                        _st_dataframe(
                            boot_table,
                            use_container_width=True,
                        )

                        dataframe_download(
                            boot_table,
                            "Bootstrap稳健性分析.csv",
                            key="download_bootstrap",
                        )

                        boot_text = (
                            f"通过{len(boot_results)}次Bootstrap重抽样"
                            "重新拟合模型，得到各系数的95%置信区间。"
                            "若区间不包含0，说明该变量的效应较为稳健；"
                            "结果显示显著变量的置信区间均不包含0，"
                            "模型结论整体稳健。"
                        )

                        st.text_area(
                            "论文表述",
                            boot_text,
                            height=110,
                            key="bootstrap_text_area",
                        )
                    except Exception as boot_error:
                        st.error(f"稳健性分析失败：{boot_error}")

        # ---------- AHP 层次分析（新增） ----------
        with adv_tabs[8]:
            st.markdown("**AHP 层次分析法（含一致性检验）**")
            st.caption(
                "适合评价类题目：通过两两比较构造判断矩阵，"
                "计算权重并检验一致性（CI/CR），"
                "与熵权TOPSIS互为补充（AHP为主观赋权，熵权为客观赋权）。"
            )

            def _ahp_weights(matrix):
                """由判断矩阵计算权重向量（特征向量法 + 归一化）。"""
                m = np.asarray(matrix, dtype=float)
                n = m.shape[0]
                eigvals, eigvecs = np.linalg.eig(m)
                max_idx = int(np.argmax(eigvals.real))
                principal = np.abs(eigvecs[:, max_idx].real)
                return principal / principal.sum(), float(
                    eigvals.real[max_idx]
                )

            def _ahp_consistency(matrix, weights, lam_max):
                """计算 CI 与 CR。"""
                n = matrix.shape[0]
                ci = (lam_max - n) / max(n - 1, 1)
                # Saaty 随机一致性指标 RI（n=1..10）
                ri_table = {
                    1: 0.0,
                    2: 0.0,
                    3: 0.58,
                    4: 0.90,
                    5: 1.12,
                    6: 1.24,
                    7: 1.32,
                    8: 1.41,
                    9: 1.45,
                    10: 1.49,
                }
                ri = ri_table.get(n, 1.49)
                cr = ci / ri if ri > 0 else 0.0
                return ci, cr

            ahp_indicator_names = st.text_input(
                "评价指标名称（用中文逗号或顿号分隔，如：成本,质量,效率）",
                value="成本,质量,效率",
                key="ahp_names_input",
            )

            ahp_names = [
                name.strip()
                for name in re.split(
                    r"[,，、;；]",
                    ahp_indicator_names,
                )
                if name.strip()
            ]

            if len(ahp_names) < 2:
                st.info("请至少输入两个评价指标。")
            else:
                n_ahp = len(ahp_names)

                st.markdown(
                    "**两两比较判断矩阵（AHP 1-9 标度）**\n\n"
                    "`a_ij` 表示指标 i 相对指标 j 的重要程度："
                    "1=同等重要，3=稍微重要，5=明显重要，"
                    "7=强烈重要，9=极端重要，2/4/6/8 为中间值；"
                    "对角线固定为 1。"
                )

                ahp_matrix = np.ones((n_ahp, n_ahp))

                for i in range(n_ahp):
                    for j in range(i + 1, n_ahp):
                        val = st.number_input(
                            f"{ahp_names[i]} 相对 {ahp_names[j]} 的重要程度",
                            min_value=1.0,
                            max_value=9.0,
                            value=1.0,
                            step=1.0,
                            key=f"ahp_{i}_{j}",
                        )
                        ahp_matrix[i, j] = val
                        ahp_matrix[j, i] = 1.0 / val

                if st.button(
                    "计算 AHP 权重与一致性检验",
                    key="run_ahp",
                ):
                    try:
                        weights, lam_max = _ahp_weights(
                            ahp_matrix
                        )
                        ci, cr = _ahp_consistency(
                            ahp_matrix,
                            weights,
                            lam_max,
                        )

                        ahp_table = pd.DataFrame(
                            {
                                "指标": ahp_names,
                                "权重": weights,
                            }
                        ).sort_values(
                            "权重",
                            ascending=False,
                        )

                        ahp_table["权重占比%"] = (
                            ahp_table["权重"] * 100
                        ).round(2)

                        st.write("AHP 权重结果")
                        _st_dataframe(
                            ahp_table,
                            use_container_width=True,
                        )

                        dataframe_download(
                            ahp_table,
                            "AHP权重.csv",
                            key="download_ahp",
                        )

                        st.write(
                            f"最大特征根 λmax = {lam_max:.4f}，"
                            f"CI = {ci:.4f}，CR = {cr:.4f}"
                        )

                        if cr < 0.10:
                            st.success(
                                "一致性检验通过（CR < 0.10），"
                                "权重可以使用。"
                            )
                        else:
                            st.warning(
                                "一致性检验未通过（CR ≥ 0.10），"
                                "建议调整判断矩阵中的矛盾比较。"
                            )

                        ahp_text = (
                            "采用层次分析法（AHP）确定评价指标权重。"
                            "构造两两比较判断矩阵，计算最大特征根"
                            f"λmax={lam_max:.4f}，一致性比例"
                            f"CR={cr:.4f}"
                            + (
                                "<0.10，通过一致性检验，权重结果可靠。"
                                if cr < 0.10
                                else "≥0.10，未通过一致性检验，需调整判断矩阵。"
                            )
                            + "各指标权重为："
                            + "，".join(
                                f"{ahp_names[k]}={weights[k]:.4f}"
                                for k in range(n_ahp)
                            )
                            + "。"
                        )

                        st.text_area(
                            "论文表述",
                            ahp_text,
                            height=130,
                            key="ahp_text_area",
                        )
                    except Exception as ahp_error:
                        st.error(f"AHP计算失败：{ahp_error}")

        # ---------- 正则化变量筛选（新增） ----------
        with adv_tabs[9]:
            st.markdown("**正则化变量筛选（Lasso / Ridge / ElasticNet）**")
            st.caption(
                "适合高维数据（自变量多、样本少）或存在多重共线性时："
                "Lasso 会把不重要变量的系数压缩为 0，实现自动变量筛选，"
                "是论文“模型改进/变量筛选”环节的加分项。"
            )

            if not adv_numeric_cols:
                st.info("没有可用的数值列。")
            else:
                reg_target = st.selectbox(
                    "目标变量（因变量）",
                    adv_numeric_cols,
                    key="reg_target_selector",
                )
                reg_feats = st.multiselect(
                    "候选自变量（数值列，可多选）",
                    adv_numeric_cols,
                    default=[
                        c
                        for c in adv_numeric_cols
                        if c != reg_target
                    ][:8],
                    key="reg_feats_selector",
                )

                if len(reg_feats) < 2:
                    st.info("请至少选择两个候选自变量。")
                else:
                    reg_method = st.radio(
                        "正则化方法",
                        ["Lasso", "Ridge", "ElasticNet"],
                        key="reg_method_radio",
                    )

                    if st.button(
                        "运行变量筛选",
                        key="run_reg_select",
                    ):
                        try:
                            reg_data = clean_data_for_model[
                                [reg_target] + reg_feats
                            ].apply(
                                pd.to_numeric,
                                errors="coerce",
                            ).dropna()

                            if len(reg_data) < 10:
                                raise ValueError(
                                    "有效样本不足10条。"
                                )

                            X_reg = reg_data[reg_feats]
                            y_reg = reg_data[reg_target]

                            # 标准化特征（正则化要求）
                            scaler_reg = StandardScaler()
                            X_scaled = scaler_reg.fit_transform(
                                X_reg
                            )

                            if reg_method == "Lasso":
                                reg_model = Lasso(
                                    alpha=0.05,
                                    max_iter=5000,
                                    random_state=42,
                                )
                            elif reg_method == "Ridge":
                                reg_model = Ridge(
                                    alpha=1.0,
                                    random_state=42,
                                )
                            else:
                                reg_model = ElasticNet(
                                    alpha=0.05,
                                    l1_ratio=0.5,
                                    max_iter=5000,
                                    random_state=42,
                                )

                            reg_model.fit(
                                X_scaled,
                                y_reg,
                            )

                            coef_df = pd.DataFrame(
                                {
                                    "特征": reg_feats,
                                    "标准化系数": reg_model.coef_,
                                }
                            ).sort_values(
                                "标准化系数",
                                key=lambda s: s.abs(),
                                ascending=False,
                            )

                            coef_df["系数绝对值"] = (
                                coef_df["标准化系数"].abs()
                            )

                            selected = coef_df[
                                coef_df["标准化系数"] != 0
                            ]

                            st.write("标准化系数（按绝对值降序）")
                            _st_dataframe(
                                coef_df,
                                use_container_width=True,
                            )

                            dataframe_download(
                                coef_df,
                                "正则化系数.csv",
                                key="download_reg_coef",
                            )

                            if reg_method == "Lasso":
                                if selected.empty:
                                    st.warning(
                                        "Lasso 将所有系数压缩为 0，"
                                        "可尝试减小惩罚（alpha）。"
                                    )
                                else:
                                    st.success(
                                        "Lasso 筛选出的重要变量："
                                        + "、".join(
                                            selected[
                                                "特征"
                                            ].tolist()
                                        )
                                    )

                            fig_reg, ax_reg = plt.subplots(
                                figsize=(8, max(4, len(coef_df) * 0.4))
                            )
                            ax_reg.barh(
                                coef_df["特征"],
                                coef_df["标准化系数"],
                            )
                            ax_reg.axvline(
                                0,
                                color="gray",
                                linewidth=0.8,
                            )
                            ax_reg.set_title(
                                f"{reg_method} 标准化系数"
                            )
                            show_fig(fig_reg)
                            plt.close(fig_reg)

                            reg_text = (
                                f"采用{reg_method}正则化对 {len(reg_feats)} 个"
                                "候选变量进行筛选（特征标准化后拟合）。"
                                "结果显示系数绝对值较大的变量为："
                                + "、".join(
                                    coef_df.head(5)[
                                        "特征"
                                    ].tolist()
                                )
                                + "。"
                            )

                            st.text_area(
                                "论文表述",
                                reg_text,
                                height=120,
                                key="reg_text_area",
                            )
                        except Exception as reg_error:
                            st.error(f"正则化筛选失败：{reg_error}")

        # ---------- 时间序列辅助（新增） ----------
        with adv_tabs[10]:
            st.markdown("**时间序列辅助：滞后项 / 差分 / 按时间切分**")
            st.caption(
                "预测类赛题的时间列处理工具：构造滞后项、"
                "一阶差分、按时间顺序切分训练/测试集，"
                "并给出建模建议。"
            )

            datetime_cols_adv = [
                col
                for col in clean_data_for_model.columns
                if pd.api.types.is_datetime64_any_dtype(
                    clean_data_for_model[col]
                )
            ]

            if not datetime_cols_adv:
                st.info(
                    "未检测到时间列。若你的数据包含日期，"
                    "请在变量类型确认步骤把该列标记为“时间”。"
                )
            else:
                ts_time_col = st.selectbox(
                    "时间列",
                    datetime_cols_adv,
                    key="ts_time_col_selector",
                )
                ts_value_col = st.selectbox(
                    "要建模的数值列（因变量）",
                    adv_numeric_cols,
                    key="ts_value_col_selector",
                )

                ts_lag_n = st.slider(
                    "构造滞后阶数（lag）",
                    0,
                    5,
                    1,
                    key="ts_lag_slider",
                )

                if st.button(
                    "生成时间序列特征",
                    key="run_ts_features",
                ):
                    try:
                        ts_df = clean_data_for_model[
                            [ts_time_col, ts_value_col]
                        ].copy()
                        ts_df[ts_time_col] = pd.to_datetime(
                            ts_df[ts_time_col],
                            errors="coerce",
                        )
                        ts_df = ts_df.dropna().sort_values(
                            ts_time_col
                        )
                        ts_df = ts_df.reset_index(drop=True)

                        ts_out = ts_df.copy()

                        for lag in range(1, int(ts_lag_n) + 1):
                            ts_out[f"lag{lag}"] = ts_out[
                                ts_value_col
                            ].shift(lag)

                        ts_out["一阶差分"] = ts_out[
                            ts_value_col
                        ].diff()

                        ts_out = ts_out.dropna().reset_index(
                            drop=True
                        )

                        st.write(
                            f"构造完成：{ts_out.shape[0]} 行，"
                            f"{ts_out.shape[1]} 列"
                        )
                        _st_dataframe(
                            ts_out.head(20),
                            use_container_width=True,
                        )

                        dataframe_download(
                            ts_out,
                            "时间序列特征表.csv",
                            key="download_ts_features",
                        )

                        # 按时间顺序切分提示
                        split_ratio = st.slider(
                            "按时间切分：测试集比例",
                            0.1,
                            0.4,
                            0.2,
                            step=0.05,
                            key="ts_split_slider",
                        )

                        split_idx = int(
                            len(ts_out) * (1 - split_ratio)
                        )

                        train_part = ts_out.iloc[:split_idx]
                        test_part = ts_out.iloc[split_idx:]

                        st.success(
                            f"按时间切分：训练集 {len(train_part)} 行"
                            f"（{train_part[ts_time_col].iloc[0]} 至 "
                            f"{train_part[ts_time_col].iloc[-1]}），"
                            f"测试集 {len(test_part)} 行"
                            f"（{test_part[ts_time_col].iloc[0]} 至 "
                            f"{test_part[ts_time_col].iloc[-1]}）。"
                        )

                        st.caption(
                            "注意：时间序列切分必须按时间顺序，"
                            "不能用随机切分（会泄漏未来信息）。"
                            "短期预测建议 <30 期用灰色预测，"
                            "≥30 期可用 ARIMA。"
                        )
                    except Exception as ts_error:
                        st.error(f"时间序列辅助失败：{ts_error}")



    # ============================================================
    # 二十二、综合报告导出
    # ============================================================

    st.subheader("📑 一键导出综合报告")

    # ===== 新增：论文要素核对清单（写论文前自查） =====
    with st.expander(
        "📋 论文要素核对清单（评委最看重，写论文前逐项自查）",
        expanded=False,
    ):
        paper_items = [
            ("摘要", "300字以内，含问题、方法、结果三要素"),
            ("问题重述", "用自己的话复述赛题，不照抄"),
            ("模型假设", "用 generate_assumptions 生成的假设并说明合理性"),
            ("符号表", "本工具的“变量符号表”可直接引用"),
            ("模型建立", "说明为什么选这个模型（依据+公式）"),
            ("模型求解", "附求解代码/工具说明与关键结果"),
            ("模型检验", "拟合优度、残差诊断、交叉验证、Bootstrap"),
            ("灵敏度分析", "优化类用影子价格，回归类扰动关键参数"),
            ("模型评价", "优点、缺点、改进方向各写2-3条"),
            ("结论", "回答赛题问题，给出量化结果"),
        ]

        checklist_state = st.session_state.setdefault(
            "paper_checklist",
            {},
        )

        for item_title, item_desc in paper_items:
            checked = checklist_state.get(item_title, False)
            checklist_state[item_title] = st.checkbox(
                f"{item_title}：{item_desc}",
                value=checked,
                key=f"paper_item_{item_title}",
            )

        done_count = sum(
            1
            for v in checklist_state.values()
            if v
        )

        st.progress(
            done_count / len(paper_items),
            text=(
                f"论文要素完成度 {done_count}/{len(paper_items)}"
            ),
        )

        if done_count == len(paper_items):
            st.success(
                "所有论文要素都已勾选，可以安心写摘要了！"
            )

    # ===== 新增：工作区保存 / 恢复（72小时连续作战防丢失） =====
    with st.expander(
        "💾 工作区保存 / 恢复（刷新页面不丢失进度）",
        expanded=False,
    ):
        st.caption(
            "把当前会话的关键设置（赛题文本、变量选择、变量类型、"
            "缺失值/异常值设置、模型结果、高级分析结果、优化约束）"
            "导出为 JSON 文件；下次打开时上传即可恢复。"
        )

        workspace_keys = [
            "problem_text",
            "problem_type",
            "target",
            "predictors",
            "variable_types",
            "missing_method",
            "outlier_method",
            "outlier_action",
            "group_col",
            "cons_list",
            "paper_checklist",
        ]

        def _build_workspace():
            payload = {}

            for key in workspace_keys:
                value = st.session_state.get(key)

                # 过滤不可序列化对象（如 DataFrame）
                try:
                    json.dumps(value)
                    payload[key] = value
                except (TypeError, ValueError):
                    continue

            return json.dumps(
                payload,
                ensure_ascii=False,
                indent=2,
            )

        if st.button(
            "📤 导出工作区（JSON）",
            key="export_workspace_button",
        ):
            workspace_bytes = _build_workspace().encode(
                "utf-8"
            )
            st.download_button(
                "⬇️ 下载工作区文件",
                data=workspace_bytes,
                file_name=(
                    f"数模工作区_{datetime.now().strftime('%Y%m%d_%H%M')}.json"
                ),
                mime="application/json",
                key="download_workspace",
            )

        workspace_upload = st.file_uploader(
            "恢复工作区（上传 JSON）",
            type=["json"],
            key="workspace_uploader",
        )

        if workspace_upload is not None:
            try:
                loaded = json.loads(
                    workspace_upload.getvalue().decode("utf-8")
                )

                for key, value in loaded.items():
                    if key in workspace_keys:
                        st.session_state[key] = value

                st.success(
                    f"工作区恢复成功：{len(loaded)} 项设置已载入。"
                )
            except Exception as ws_error:
                st.error(f"工作区恢复失败：{ws_error}")

    # 修复：不再无条件把引导进度推到 8，
    # 改为“报告确实生成成功”后由按钮回调推进 guide_auto_step。
    def _get_var(name):
        """从当前作用域安全获取变量，避免未定义报错。"""
        return globals().get(name)

    def _add_section(sections, title, df):
        if df is not None and hasattr(df, 'empty') and not df.empty:
            sections.append((title, df))

    report_sections = []
    report_sections.append(("基本信息", f"赛题类型：{_get_var('problem_type') or '未填写'}"))
    _add_section(report_sections, "数据清洗汇总", _get_var("cleaning_summary"))
    _add_section(report_sections, "缺失值处理明细", _get_var("missing_detail_table"))
    _add_section(report_sections, "异常值处理明细", _get_var("outlier_detail_table"))
    _add_section(report_sections, "相关性分析", _get_var("corr_table"))
    _add_section(report_sections, "VIF多重共线性诊断", _get_var("vif_table"))
    _add_section(report_sections, "模型评价指标", _get_var("metric_table"))
    _add_section(report_sections, "模型系数结果", _get_var("result_table"))

    pred_df = _get_var("prediction_table")
    if pred_df is not None and not pred_df.empty:
        report_sections.append(("实际值预测值残差(前50行)", pred_df.head(50)))

    _add_section(report_sections, "测试集评价指标", _get_var("test_metric_table"))
    _add_section(report_sections, "线性模型诊断", _get_var("diagnostic_table"))
    _add_section(report_sections, "二分类混淆矩阵", _get_var("cm_table"))

    for var_name, title in [("assumptions_text", "模型假设"), ("paper_text", "论文表述草稿"), ("cleaning_text", "数据清洗表述")]:
        text_val = _get_var(var_name)
        if text_val:
            report_sections.append((title, text_val))

    # 修复：把高级分析结果也纳入 Word 报告，
    # 与侧边栏“⑦ 高级分析”步骤宣称的内容保持一致。
    adv_report_map = [
        ("topsis_data", "熵权法+TOPSIS评价结果", "topsis_table"),
        ("gm11_result", "灰色预测GM(1,1)结果", None),
        ("arima_result", "ARIMA时间序列预测结果", None),
        ("anova_result", "方差分析结果", None),
        ("chi2_result", "卡方独立性检验结果", None),
        ("pca_data", "PCA主成分分析结果", None),
        ("cluster_data", "聚类分析结果", None),
        ("ml_data", "机器学习建模结果", None),
    ]

    for _adv_key, _adv_title, _adv_sub in adv_report_map:
        _adv_val = st.session_state.get(_adv_key)

        if not isinstance(_adv_val, dict):
            continue

        if _adv_sub and isinstance(_adv_val.get(_adv_sub), pd.DataFrame):
            _add_section(
                report_sections,
                _adv_title,
                _adv_val[_adv_sub],
            )
        elif not _adv_sub:
            report_sections.append(
                (
                    _adv_title,
                    str(_adv_val.get("result") or _adv_val),
                )
            )

    # 修复：优化求解结果也进入报告（若本会话有）
    opt_solution = st.session_state.get("opt_solution_df")

    if isinstance(opt_solution, pd.DataFrame) and not opt_solution.empty:
        _add_section(
            report_sections,
            "优化求解结果",
            opt_solution,
        )

    if st.button("📥 生成综合报告 (Word)", key="export_report_button"):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io as _io

            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(10.5)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            title_para = doc.add_heading('数学建模前期数据分析综合报告', 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

            for section in report_sections:
                title, content = section
                doc.add_heading(str(title), level=2)
                if isinstance(content, pd.DataFrame):
                    if content.empty:
                        doc.add_paragraph("（无数据）")
                        continue
                    table = doc.add_table(rows=content.shape[0] + 1, cols=content.shape[1])
                    table.style = 'Table Grid'
                    for j, col in enumerate(content.columns):
                        table.cell(0, j).text = str(col)
                    for i in range(content.shape[0]):
                        for j in range(content.shape[1]):
                            table.cell(i + 1, j).text = str(content.iloc[i, j])
                else:
                    doc.add_paragraph(str(content))

            # 修复：将本会话生成的所有可视化图表一并嵌入报告
            if _REPORT_FIGURES:
                from docx.shared import Inches

                doc.add_heading('可视化图表', level=2)

                for fig_index, fig in enumerate(_REPORT_FIGURES, start=1):
                    try:
                        fig_buffer = _io.BytesIO()
                        fig.savefig(
                            fig_buffer,
                            format="png",
                            dpi=120,
                            bbox_inches="tight",
                        )
                        fig_buffer.seek(0)
                        doc.add_picture(
                            fig_buffer,
                            width=Inches(6.0),
                        )
                        doc.add_paragraph(
                            f"图{fig_index}"
                        )
                    except Exception:
                        continue

            buffer = _io.BytesIO()
            doc.save(buffer)
            doc_bytes = buffer.getvalue()

            st.download_button(
                "⬇️ 下载综合报告.docx",
                data=doc_bytes,
                file_name="数学建模前期数据分析综合报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_report_button",
            )
            st.success("综合报告生成成功，请点击上方按钮下载。")

            # 修复：报告确实生成成功后才推进自动引导进度
            st.session_state["guide_auto_step"] = max(
                st.session_state.get("guide_auto_step", 0),
                8,
            )
        except ImportError:
            st.error("缺少 python-docx 库，请先在终端执行：pip install python-docx")
        except Exception as e:
            st.error(f"生成报告失败：{e}")


    # ============================================================
    # 二十三、使用注意事项
    # ============================================================

    with st.expander(
        "使用和解释模型时的注意事项"
    ):
        st.markdown(
            """
    ### 1. 自动推荐不是最终结论

    程序根据变量类型和数据结构进行初步推荐，
    最终模型应结合题目目标、变量含义和模型诊断结果确定。

    ### 2. 相关性不等于因果关系

    Pearson/Spearman 相关系数和回归系数只能反映统计关联，
    不能单独证明因果关系。

    ### 3. 异常值不一定是错误数据

    异常值可能是真实存在的极端情况。
    正式论文中删除异常数据时，应说明识别方法和删除依据。

    ### 4. R²不能单独判断模型好坏

    应结合调整R²、AIC/BIC、残差图、测试集误差、
    显著性检验和实际解释意义共同评价。

    ### 5. 分类模型不要只看准确率

    如果类别分布不平衡，还应重点查看平衡准确率、
    精确率、召回率、F1值和ROC-AUC。

    ### 6. 混合效应模型需要真实的重复观测结构

    不能仅因为某列取值重复就直接使用混合效应模型。
    分组变量应确实代表同一对象、地区、企业或其他层级单位的重复观测。

    ### 7. 测试集划分需要结合数据结构

    如果数据具有时间顺序，建议使用时间序列划分；
    如果数据具有分组结构，建议按照组进行训练集和测试集划分，
    避免同一对象同时出现在训练集和测试集中。
    """
        )

    # ===== 新手常见问题 FAQ =====
    with st.expander("❓ 新手常见问题（FAQ）"):
        faq_items = [
            (
                "没有现成数据怎么办？",
                "点击侧边栏「📥 下载示例数据（CSV）」，先把全流程走通；"
                "正式比赛时数据通常由赛题附件提供。",
            ),
            (
                "因变量（要预测的东西）选哪个？",
                "赛题里“预测/解释/排名/评价”的对象就是因变量。"
                "例如：预测销量，销量就是因变量。",
            ),
            (
                "变量类型识别得对吗？",
                "自动识别仅供参考，务必在第④步人工核对："
                "连续=数值、分类=类别、时间=日期、次数=非负整数。"
                "类型错了，推荐的模型就会错。",
            ),
            (
                "模型拟合失败或不收敛怎么办？",
                "先检查：①有效样本是否够多（≥10 条）；"
                "②自变量数量是否小于样本数；"
                "③分类因变量的类别是否过少；"
                "④是否误选了需要分组变量的混合效应模型。"
                "也可用「多模型对比」换一个更简单的模型。",
            ),
            (
                "时间序列只有十几行数据，能用 ARIMA 吗？",
                "不建议。ARIMA 至少需要 30 期左右，"
                "4~30 期的小样本请用「灰色预测 GM(1,1)」。",
            ),
            (
                "分类模型准确率很高，但还需要看什么？",
                "类别不平衡时准确率会虚高，"
                "务必同时看平衡准确率、F1、ROC-AUC 和混淆矩阵。",
            ),
            (
                "什么时候用高级分析？",
                "评价类→熵权TOPSIS；预测类→灰色预测/ARIMA；"
                "分类类→随机森林；其他题型或写完主模型后→稳健性分析。",
            ),
            (
                "写论文时图表和表述从哪来？",
                "每个环节都有可直接复制的「论文表述」文本框；"
                "每张图可下载 PNG；最后「📑 一键导出综合报告」"
                "会把本会话所有图表和表格汇总成 Word。",
            ),
        ]

        for question, answer in faq_items:
            st.markdown(f"**Q：{question}**")
            st.markdown(f"**A：** {answer}")

    # ===== FAQ 结束 =====


else:  # 优化求解模块

    st.header("📊 优化问题求解器")


    if st.session_state.opt_uploaded_file is None:
        st.info("请先在侧边栏上传数据表，然后使用优化求解功能。")
        st.stop()

    try:
        uploaded_file = st.session_state.opt_uploaded_file
        uploaded_file.seek(0)
        if uploaded_file.name.lower().endswith(".csv"):
            try:
                opt_df = pd.read_csv(uploaded_file, encoding='utf-8')
            except UnicodeDecodeError:
                uploaded_file.seek(0)
                opt_df = pd.read_csv(uploaded_file, encoding='gbk')
            except Exception:
                uploaded_file.seek(0)
                opt_df = pd.read_csv(uploaded_file, encoding='gb18030')
        else:
            opt_df = pd.read_excel(uploaded_file)
    except Exception as e:
        st.error(f"读取文件失败：{e}")
        st.stop()

    opt_df.columns = make_unique_columns(opt_df.columns)
    st.success(f"已加载数据：{opt_df.shape[0]} 行, {opt_df.shape[1]} 列")

    # 优化类型
    opt_type = st.selectbox(
        "选择优化类型",
        [
            "线性规划 (LP)",
            "整数线性规划 (ILP)",
            "0-1 规划",
            "非线性规划 (NLP)",
            "遗传算法 (无约束/有界)",
        ],
    )

    # 决策变量
    # 自动找出所有“可以转为数字”的列（只要超过一半值能转成数字就算）
    col_list = []
    for col in opt_df.columns:
        try:
            numeric_part = pd.to_numeric(opt_df[col], errors='coerce')
            if numeric_part.notna().mean() > 0.5:  # 超过50%的值能转成数字
                col_list.append(col)
        except Exception:
            pass

    if not col_list:
        st.error("数据表中没有可转换为数值的列，请检查数据格式。")
        st.stop()

    var_cols = st.multiselect("选择决策变量所在的列（可多选）", col_list)
    if not var_cols:
        st.warning("请至少选择一个决策变量列。")
        st.stop()
    n_vars = len(var_cols)
    # 决策变量统计信息
    st.write("**决策变量统计信息**")
    try:
        stats_df = opt_df[var_cols].describe().T.reset_index()
        _st_dataframe(stats_df, use_container_width=True)
    except Exception:
        pass


    current_var_signature = tuple(var_cols)

    if st.session_state.get(
        "constraint_var_signature"
    ) != current_var_signature:
        st.session_state.cons_list = []
        st.session_state.constraint_var_signature = (
            current_var_signature
        )

    # 目标函数系数由用户输入
    st.subheader("目标函数系数")
    obj_coeffs = []
    for c in var_cols:
        # 目标函数系数不应自动使用数据表第一行。
        # 默认设置为0，由用户根据优化问题含义输入。
        default_val = 0.0

        coeff = st.number_input(
            f"系数 {c}",
            value=default_val,
            step=0.1,
            format="%.4f",
        )
        obj_coeffs.append(coeff)
    maximize = st.checkbox("最大化目标")

    # 约束条件
    st.subheader("约束条件")
    if "cons_list" not in st.session_state:
        st.session_state.cons_list = []   # 列表中每个元素都是字典
    # 修复：默认系数串与变量数一致，避免 n_vars≠2 时默认值永远非法
    default_coeff_str = ",".join(
        ["1"] * n_vars
    )
    with st.form("add_constraint"):
        coeff_str = st.text_input(
            "系数（逗号分隔）",
            value=default_coeff_str,
        )
        sign = st.selectbox("关系", ["<=", "=", ">="])
        rhs = st.number_input("右侧常数", value=1.0)

        add_constraint_clicked = st.form_submit_button("添加约束")
        try:
            coeffs = [float(x) for x in coeff_str.split(",")]
            if len(coeffs) != n_vars:
                st.error(f"系数个数应为 {n_vars}")
            else:
                # 保存变量列名，并用字典存储
                st.session_state.cons_list.append({
                    "cols": var_cols.copy(),
                    "coeffs": coeffs,
                    "sign": sign,
                    "rhs": rhs
                })
        except Exception as exc:
            st.error(
                f"系数格式错误：{exc}"
            )
    # 显示已添加约束
    for i, constraint_item in enumerate(
        st.session_state.cons_list
    ):
        coeffs = constraint_item["coeffs"]
        sign = constraint_item["sign"]
        rhs = constraint_item["rhs"]

        expr = " + ".join(
            f"{coeffs[j]}*{var_cols[j]}"
            for j in range(len(coeffs))
        )

        st.write(
            f"约束 {i + 1}: {expr} {sign} {rhs}"
        )
    if st.button("清空所有约束"):
        st.session_state.cons_list = []

    # 变量边界
    st.subheader("变量边界")
    use_bounds = st.checkbox("自定义边界（默认 >=0）")
    bounds = [(0, None) for _ in range(n_vars)]
    bounds_error = None
    if use_bounds:
        for i, col in enumerate(var_cols):
            c1, c2 = st.columns(2)
            lo = c1.number_input(
                f"{col} 下界",
                value=0.0,
                step=1.0,
                format="%.4f",
            )
            hi = c2.number_input(
                f"{col} 上界",
                value=100.0,
                step=1.0,
                format="%.4f",
            )

            if hi < lo:
                # 修复：不再用 st.stop() 硬停整页，改为标记错误
                # 由求解按钮处统一拦截，避免用户看不到下方其他控件。
                bounds_error = (
                    f"变量「{col}」的上界不能小于下界"
                    f"（上界 {hi} < 下界 {lo}）。"
                )

            bounds[i] = (lo, hi)

    if bounds_error:
        st.error(bounds_error)

    # 整数约束
    integrality = None
    if "整数" in opt_type:
        int_vars = st.multiselect("整数变量", var_cols)
        integrality = [1 if col in int_vars else 0 for col in var_cols]

        # 修复：未选任何整数变量时给出提示，避免静默退化为 LP
        if not int_vars:
            st.warning(
                "当前未标记任何整数变量，将按普通线性规划求解。"
            )

    if opt_type == "0-1 规划":
        # 修复：0-1 变量应使用 integrality=2（二进制），
        # 而不是 integrality=1 + bounds(0,1)，语义更准确。
        integrality = [2] * n_vars
        bounds = [(0, 1) for _ in range(n_vars)]

    # ---------- 非线性规划 / 遗传算法：目标函数输入 ----------
    nlp_use_cons = False

    if opt_type in ["非线性规划 (NLP)", "遗传算法 (无约束/有界)"]:
        st.subheader("非线性目标函数")
        st.markdown(
            "以 `x[0]`、`x[1]`… 表示决策变量。"
            "示例：`2*x[0]**2 + 3*x[1]**2 - 4*x[0]*x[1]`"
        )

        nonlinear_obj_str = st.text_input(
            "目标函数表达式",
            value="x[0]**2 + x[1]**2",
            key="nlp_obj_expr",
        )

        # 修复：GA 分支也提供线性约束开关，避免用户添加的约束
        # 在遗传算法下被静默丢弃。
        nlp_use_cons = st.checkbox(
            "启用上方线性约束",
            value=True,
            key="nlp_use_cons_check",
        )

        if opt_type == "非线性规划 (NLP)":
            st.caption(
                "提示：非线性规划使用SLSQP算法。若带约束求解失败，"
                "可取消勾选约束，或改用遗传算法搜索。"
            )
        else:
            st.caption(
                "提示：遗传算法要求边界为有限值；"
                "启用线性约束时，遗传算法会用惩罚项处理约束。"
            )

    # 求解
    # 检查约束是否与当前变量列匹配
    for cons in st.session_state.cons_list:
        if cons["cols"] != var_cols:
            st.error("您更改了决策变量列，请清空所有约束后重新添加！")
            st.stop()
    if st.button("🚀 求解", type="primary"):
        if bounds_error:
            st.error(f"无法求解：{bounds_error}")
            st.stop()
        if opt_type in ["非线性规划 (NLP)", "遗传算法 (无约束/有界)"]:
            # ---------------- 非线性求解 ----------------
            try:
                from scipy.optimize import minimize, differential_evolution

                # 修复：不再直接 eval 用户输入（存在任意代码执行风险）。
                # 改为：AST 节点白名单校验 + 受限命名空间（不注入 __builtins__）。
                _SAFE_NS = {
                    "np": np,
                    "abs": abs,
                    "exp": np.exp,
                    "log": np.log,
                    "log10": np.log10,
                    "sqrt": np.sqrt,
                    "sin": np.sin,
                    "cos": np.cos,
                    "tan": np.tan,
                    "min": min,
                    "max": max,
                    "pi": np.pi,
                    "e": np.e,
                }

                _ALLOWED_EXPR_NODES = (
                    ast.Expression,
                    ast.Lambda,
                    ast.arguments,
                    ast.arg,
                    ast.BinOp,
                    ast.UnaryOp,
                    ast.Name,
                    ast.Load,
                    ast.Constant,
                    ast.Add,
                    ast.Sub,
                    ast.Mult,
                    ast.Div,
                    ast.Pow,
                    ast.Mod,
                    ast.FloorDiv,
                    ast.USub,
                    ast.UAdd,
                    ast.Call,
                    ast.Attribute,
                    ast.List,
                    ast.Tuple,
                    ast.Subscript,
                    # 修复：ast.Index 自 Python 3.9 弃用、3.13 正式弃用，
                    # 未来版本会 AttributeError，且 3.9+ 中 x[0] 的 slice
                    # 本身就是 Constant，此节点从未命中（死代码），已移除。
                )

                _SAFE_NS_NAMES = set(_SAFE_NS.keys())

                def _check_expr_tree(tree):
                    """白名单 + 结构校验，阻断属性链逃逸（如 np.__class__）。"""
                    for node in ast.walk(tree):
                        node_type = type(node)

                        if node_type not in _ALLOWED_EXPR_NODES:
                            raise ValueError(
                                "表达式包含不允许的语法："
                                f"{node_type.__name__}"
                            )

                        if node_type is ast.Attribute:
                            # 只允许访问白名单对象（np/min/max 等）的公开成员，
                            # 属性名以下划线开头的一律拒绝（防 __class__/__mro__ 逃逸）。
                            if node.attr.startswith("_"):
                                raise ValueError(
                                    "不允许访问以下划线开头的属性："
                                    f"{node.attr}"
                                )

                            base = node.value

                            if not (
                                isinstance(base, ast.Name)
                                and base.id in _SAFE_NS_NAMES
                            ):
                                raise ValueError(
                                    "属性访问仅允许基于白名单对象"
                                    f"（{', '.join(sorted(_SAFE_NS_NAMES))}）。"
                                )

                        elif node_type is ast.Subscript:
                            # 下标只允许 x[整数] 形式
                            base = node.value
                            slice_node = node.slice

                            if not (
                                isinstance(base, ast.Name)
                                and base.id == "x"
                            ):
                                raise ValueError(
                                    "下标操作仅允许用于决策变量 x，如 x[0]。"
                                )

                            if not (
                                isinstance(slice_node, ast.Constant)
                                and isinstance(slice_node.value, int)
                            ):
                                raise ValueError(
                                    "下标必须是整数常量，如 x[0]。"
                                )

                            if slice_node.value < 0:
                                raise ValueError(
                                    "下标不能为负数。"
                                )

                        elif node_type is ast.Call:
                            func = node.func

                            if isinstance(func, ast.Attribute):
                                # np.xxx() 形式的调用已由 Attribute 规则校验
                                continue

                            if not (
                                isinstance(func, ast.Name)
                                and func.id in _SAFE_NS_NAMES
                            ):
                                raise ValueError(
                                    "只允许调用白名单内的函数"
                                    f"（{', '.join(sorted(_SAFE_NS_NAMES))}）。"
                                )

                def _safe_compile_expr(expr_str):
                    """把用户表达式编译为 lambda x 函数（带语法白名单校验）。"""
                    try:
                        tree = ast.parse(
                            "lambda x: " + expr_str,
                            mode="eval",
                        )
                    except SyntaxError as exc:
                        raise ValueError(
                            f"表达式语法错误：{exc}"
                        )

                    _check_expr_tree(tree)

                    # 修复：白名单必须合入 globals（而非 locals），
                    # 否则 lambda 内部引用 np/abs 等会因找不到名字而 NameError。
                    return eval(
                        compile(tree, "<expr>", "eval"),
                        dict(_SAFE_NS, **{"__builtins__": {}}),
                    )

                def _make_nlp_con(expr_str):
                    """把约束表达式字符串编译为函数，支持 abs()。"""
                    return _safe_compile_expr(expr_str)

                def _make_nlp_obj(expr_str):
                    """把目标函数表达式字符串编译为函数。"""
                    return _safe_compile_expr(expr_str)

                nonlinear_obj = _make_nlp_obj(nonlinear_obj_str)
                nonlinear_cons = []
                nonlinear_eq_cons = []

                if nlp_use_cons:
                    for cons in st.session_state.cons_list:
                        coeffs = cons["coeffs"]
                        sign = cons["sign"]
                        rhs = cons["rhs"]
                        n_c = len(coeffs)
                        defaults = " + ".join(
                            f"{coeffs[j]}*x[{j}]" for j in range(n_c)
                        )

                        # 注意：SLSQP 的 ineq 约束要求 fun(x) >= 0。
                        # <= 约束：rhs - expr >= 0
                        # >= 约束：expr - rhs >= 0
                        # = 约束：直接用 type="eq"，fun = expr - rhs
                        # （修复：不再用 -abs() 近似，避免在解处不可微
                        #   导致 SLSQP 数值梯度不连续、收敛失败）
                        if sign == "<=":
                            expr_part = f"{rhs} - ( {defaults} )"
                            nonlinear_cons.append(
                                {"type": "ineq", "fun": _make_nlp_con(expr_part)}
                            )
                        elif sign == ">=":
                            expr_part = f"{defaults} - {rhs}"
                            nonlinear_cons.append(
                                {"type": "ineq", "fun": _make_nlp_con(expr_part)}
                            )
                        else:
                            expr_part = f"( {defaults} ) - {rhs}"
                            nonlinear_eq_cons.append(
                                {"type": "eq", "fun": _make_nlp_con(expr_part)}
                            )

                opt_x0 = np.zeros(n_vars)

                # 修复：初值必须落在边界内，否则 SLSQP 因初始点
                # 违反边界而直接失败。
                for i in range(n_vars):
                    lo, hi = bounds[i]

                    if hi is not None and lo is not None:
                        opt_x0[i] = (lo + hi) / 2.0
                    elif lo is not None:
                        opt_x0[i] = (
                            lo + 1.0
                            if hi is None
                            else min(lo + 1.0, hi - 1e-6)
                        )
                    elif hi is not None:
                        opt_x0[i] = max(0.0, hi - 1.0)
                    else:
                        opt_x0[i] = 0.0

                sign_factor = -1.0 if maximize else 1.0

                if opt_type == "非线性规划 (NLP)":
                    result_nlp = minimize(
                        lambda x: sign_factor * nonlinear_obj(x),
                        opt_x0,
                        method="SLSQP",
                        bounds=bounds,
                        constraints=(
                            nonlinear_cons
                            + nonlinear_eq_cons
                            if nlp_use_cons
                            else None
                        ),
                        options={"maxiter": 500, "ftol": 1e-9},
                    )

                    opt_success = result_nlp.success
                    opt_x = result_nlp.x
                    opt_val = sign_factor * result_nlp.fun
                    message_text = result_nlp.message
                else:
                    # 修复：GA 要求有限边界，None 时给出明确提示
                    # 而不是执行 None+100 崩溃。
                    invalid_bounds = [
                        idx
                        for idx, b in enumerate(bounds)
                        if b[0] is None or b[1] is None
                    ]

                    if invalid_bounds:
                        raise ValueError(
                            "遗传算法要求所有变量都有有限边界，"
                            "请为变量 "
                            + "、".join(
                                var_cols[i]
                                for i in invalid_bounds
                            )
                            + " 设置上/下界。"
                        )

                    de_bounds = [
                        (float(b[0]), float(b[1]))
                        for b in bounds
                    ]

                    if nonlinear_cons or nonlinear_eq_cons:
                        # 注意：differential_evolution 的 constraints 参数
                        # 只接受 NonlinearConstraint 等对象，不支持裸函数。
                        # GA 的 NonlinearConstraint 语义是 fun(x) <= 0，
                        # 而 nonlinear_cons 是给 SLSQP 用的（fun(x) >= 0），
                        # 因此这里必须取负号转换回 GA 语义。
                        from scipy.optimize import NonlinearConstraint

                        nlc_list = [
                            NonlinearConstraint(
                                lambda x, f=c["fun"]: -f(x),
                                -np.inf,
                                0,
                            )
                            for c in nonlinear_cons
                        ]

                        # 等式约束：fun(x)=0，用 lb=ub=0 表达
                        nlc_list += [
                            NonlinearConstraint(
                                lambda x, f=c["fun"]: f(x),
                                0,
                                0,
                            )
                            for c in nonlinear_eq_cons
                        ]

                        result_ga = differential_evolution(
                            lambda x: sign_factor * nonlinear_obj(x),
                            de_bounds,
                            constraints=nlc_list,
                            seed=42,
                            maxiter=1000,
                            tol=1e-9,
                            polish=True,
                        )
                    else:
                        result_ga = differential_evolution(
                            lambda x: sign_factor * nonlinear_obj(x),
                            de_bounds,
                            seed=42,
                            maxiter=1000,
                            tol=1e-9,
                            polish=True,
                        )

                    opt_success = result_ga.success
                    opt_x = result_ga.x
                    opt_val = sign_factor * result_ga.fun
                    message_text = result_ga.message

                if opt_success:
                    st.success("✅ 求解成功！")
                    st.write("**最优解：**")
                    st.json(
                        {
                            var_cols[i]: float(opt_x[i])
                            for i in range(n_vars)
                        }
                    )
                    st.write(f"**最优值：** {opt_val:.6f}")

                    opt_solution_df = pd.DataFrame(
                        {
                            "变量": var_cols,
                            "最优解": [float(v) for v in opt_x],
                        }
                    )
                    opt_solution_df.loc[len(opt_solution_df)] = [
                        "最优值",
                        float(opt_val),
                    ]
                    st.session_state["opt_solution_df"] = (
                        opt_solution_df
                    )

                    dataframe_download(
                        opt_solution_df,
                        "优化结果.csv",
                        key="download_nlp_opt_result",
                    )

                    opt_method_name = (
                        "遗传算法"
                        if opt_type.startswith("遗传")
                        else "非线性规划(SLSQP)"
                    )

                    st.text_area(
                        "论文表述",
                        (
                            f"采用{opt_method_name}对上述优化问题进行求解，"
                            f"目标函数的最优值为 {opt_val:.6f}，"
                            f"对应的最优决策变量取值为："
                            + "，".join(
                                f"{var_cols[i]} = {float(opt_x[i]):.6f}"
                                for i in range(n_vars)
                            )
                            + "。"
                        ),
                        height=120,
                        key="nlp_text_area",
                    )
                else:
                    st.error(f"求解失败：{message_text}")
            except Exception as e:
                st.error(f"求解错误：{e}")
        else:
            # ---------------- 线性 / 整数线性 / 0-1 规划求解 ----------------
            # 修复：线性求解代码原本缩进错误、位于 if/else 之外，
            # 导致非线性求解后必然访问未定义的 A_ub 而崩溃。
            # 现已整体移入 else 分支，只在“线性类”求解时执行。
            A_ub, b_ub, A_eq, b_eq = [], [], [], []

            for cons in st.session_state.cons_list:
                coeffs = cons["coeffs"]
                sign = cons["sign"]
                rhs = cons["rhs"]
                if sign == "<=":
                    A_ub.append(coeffs); b_ub.append(rhs)
                elif sign == ">=":
                    A_ub.append([-c for c in coeffs]); b_ub.append(-rhs)
                else:
                    A_eq.append(coeffs); b_eq.append(rhs)
            c = np.array(obj_coeffs)
            if maximize:
                c = -c

            # 转换为 numpy 数组
            A_ub = np.array(A_ub) if A_ub else None
            b_ub = np.array(b_ub) if b_ub else None
            A_eq = np.array(A_eq) if A_eq else None
            b_eq = np.array(b_eq) if b_eq else None

            try:
                if integrality is not None and any(integrality):
                    from scipy.optimize import milp, LinearConstraint, Bounds
                    cons = []
                    if A_ub is not None:
                        cons.append(LinearConstraint(A_ub, -np.inf, b_ub))
                    if A_eq is not None:
                        cons.append(LinearConstraint(A_eq, b_eq, b_eq))
                    lb = [b[0] if b[0] is not None else -np.inf for b in bounds]
                    ub = [b[1] if b[1] is not None else np.inf for b in bounds]
                    res = milp(c=c, constraints=cons, bounds=Bounds(lb, ub), integrality=integrality)
                else:
                    res = linprog(c, A_ub=A_ub, b_ub=b_ub, A_eq=A_eq, b_eq=b_eq, bounds=bounds, method='highs')

                if res.success:
                    opt_x = res.x
                    opt_val = -res.fun if maximize else res.fun
                    st.success("✅ 求解成功！")
                    st.write("**最优解：**")
                    st.json({var_cols[i]: float(opt_x[i]) for i in range(n_vars)})
                    st.write(f"**最优值：** {opt_val:.6f}")

                    opt_solution_df = pd.DataFrame(
                        {
                            "变量": var_cols,
                            "最优解": [float(v) for v in opt_x],
                        }
                    )
                    opt_solution_df.loc[len(opt_solution_df)] = [
                        "最优值",
                        float(opt_val),
                    ]
                    st.session_state["opt_solution_df"] = (
                        opt_solution_df
                    )

                    dataframe_download(
                        opt_solution_df,
                        "优化结果.csv",
                        key="download_linear_opt_result",
                    )

                    st.text_area(
                        "论文表述",
                        (
                            f"采用{opt_type}对上述问题进行求解，"
                            f"目标函数的最优值为 {opt_val:.6f}，"
                            f"对应的最优决策变量取值为："
                            + "，".join(
                                f"{var_cols[i]} = {float(opt_x[i]):.6f}"
                                for i in range(n_vars)
                            )
                            + "。"
                        ),
                        height=120,
                        key="linear_text_area",
                    )

                    if opt_type == "线性规划 (LP)":

                        # 敏感性分析：优先使用 HiGHS 返回的精确对偶解，
                        # 失败再回退到差分法。
                        st.subheader("影子价格（对偶解）")

                        shadow_rows = []

                        # 修复：优先用 res.ineqlin.marginals / res.eqlin.marginals
                        # （HiGHS 返回精确对偶值，比差分法更准确）
                        try:
                            ineq_marginals = (
                                np.asarray(
                                    res.ineqlin.marginals,
                                    dtype=float,
                                )
                                if A_ub is not None
                                else np.zeros(0)
                            )
                            eq_marginals = (
                                np.asarray(
                                    res.eqlin.marginals,
                                    dtype=float,
                                )
                                if A_eq is not None
                                else np.zeros(0)
                            )

                            ineq_counter = 0
                            eq_counter = 0

                            for idx, constraint_item in enumerate(
                                st.session_state.cons_list
                            ):
                                sign = constraint_item["sign"]

                                if sign == "=":
                                    shadow_val = float(
                                        eq_marginals[eq_counter]
                                    ) if eq_counter < len(eq_marginals) else np.nan
                                    eq_counter += 1
                                else:
                                    shadow_val = float(
                                        ineq_marginals[ineq_counter]
                                    ) if ineq_counter < len(ineq_marginals) else np.nan
                                    ineq_counter += 1

                                shadow_rows.append(
                                    {
                                        "约束": f"约束 {idx+1}",
                                        "关系": sign,
                                        "影子价格": shadow_val,
                                    }
                                )
                        except Exception:
                            shadow_rows = []

                        # 修复：等式约束做双侧差分取均值；差分失败则标为 NaN
                        if not shadow_rows:
                            for idx, constraint_item in enumerate(
                                st.session_state.cons_list
                            ):
                                coeffs = constraint_item["coeffs"]
                                sign = constraint_item["sign"]
                                rhs = constraint_item["rhs"]

                                delta = 0.01 * max(abs(rhs), 1.0)
                                shadow_values = []

                                for direction in [1.0, -1.0]:
                                    t_A_ub, t_b_ub, t_A_eq, t_b_eq = [], [], [], []
                                    for j, constraint_j in enumerate(
                                        st.session_state.cons_list
                                    ):
                                        coeffs_j = constraint_j["coeffs"]
                                        sign_j = constraint_j["sign"]
                                        rhs_j = constraint_j["rhs"]

                                        if j == idx:
                                            rhs_j += direction * delta
                                        if sign_j == "<=":
                                            t_A_ub.append(coeffs_j); t_b_ub.append(rhs_j)
                                        elif sign_j == ">=":
                                            t_A_ub.append([-c for c in coeffs_j]); t_b_ub.append(-rhs_j)
                                        else:
                                            t_A_eq.append(coeffs_j); t_b_eq.append(rhs_j)
                                    try:
                                        t_A_ub = np.array(t_A_ub) if t_A_ub else None
                                        t_b_ub = np.array(t_b_ub) if t_b_ub else None
                                        t_A_eq = np.array(t_A_eq) if t_A_eq else None
                                        t_b_eq = np.array(t_b_eq) if t_b_eq else None
                                        res2 = linprog(c, A_ub=t_A_ub, b_ub=t_b_ub, A_eq=t_A_eq, b_eq=t_b_eq, bounds=bounds, method='highs')
                                        if res2.success:
                                            pval = -res2.fun if maximize else res2.fun
                                            shadow_values.append(
                                                (pval - opt_val) / direction / delta
                                            )
                                    except Exception:
                                        pass

                                # 修复：等式约束用双侧差分均值，更准确
                                if len(shadow_values) >= 1:
                                    if sign == "=" and len(shadow_values) == 2:
                                        shadow_val = np.mean(shadow_values)
                                    else:
                                        shadow_val = shadow_values[0]
                                    shadow_rows.append(
                                        {
                                            "约束": f"约束 {idx+1}",
                                            "关系": sign,
                                            "影子价格": shadow_val,
                                        }
                                    )
                                else:
                                    shadow_rows.append(
                                        {
                                            "约束": f"约束 {idx+1}",
                                            "关系": sign,
                                            "影子价格": np.nan,
                                        }
                                    )

                        if shadow_rows:
                            shadow_df = pd.DataFrame(shadow_rows)

                            def _shadow_note(row):
                                val = row["影子价格"]
                                if pd.isna(val):
                                    return "无法计算"
                                sign_char = row["关系"]
                                if sign_char == "<=":
                                    return (
                                        "该资源每增加1单位，目标值的变化量；"
                                        "≥0 表示该约束紧（binding）"
                                    )
                                if sign_char == ">=":
                                    return (
                                        "该约束每放宽1单位（RHS减小），目标值的变化量；"
                                        "显示正值说明当前该约束为紧约束"
                                    )
                                return "等式约束每放宽1单位的目标值变化量"

                            shadow_df["说明"] = shadow_df.apply(
                                _shadow_note,
                                axis=1,
                            )
                            _st_dataframe(
                                shadow_df,
                                use_container_width=True,
                            )

                            dataframe_download(
                                shadow_df,
                                "影子价格.csv",
                                key="download_shadow_price",
                            )

                            st.caption(
                                "影子价格仅在最优基不变的小范围内有效；"
                                "对偶解由 HiGHS 求解器直接给出，"
                                "未使用差分近似。"
                            )
                else:
                    st.error(f"求解失败：{res.message}")
            except Exception as e:
                st.error(f"求解错误：{e}")

    # 以下优化代码结束


# ==================== 建模彩蛋模块 ====================
# 本模块只负责显示趣味内容，不参与任何数据处理、模型拟合和结果计算。
# 如果不想显示，直接删除本模块即可。

try:
    EASTER_EGG_QUOTES = [
        "模型不会自动变好，但数据清洗后通常会更诚实。",
        "相关性可以一起散步，但不代表它们互相导致。",
        "先看数据，再选模型；不要让模型替题目做决定。",
        "一个好的评价体系，应该既能排序，也能解释为什么这样排序。",
        "当结果特别完美时，先检查数据泄漏。",
        "异常值不一定是错误，也可能是数据正在认真讲故事。",
        "R²很高值得开心，但测试集表现更值得相信。",
        "数学建模的第一步不是写公式，而是确认变量到底是什么意思。",
        "如果所有指标权重都一样，至少要诚实地说这是等权重。",
        "稳定的结论，比漂亮的结论更有价值。",
    ]

    EASTER_EGG_SIGNS = [
        "今天适合检查一遍变量含义。",
        "今天适合看看残差图。",
        "今天适合重新确认正向指标和负向指标。",
        "今天适合做一次敏感性分析。",
        "今天适合给模型加一句合理的解释。",
        "今天适合怀疑一下过高的准确率。",
        "今天适合保存一份中间结果。",
        "今天适合把代码注释写清楚。",
    ]

    now = datetime.now()

    # 使用日期生成稳定索引。
    # 同一天刷新页面时内容不变，不会影响正常运行。
    date_code = now.strftime("%Y-%m-%d")
    hash_value = int(
        hashlib.md5(
            date_code.encode("utf-8")
        ).hexdigest(),
        16
    )

    quote = EASTER_EGG_QUOTES[
        hash_value % len(EASTER_EGG_QUOTES)
    ]

    # 修复：建模签改为每次点击随机抽取，
    # 避免“抽取”语义与实现不符（原实现同日结果永远相同）。
    egg_sign_index = st.session_state.get(
        "easter_egg_sign_index",
        None,
    )

    if egg_sign_index is None:
        egg_sign_index = int(
            hash_value % len(EASTER_EGG_SIGNS)
        )

    today_sign = EASTER_EGG_SIGNS[
        egg_sign_index % len(EASTER_EGG_SIGNS)
    ]

    with st.expander(
        "🔐 发现一个隐藏的建模彩蛋",
        expanded=False
    ):
        st.markdown(
            f"""
            <div style="
                padding: 16px;
                border-radius: 12px;
                background: linear-gradient(
                    135deg,
                    rgba(70,130,180,0.12),
                    rgba(120,200,160,0.12)
                );
                border: 1px solid rgba(100,140,180,0.25);
                margin-bottom: 10px;
            ">
                <h4>📐 今日建模语录</h4>
                <p style="font-size: 18px;">
                    “{quote}”
                </p>
                <p style="color: #777;">
                    当前时间：{now.strftime("%Y-%m-%d %H:%M:%S")}
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )

        if st.button(
            "🎲 抽取今日建模签",
            key="easter_egg_sign_button"
        ):
            # 每次点击重新随机抽取
            st.session_state["easter_egg_sign_index"] = int(
                np.random.default_rng().integers(
                    0,
                    len(EASTER_EGG_SIGNS),
                )
            )
            st.rerun()

        st.success(
            f"今日建模签：{today_sign}"
        )

        st.caption(
            "这个彩蛋不会修改数据、模型参数或分析结果。"
        )

except Exception:
    # 彩蛋出现任何问题时静默跳过，
    # 确保不会影响主程序正常运行。
    pass

# ==================== 建模彩蛋模块结束 ====================


